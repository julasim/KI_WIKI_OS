#!/usr/bin/env python3
"""
ki_wiki_bot.py — Telegram-Bot für KI_WIKI_Vault.

Single-file, single-user. Vault-Operations via LLM-Tool-Use.
Voice → faster-whisper (lokal). Photo → Vision-LLM. URL → trafilatura.

ENV:
  TG_TOKEN, ALLOWED_USER_ID, OPENROUTER_API_KEY,
  VAULT_PATH (default /vault),
  LLM_MODEL, VISION_MODEL,
  WHISPER_MODEL (small), WHISPER_DEVICE (cpu), WHISPER_LANG (de)
"""

import os
import re
import json
import time
import threading
import base64
import shutil
import asyncio
import logging
import subprocess
import tempfile
from datetime import date, datetime, time as dtime, timedelta
from pathlib import Path
from typing import Optional
from zoneinfo import ZoneInfo

import frontmatter
import trafilatura
from openai import OpenAI
from faster_whisper import WhisperModel
from telegram import Update, constants
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
)

# ============================================================================
# Config
# ============================================================================

VAULT = Path(os.environ.get("VAULT_PATH", "/vault"))
# 0 = Setup-Modus: Bot meldet beim ersten Kontakt die User-ID,
#                  user editiert .env und startet neu.
ALLOWED_USER_ID = int(os.environ.get("ALLOWED_USER_ID", "0") or "0")
TG_TOKEN = os.environ["TG_TOKEN"]

# LLM-Provider — beliebige OpenAI-API-kompatible Endpoints:
#   OpenRouter      → https://openrouter.ai/api/v1
#   Ollama Cloud    → https://ollama.com/v1
#   OpenAI direkt   → https://api.openai.com/v1
#   Lokales Ollama  → http://ollama:11434/v1
# OPENROUTER_API_KEY bleibt als Fallback für bestehende Installs.
LLM_API_KEY = os.environ.get("LLM_API_KEY") or os.environ.get("OPENROUTER_API_KEY")
if not LLM_API_KEY:
    raise RuntimeError("LLM_API_KEY (oder OPENROUTER_API_KEY) muss gesetzt sein.")
LLM_BASE_URL = os.environ.get("LLM_BASE_URL", "https://openrouter.ai/api/v1")
LLM_MODEL = os.environ.get("LLM_MODEL", "anthropic/claude-sonnet-4-5")
# Provider-Detection für Format-Kompatibilität:
# - Anthropic-direkt + OpenRouter→Anthropic: cache_control + content-as-list erlaubt
# - Gemini OpenAI-Compat + reine OpenAI: content muss String sein, kein cache_control
# - Ollama: content muss String sein
_LLM_BASE_LOWER = LLM_BASE_URL.lower()
_LLM_MODEL_LOWER = LLM_MODEL.lower()
USE_ANTHROPIC_CACHE = (
    "anthropic" in _LLM_BASE_LOWER
    or ("openrouter" in _LLM_BASE_LOWER and _LLM_MODEL_LOWER.startswith("anthropic/"))
)
# Vision kann optional ein anderer Provider sein (z.B. wenn Haupt-LLM kein Vision kann)
VISION_API_KEY = os.environ.get("VISION_API_KEY", LLM_API_KEY)
VISION_BASE_URL = os.environ.get("VISION_BASE_URL", LLM_BASE_URL)
VISION_MODEL = os.environ.get("VISION_MODEL", LLM_MODEL)
WHISPER_MODEL_SIZE = os.environ.get("WHISPER_MODEL", "small")
WHISPER_DEVICE = os.environ.get("WHISPER_DEVICE", "cpu")
WHISPER_LANG = os.environ.get("WHISPER_LANG", "de")

# Daily-Briefing
try:
    BRIEFING_HOUR = int(os.environ.get("BRIEFING_HOUR", "0") or "0")
except ValueError:
    BRIEFING_HOUR = 0
try:
    SUGGESTION_HOUR = int(os.environ.get("SUGGESTION_HOUR", "0") or "0")
except ValueError:
    SUGGESTION_HOUR = 0
TIMEZONE = ZoneInfo(os.environ.get("TIMEZONE", "Europe/Vienna"))

TG_MAX_MESSAGE = 3800

LIFE = VAULT / "10_Life"
DAILY_DIR = LIFE / "daily"
TASKS_DIR = LIFE / "tasks"
NOTES_DIR = LIFE / "notes"
MEETINGS_DIR = LIFE / "meetings"
AREAS_DIR = LIFE / "areas"
PROJECTS_DIR = VAULT / "05_Projects"  # Schema-konform: Projekte sind eigener Top-Level
TEMPLATES_DIR = VAULT / "08_Templates"
ATTACHMENTS_DIR = VAULT / "09_Attachments"
RAW_ARTICLES_DIR = VAULT / "01_Raw" / "articles"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
log = logging.getLogger("ki_wiki_bot")


# ─── Log-Sanitization: TG_TOKEN aus HTTPX-INFO-Logs maskieren ───────────────
# python-telegram-bot embeddet den Token in die HTTP-URL. httpx loggt URLs
# bei INFO-Level → Token leakt in Log-Files / Container-Output / Bug-Reports.
# Filter wandelt /bot<TOKEN>/method → /bot[REDACTED]/method.
class _TokenMaskingFilter(logging.Filter):
    _TG_TOKEN_RE = re.compile(r"/bot\d{7,15}:[A-Za-z0-9_\-]{30,}/")

    def filter(self, record: logging.LogRecord) -> bool:
        try:
            msg = record.getMessage()
            if "/bot" in msg:
                masked = self._TG_TOKEN_RE.sub("/bot[REDACTED]/", msg)
                if masked != msg:
                    record.msg = masked
                    record.args = ()
        except Exception:
            pass
        return True


# Auf root-Handler hängen → fängt alle Logger inkl. httpx + telegram.ext
for _h in logging.getLogger().handlers:
    _h.addFilter(_TokenMaskingFilter())

# ============================================================================
# LLM client (OpenRouter, OpenAI-API-kompatibel)
# ============================================================================

llm = OpenAI(
    base_url=LLM_BASE_URL,
    api_key=LLM_API_KEY,
    default_headers={
        "HTTP-Referer": "https://github.com/julasim/KI_WIKI_OS",
        "X-Title": "KI Wiki Bot",
    },
)
# Separater Vision-Client falls Vision via anderen Provider läuft
vision_llm = (
    llm if (VISION_BASE_URL == LLM_BASE_URL and VISION_API_KEY == LLM_API_KEY)
    else OpenAI(base_url=VISION_BASE_URL, api_key=VISION_API_KEY)
)

# ============================================================================
# Whisper (einmal beim Start laden, im Speicher halten)
# ============================================================================

log.info(f"Loading Whisper '{WHISPER_MODEL_SIZE}' on {WHISPER_DEVICE}...")
whisper = WhisperModel(WHISPER_MODEL_SIZE, device=WHISPER_DEVICE, compute_type="int8")
log.info("Whisper loaded.")

# ============================================================================
# Helpers
# ============================================================================

def slugify(s: str, max_len: int = 50) -> str:
    """kebab-case slug, deutsche Umlaute behandelt."""
    s = s.lower().strip()
    s = (s.replace("ä", "ae").replace("ö", "oe")
         .replace("ü", "ue").replace("ß", "ss"))
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s[:max_len] or "untitled"


def safe_path(rel_path: str) -> Path:
    """Resolve rel_path inside VAULT, prevent traversal.

    Nutzt is_relative_to (Python 3.9+) — startswith() ist anfällig für
    Prefix-Bug: VAULT='/x/vault' würde '/x/vault_evil/secret' akzeptieren.
    Vault selbst ist erlaubt (Edge-Case bei rel_path='').
    """
    vault_resolved = VAULT.resolve()
    p = (VAULT / rel_path).resolve()
    try:
        # is_relative_to True wenn p == vault_resolved oder echter Subpath
        if not p.is_relative_to(vault_resolved):
            raise ValueError(f"Path traversal: {rel_path}")
    except AttributeError:
        # Fallback für Python < 3.9 — separator-aware prefix-check
        v_str = str(vault_resolved)
        p_str = str(p)
        if p_str != v_str and not p_str.startswith(v_str + os.sep):
            raise ValueError(f"Path traversal: {rel_path}")
    return p


def iter_vault_md(root: Optional[Path] = None,
                  recursive: bool = True,
                  skip_noise: bool = True):
    """Generator über .md-Files mit geladenem Frontmatter.

    Konsolidiert das Vault-Walk-Pattern aus 5+ Stellen — vorher überall
    selbst implementiert, mit leicht unterschiedlichen Skip-Listen und
    Error-Handling. Jetzt einheitlich.

    root: Wurzel — Default VAULT (ganzes Vault). Bei z.B. TASKS_DIR genügt
          recursive=False (flache Liste).
    recursive: True → rglob, False → glob.
    skip_noise: skippt VAULT_NOISE_DIRS (Templates/Tools/Meta/Archive)
                UND VAULT_NOISE_FILE_NAMES (System-Doku, _index.md, CONTEXT.md).
                Nur relevant wenn Pfade unter VAULT liegen.

    Yields: (path, frontmatter.Post)-Tuples. Files mit kaputtem YAML werden
            geloggt + übersprungen — kein Crash.
    """
    if root is None:
        root = VAULT
    iter_method = root.rglob if recursive else root.glob
    for path in iter_method("*.md"):
        if skip_noise:
            # System-Files überall skippen (CLAUDE/MOC/_index/CONTEXT etc.)
            # — diese sind absichtlich ohne Frontmatter bzw. Doku-Files mit
            # Wikilink-Beispielen die als "broken link" missverstanden würden.
            # Project-READMEs werden hier NICHT geskippt (sind User-Content
            # mit Frontmatter type:project).
            if _is_system_file(path):
                continue
            try:
                rel = path.relative_to(VAULT)
                if rel.parts and rel.parts[0] in VAULT_NOISE_DIRS:
                    continue
            except ValueError:
                # Pfad nicht unter VAULT (sollte nicht vorkommen) → skippen
                continue
        try:
            post = frontmatter.load(path)
        except Exception as e:
            log.warning(f"frontmatter parse failed: {path.name}: {e}")
            continue
        yield path, post


def find_project_dir(slug: str) -> Optional[Path]:
    """Findet einen Projekt-Ordner via rekursiver Suche unter 05_Projects/.

    Erlaubt beliebige Verschachtelung (Subprojekte). Gibt None bei Nicht-Treffer
    oder Mehrdeutigkeit zurück (in dem Fall sollte Caller einen Fehler werfen).
    Slug wird vorher normalisiert ('project-' Präfix entfernt, lowercase).
    """
    if not slug:
        return None
    slug = slug.strip().lower()
    if slug.startswith("project-"):
        slug = slug[len("project-"):]
    if not PROJECTS_DIR.exists():
        return None
    matches = [d for d in PROJECTS_DIR.rglob(slug) if d.is_dir() and d.name == slug]
    if len(matches) == 1:
        return matches[0]
    return None  # 0 Treffer oder mehrdeutig


def list_all_project_slugs() -> list[tuple[str, str]]:
    """Liefert (slug, relative_path)-Liste aller Projekte (rekursiv).

    Brauchen wir für Mehrdeutigkeits-Diagnose und List-Outputs.
    """
    if not PROJECTS_DIR.exists():
        return []
    out = []
    for d in sorted(PROJECTS_DIR.rglob("*")):
        if not d.is_dir():
            continue
        # Nur Projekt-Dirs zählen — die haben README.md mit type: project
        readme = d / "README.md"
        if not readme.exists():
            continue
        try:
            post = frontmatter.load(readme)
            if post.metadata.get("type") == "project":
                rel = d.relative_to(VAULT).as_posix()
                out.append((d.name, rel))
        except Exception:
            continue
    return out


def atomic_write(path: Path, content: str) -> None:
    """Write atomic: tmp + rename."""
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(content, encoding="utf-8")
    os.replace(tmp, path)


def today_iso() -> str:
    """Heute als ISO-String in BOT-Lokalzeit (Europe/Vienna o.ä.).

    KRITISCH: nutzt TIMEZONE statt date.today(). Container läuft typisch
    UTC — `date.today()` liefert nach 22:00 Wien-Zeit schon morgen.
    Folge wäre: Tasks landen in falscher Daily, recurring Tasks
    reaktivieren nie (last_completed = morgen → _is_recurrence_due False).
    """
    return datetime.now(TIMEZONE).date().isoformat()


def load_template(name: str) -> str:
    """Read template file (e.g. 'daily' → daily_template.md)."""
    return (TEMPLATES_DIR / f"{name}_template.md").read_text(encoding="utf-8")


def ocr_image(image_path: Path) -> str:
    """OCR via Tesseract (deutsch+englisch). Leerer String wenn nichts erkennbar."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        return ""
    try:
        img = Image.open(str(image_path))
        text = pytesseract.image_to_string(img, lang="deu+eng").strip()
        return text
    except Exception as e:
        log.warning(f"OCR failed for {image_path}: {e}")
        return ""


def extract_docx_text(docx_path: Path) -> tuple:
    """Extrahiert Text + Tabellen + Metadaten aus .docx via python-docx.

    Tabellen werden als Markdown-Tabellen gerendert (für Volltext-Suche
    + Lesbarkeit im .md-Wrapper).

    Returns: (text, metadata_dict, paragraph_count)
    metadata_dict hat title, author, subject (alle optional).
    """
    try:
        import docx  # type: ignore  # python-docx
    except ImportError:
        return "(python-docx nicht installiert)", {}, 0

    try:
        doc = docx.Document(str(docx_path))

        # Body: Paragraphs + Tabellen IN REIHENFOLGE der Doc-Struktur.
        # docx hat doc.paragraphs und doc.tables separat — Reihenfolge geht
        # verloren wenn man sie einzeln iteriert. Lösung: über doc.element.body.
        from docx.oxml.ns import qn  # type: ignore
        text_parts = []
        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):  # Paragraph
                # Text-Inhalt sammeln (alle text-runs)
                runs = child.findall(".//" + qn("w:t"))
                line = "".join(r.text or "" for r in runs).strip()
                if line:
                    text_parts.append(line)
            elif tag == qn("w:tbl"):  # Tabelle
                # Markdown-Tabelle bauen
                rows_text = []
                for row in child.findall(qn("w:tr")):
                    cells = []
                    for cell in row.findall(qn("w:tc")):
                        cell_runs = cell.findall(".//" + qn("w:t"))
                        cell_text = "".join(r.text or "" for r in cell_runs).strip()
                        # Pipes in Zellen escapen
                        cells.append(cell_text.replace("|", "\\|").replace("\n", " "))
                    rows_text.append(cells)
                if rows_text:
                    n_cols = max(len(r) for r in rows_text)
                    rows_text = [r + [""] * (n_cols - len(r)) for r in rows_text]
                    md_table = ["| " + " | ".join(rows_text[0]) + " |",
                                "|" + "|".join(["---"] * n_cols) + "|"]
                    for r in rows_text[1:]:
                        md_table.append("| " + " | ".join(r) + " |")
                    text_parts.append("\n" + "\n".join(md_table) + "\n")

        # Core-Metadaten
        cp = doc.core_properties
        meta = {
            "title": (cp.title or "").strip(),
            "author": (cp.author or "").strip(),
            "subject": (cp.subject or "").strip(),
        }
        para_count = sum(1 for c in body.iterchildren() if c.tag == qn("w:p"))

        text = "\n\n".join(text_parts).strip()
        return (text or "(kein Text extrahiert)", meta, para_count)
    except Exception as e:
        log.exception(f"DOCX extract failed for {docx_path}")
        return f"(Extraktions-Fehler: {e})", {}, 0


def extract_pdf_text(pdf_path: Path, max_pages: int = 200) -> tuple:
    """Extrahiert Text + Metadaten aus PDF via pymupdf.

    Returns: (text, metadata_dict, total_pages)
    metadata_dict hat title, author, subject, keywords (alle optional).
    Bei sehr großen PDFs werden nur die ersten max_pages extrahiert.
    """
    try:
        import pymupdf  # type: ignore
    except ImportError:
        return "(pymupdf nicht installiert)", {}, 0

    try:
        doc = pymupdf.open(str(pdf_path))
        total = doc.page_count
        pages_to_read = min(total, max_pages)

        text_parts = []
        for i in range(pages_to_read):
            page_text = doc[i].get_text("text")
            if page_text.strip():
                text_parts.append(f"\n## Seite {i+1}\n\n{page_text.strip()}")

        if total > max_pages:
            text_parts.append(
                f"\n\n_(PDF hat {total} Seiten, nur die ersten {max_pages} extrahiert)_"
            )

        meta_raw = doc.metadata or {}
        meta = {k: (meta_raw.get(k) or "").strip() for k in ("title", "author", "subject", "keywords")}
        doc.close()

        return ("\n".join(text_parts).strip() or "(kein Text extrahiert)", meta, total)
    except Exception as e:
        log.exception(f"PDF extract failed for {pdf_path}")
        return f"(Extraktions-Fehler: {e})", {}, 0


# ============================================================================
# Tool implementations
# ============================================================================

def ensure_daily() -> Path:
    """Create today's daily file if missing, return path.

    TOCTOU-safe: atomic create-or-skip via os.O_EXCL. Verhindert dass
    parallele Aufrufe (z.B. recurring_task_reset_job + User-Message
    gleichzeitig) die Daily mit Template-Inhalt überschreiben nachdem
    der erste Caller schon edits gemacht hat.
    """
    today = today_iso()
    path = DAILY_DIR / f"{today}.md"
    if path.exists():
        return path
    template = load_template("daily")
    content = template.replace("{{date:YYYY-MM-DD}}", today)
    post = frontmatter.loads(content)
    post["id"] = f"daily-{today}"
    post["title"] = today
    body = frontmatter.dumps(post) + "\n"
    DAILY_DIR.mkdir(parents=True, exist_ok=True)
    # O_EXCL: atomic create-only. Wenn ein anderer Task schon angelegt
    # hat → FileExistsError → wir nutzen den bestehenden File.
    try:
        fd = os.open(str(path), os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o644)
        try:
            os.write(fd, body.encode("utf-8"))
        finally:
            os.close(fd)
        log.info(f"Created daily: {path.name}")
    except FileExistsError:
        log.debug(f"daily {path.name} bereits angelegt (parallel race) — nutze existierendes")
    return path


VALID_SECTIONS = {"Heute", "Notizen & Gedanken", "Offen / Einsortieren", "Abends"}

# Platzhalter aus dem daily_template, die beim ersten Append entfernt werden sollen
EMPTY_PLACEHOLDERS = {"- [ ]", "- [x]", "-", "•", "- Was lief gut?", "- Was nehme ich mit?"}


def _clean_section_body(body: str) -> str:
    """Entfernt nur-Platzhalter-Zeilen wenn die Sektion sonst leer wäre."""
    lines = body.split("\n")
    real_content = [l for l in lines if l.strip() and l.strip() not in EMPTY_PLACEHOLDERS]
    if not real_content:
        return ""  # Sektion war komplett leer (nur Template-Platzhalter)
    return "\n".join(l for l in lines if l.strip())  # Auch leere Zeilen weg


# ─── Auto-Linking ───────────────────────────────────────────────────────────
# Wandelt erkannte Vault-IDs/Titles im Schreiber-Output in [[wikilinks]] um.
# Conservative-by-default: nur exakte Matches, Stop-Wörter ausgenommen, Code/
# bestehende Links unangetastet, mit Cache und expliziter Invalidierung.

# Cache: (timestamp, phrase_lower → canonical_id)
_AUTO_LINK_CACHE: tuple = (0.0, {})
_AUTO_LINK_TTL_SEC = 300  # 5 Min — bei Schreibvorgang explizit invalidiert
_AUTO_LINK_LOCK = threading.Lock()

# Verzeichnisse die NIE für Auto-Link/Listings/Indexes gescannt werden
# (Templates/Tools/Memory/Archive würden False-Positives erzeugen und
# nichts Sinnvolles beitragen). Genutzt von iter_vault_md() + auto_link.
VAULT_NOISE_DIRS = {
    ".obsidian", ".trash", "99_Archive",
    "06_Meta", "07_Tools", "08_Templates",
}

# Doku-Files im Vault-Root die KEIN Frontmatter haben sollen.
_VAULT_ROOT_SYSTEM_DOCS = {
    "README.md", "MOC.md", "CLAUDE.md", "SCHEMA.md", "COMMANDS.md", "PIPELINES.md",
}


def _is_system_file(path: Path) -> bool:
    """True wenn Datei System/Doku/Auto-Generated und nicht User-Content.

    Wird von iter_vault_md(skip_noise=True) genutzt.
    Nuanciert: Project-READMEs (05_Projects/<slug>/README.md) sind echte
    User-Content-Files mit Frontmatter (type:project) und werden NICHT geskippt.
    """
    name = path.name
    # _index.md (auto-generiert von vault_toolkit) und CONTEXT.md (Projekt-
    # Kontext, freier Text ohne Schema) überall skippen.
    if name in ("_index.md", "CONTEXT.md"):
        return True
    # Vault-Root-Doku (CLAUDE/MOC/COMMANDS/SCHEMA/PIPELINES/README) skippen
    try:
        if path.parent == VAULT and name in _VAULT_ROOT_SYSTEM_DOCS:
            return True
    except Exception:
        pass
    # README.md in Subordnern: skippen AUSSER es ist eine Project-README.
    # Project-README liegt in 05_Projects/<slug>/README.md (parts ≥ 3),
    # NICHT 05_Projects/README.md (parts == 2 — das ist Top-Level-Doku).
    if name == "README.md":
        try:
            rel = path.relative_to(VAULT)
            if (len(rel.parts) >= 3 and rel.parts[0] == "05_Projects"):
                return False  # echte Project-README → behalten
            return True       # alle anderen README → System-Doku, skip
        except ValueError:
            return True
    return False

# Häufige Wörter die Bot oft schreibt — würden sonst nervig auto-linked
# wenn jemand zufällig ein Note/Task mit so einem Title hat
_LINK_STOPWORDS = {
    # Zeit
    "heute", "morgen", "gestern", "jetzt", "abends", "morgens",
    # Generic Vault-Begriffe
    "test", "todo", "task", "tasks", "meeting", "meetings", "note", "notes",
    "projekt", "projekte", "project", "projects", "area", "areas", "daily",
    # Häufige Sätze
    "info", "link", "wichtig", "okay", "danke", "bitte", "frage", "antwort",
    # Englisch
    "what", "where", "when", "this", "that", "these", "those",
}

# Min-Length für Phrase-Kandidaten — verhindert dass jeder 3-Buchstaben-Slug
# (RAG, API, etc.) jeden Vorkommnis kapert
_LINK_MIN_LEN = 4

# Schutz-Regex: alles was wir NICHT auto-linken (geschützte Bereiche werden
# durch Sentinels ersetzt, am Ende zurückgetauscht)
_PROTECT_RE = re.compile(
    r"```.*?```"             # fenced code
    r"|`[^`\n]+`"             # inline code
    r"|\[\[[^\]\n]+\]\]"      # already-wikilink
    r"|\[[^\]\n]+\]\([^)\n]+\)"  # markdown-link
    r"|https?://\S+"          # URL
    r"|<[^>\n]+>",            # HTML-Tag
    re.DOTALL,
)


def _build_link_index() -> dict:
    """Walk vault, extrahiere {phrase_lower: canonical_id} für Auto-Linking."""
    phrase_map: dict = {}
    if not VAULT.exists():
        return phrase_map

    for md_file, post in iter_vault_md():
        meta = post.metadata
        file_id = meta.get("id")
        if not file_id or not isinstance(file_id, str):
            continue
        file_id = file_id.strip()
        if len(file_id) < _LINK_MIN_LEN:
            continue
        if file_id.lower() in _LINK_STOPWORDS:
            continue

        # Kandidaten-Phrasen sammeln
        candidates = {file_id, file_id.replace("-", " "), file_id.replace("_", " ")}

        title = meta.get("title")
        if isinstance(title, str) and len(title.strip()) >= 6:
            t = title.strip()
            if t.lower() not in _LINK_STOPWORDS:
                candidates.add(t)

        aliases = meta.get("aliases", [])
        if isinstance(aliases, list):
            for a in aliases:
                if isinstance(a, str) and len(a.strip()) >= _LINK_MIN_LEN:
                    a_clean = a.strip()
                    if a_clean.lower() not in _LINK_STOPWORDS:
                        candidates.add(a_clean)

        for c in candidates:
            key = c.strip().lower()
            if not key or key in _LINK_STOPWORDS:
                continue
            # Erste Zuordnung gewinnt — bei Konflikt (zwei Files mit ähnlicher
            # Phrase) lieber gar nicht linken als falsch linken. Wir markieren
            # mit None.
            if key in phrase_map and phrase_map[key] != file_id:
                phrase_map[key] = None
            else:
                phrase_map.setdefault(key, file_id)

    # None-Einträge entfernen (mehrdeutig)
    return {k: v for k, v in phrase_map.items() if v is not None}


def _get_link_index() -> dict:
    global _AUTO_LINK_CACHE
    now = time.time()
    ts, pmap = _AUTO_LINK_CACHE
    if now - ts <= _AUTO_LINK_TTL_SEC and pmap:
        return pmap
    with _AUTO_LINK_LOCK:
        ts, pmap = _AUTO_LINK_CACHE
        if now - ts > _AUTO_LINK_TTL_SEC or not pmap:
            try:
                pmap = _build_link_index()
            except Exception as e:
                log.warning(f"link-index build failed: {e}")
                pmap = {}
            _AUTO_LINK_CACHE = (now, pmap)
    return pmap


_AUTO_LINK_INVALIDATE_FLOOR_SEC = 30  # Cache nach Mutation max so alt — nicht sofort reset


def invalidate_link_index() -> None:
    """Cache-Reset nach Schreibvorgang — mit TTL-Floor (nicht sofort).

    Vorher: jede Mutation → Cache komplett leer → nächster auto_link-Call
    macht vollen Vault-Walk. Bei Bulk-Move (6 Files): 6× Invalidate = bei
    nächsten 6 auto_link-Calls 6× Walk im Event-Loop.

    Jetzt: Cache wird auf "läuft in 30s ab" markiert. Mehrfach-Invalidate
    während dieser 30s = no-op. Bulk-Move mit nachfolgenden Reads → max
    EIN Walk in 30s. Auto-Link tolerant gegen 30s-stale-Daten.

    Plus: invalidiert today_data_cache (gleiche Trigger-Logik).
    """
    global _AUTO_LINK_CACHE
    ts, pmap = _AUTO_LINK_CACHE
    # Setze ts so, dass Cache in FLOOR Sekunden ungültig wird
    target_ts = time.time() - _AUTO_LINK_TTL_SEC + _AUTO_LINK_INVALIDATE_FLOOR_SEC
    # Nur "altern" wenn Cache aktuell jünger ist (verhindert paradoxes Reset)
    if ts > target_ts:
        _AUTO_LINK_CACHE = (target_ts, pmap)
    # Forward-tolerant — invalidate_today_data_cache wird später definiert.
    fn = globals().get("invalidate_today_data_cache")
    if fn is not None:
        fn()


def auto_link(text: str, exclude_ids: Optional[set] = None) -> str:
    """Findet bekannte Vault-IDs im Text + ersetzt durch [[wikilinks]].

    Schützt Code, bestehende Links, URLs vor Substitution. Idempotent —
    Doppelaufruf produziert keine doppelten Links.

    exclude_ids: IDs die NICHT verlinkt werden sollen (typisch: das File
    das gerade selbst geschrieben wird → keine Self-Links).
    """
    if not text or not isinstance(text, str):
        return text

    pmap = _get_link_index()
    if not pmap:
        return text

    exclude = exclude_ids or set()

    # 1) Geschützte Bereiche stashen
    stash: list = []

    def _stash_repl(m):
        stash.append(m.group(0))
        return f"\x00LINK{len(stash)-1}\x00"

    safe_text = _PROTECT_RE.sub(_stash_repl, text)

    # 2) Phrasen sortieren — längste zuerst (damit "BBM Skript K-Blätter"
    # vor "BBM" matcht und nicht umgekehrt)
    phrases = sorted(pmap.keys(), key=len, reverse=True)

    # Cap: max 30 Substitutionen pro Aufruf gegen Pathological-Input
    sub_count = 0
    SUB_CAP = 30

    for phrase in phrases:
        if sub_count >= SUB_CAP:
            break
        target_id = pmap[phrase]
        if target_id in exclude:
            continue
        if len(phrase) < _LINK_MIN_LEN:
            continue

        # Whole-word match — Wortgrenzen mit Umlaut-Awareness:
        # negative Lookarounds auf [Buchstabe/Ziffer/Underscore/Umlaut].
        pat = re.compile(
            r"(?<![\wÀ-ſ])"
            + re.escape(phrase)
            + r"(?![\wÀ-ſ])",
            re.IGNORECASE,
        )

        def _link_repl(m):
            nonlocal sub_count
            if sub_count >= SUB_CAP:
                return m.group(0)
            sub_count += 1
            matched = m.group(0)
            # Wenn matched lowercase exakt = ID → einfacher [[id]]-Link
            if matched.lower() == target_id.lower():
                return f"[[{target_id}]]"
            # Sonst Display-Form: [[id|original-Schreibweise]]
            return f"[[{target_id}|{matched}]]"

        safe_text = pat.sub(_link_repl, safe_text)

    # 3) Stash zurücktauschen
    def _unstash(m):
        idx = int(m.group(1))
        return stash[idx] if idx < len(stash) else m.group(0)

    safe_text = re.sub(r"\x00LINK(\d+)\x00", _unstash, safe_text)
    return safe_text


def append_to_daily(section: str, text: str) -> str:
    """Append text to today's daily under the given section.

    Spezielle Logik: wenn Sektion nur Template-Platzhalter enthält (z.B. '- [ ]'
    aus daily_template), wird der Platzhalter ersetzt statt zusätzlich gestapelt.
    """
    if section not in VALID_SECTIONS:
        section = "Notizen & Gedanken"
    path = ensure_daily()
    content = path.read_text(encoding="utf-8")

    # Auto-Link bekannte Vault-IDs/Titles im neuen Text — exclude die Daily
    # selbst (verhindert Self-Reference wenn die Daily-ID im Text vorkommt).
    today_daily_id = f"daily-{today_iso()}"
    text = auto_link(text, exclude_ids={today_daily_id})

    pattern = rf"(?m)^## {re.escape(section)}\s*$"
    match = re.search(pattern, content)
    if not match:
        # Sektion existiert nicht → am Ende anhängen
        new_content = content.rstrip() + f"\n\n## {section}\n{text}\n"
    else:
        start = match.end()
        next_h2 = re.search(r"^## ", content[start:], re.MULTILINE)
        end = start + next_h2.start() if next_h2 else len(content)

        # Aktueller Body der Sektion (ohne Header)
        body_before_h = content[:start]
        body_section = content[start:end].strip("\n")
        body_after = content[end:]

        # Body von Platzhaltern bereinigen
        cleaned = _clean_section_body(body_section)

        # Neuen Text anfügen
        if cleaned:
            new_section_body = cleaned + "\n" + text
        else:
            new_section_body = text

        new_content = (
            body_before_h.rstrip("\n")
            + "\n" + new_section_body
            + "\n\n"
            + body_after.lstrip("\n")
        )

    post = frontmatter.loads(new_content)
    post["updated"] = today_iso()
    atomic_write(path, frontmatter.dumps(post) + "\n")
    return f"In Daily ({section}) eingetragen: {path.name}"


VALID_PRIORITIES = {"low", "medium", "high", "urgent"}
VALID_TASK_CONTEXTS = {"home", "work", "errand", "phone", "computer"}
# WICHTIG: getrennt von Reminder-Recurrence — Tasks haben monthly, Reminders nicht.
# Frühere Doppeldefinition als VALID_RECURRENCE hat sich überschrieben → Bug.
VALID_TASK_RECURRENCE = {"daily", "weekdays", "weekly", "monthly"}


def create_task(title: str, priority: str = "medium",
                due: Optional[str] = None, area: Optional[str] = None,
                project: Optional[str] = None, context: Optional[str] = None,
                tags: Optional[list] = None,
                recurrence: Optional[str] = None) -> str:
    """Create a new task file in 10_Life/tasks/.

    Body wird dynamisch gebaut (nicht Template-Substitution) damit
    Status/Priorität/Fälligkeit korrekt im Body erscheinen.

    recurrence: optional "daily" / "weekdays" / "weekly" / "monthly".
      Bei gesetztem recurrence wird die Task nach mark_task_done vom
      Daily-Reset-Job am nächsten passenden Tag wieder auf 'open' gesetzt.
    """
    if not title or not title.strip():
        return "Fehler: Task-Titel darf nicht leer sein."

    # Priority normalisieren — LLM schickt manchmal Abkürzungen wie "M", "high prio" etc.
    p_lower = (priority or "").lower().strip()
    p_map = {
        "u": "urgent", "urgent": "urgent",
        "h": "high", "high": "high",
        "m": "medium", "medium": "medium", "med": "medium", "normal": "medium",
        "l": "low", "low": "low",
    }
    priority = p_map.get(p_lower, "medium")

    # Due nur akzeptieren wenn ISO-Format
    if due:
        due = due.strip()
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", due):
            log.warning(f"create_task: ungültiges due-Format '{due}', ignoriert")
            due = None

    # Recurrence normalisieren + validieren
    if recurrence:
        rec_lower = recurrence.strip().lower()
        # Aliase: weekday=weekdays, week=weekly, month=monthly, day=daily
        rec_map = {
            "day": "daily", "daily": "daily", "täglich": "daily",
            "weekday": "weekdays", "weekdays": "weekdays", "werktags": "weekdays",
            "week": "weekly", "weekly": "weekly", "wöchentlich": "weekly",
            "month": "monthly", "monthly": "monthly", "monatlich": "monthly",
        }
        recurrence = rec_map.get(rec_lower)
        if recurrence is None:
            log.warning(f"create_task: ungültiges recurrence '{rec_lower}', ignoriert")

    slug = slugify(title)
    path = TASKS_DIR / f"{slug}.md"
    n = 2
    while path.exists():
        path = TASKS_DIR / f"{slug}-{n}.md"
        n += 1
    today = today_iso()
    task_id = f"t-{path.stem}"

    # Body dynamisch bauen (statt Template mit Hardcoded-Defaults)
    rec_line = f" · **Wiederholung**: {recurrence}" if recurrence else ""
    body = (
        f"# {title}\n\n"
        f"**Status**: open · **Priorität**: {priority} · **Fällig**: {due or '—'}{rec_line}\n\n"
        f"## Was\n\n"
        f"## Warum\n\n"
        f"## Subschritte\n- [ ]\n\n"
        f"## Notizen\n\n"
        f"## Log\n- {today}: angelegt\n"
    )

    # Tags filtern (nur strings, nicht-leer, deduplizieren)
    clean_tags = []
    if tags and isinstance(tags, list):
        seen = set()
        for t in tags:
            if isinstance(t, str) and t.strip() and t.strip() not in seen:
                clean_tags.append(t.strip().lower())
                seen.add(t.strip())

    fm_data = {
        "id": task_id,
        "title": title,
        "type": "task",
        "created": today,
        "updated": today,
        "tags": clean_tags,
        "status": "open",
        "priority": priority,
    }
    if due:
        fm_data["due"] = due
    if area:
        fm_data["area"] = area
    if project:
        fm_data["project"] = project
    if context:
        ctx_lower = context.lower().strip()
        if ctx_lower in VALID_TASK_CONTEXTS:
            fm_data["context"] = ctx_lower
    if recurrence:
        fm_data["recurrence"] = recurrence

    post = frontmatter.Post(body, **fm_data)
    atomic_write(path, frontmatter.dumps(post) + "\n")

    # Link in heutige Daily
    try:
        # Pipe-Syntax: title ist im Stash von auto_link geschützt → kein
        # nested-Wikilink-Bug wenn title Substring von einem anderen ID ist
        # (vorher: "Matura" im Title → auto_link nested daraus [[project-matura|...]])
        append_to_daily("Heute", f"- [ ] [[{task_id}|{title}]]")
    except Exception as e:
        log.warning(f"Daily-Link für Task fehlgeschlagen: {e}")

    invalidate_link_index()  # ruft auch invalidate_today_data_cache (siehe dort)

    extras = []
    if due:
        extras.append(f"fällig {due}")
    if project:
        extras.append(f"projekt {project}")
    if priority != "medium":
        extras.append(f"prio {priority}")
    if recurrence:
        extras.append(f"wiederholt {recurrence}")
    extra_str = f" ({', '.join(extras)})" if extras else ""
    return f"Task angelegt: [[{task_id}]]{extra_str}"


def mark_task_done(slug: str) -> str:
    """Mark task as done.

    Bei recurring Tasks (frontmatter.recurrence gesetzt): Status bleibt 'done'
    bis der recurring_task_reset_job am passenden nächsten Tag wieder auf
    'open' setzt. last_completed wird gesetzt damit Reset weiß wann es passt.
    """
    filename = slug[2:] if slug.startswith("t-") else slug
    # Path-Traversal-Schutz: slug ist LLM-kontrolliert. Nur kebab-case + .md erlauben.
    if not filename or not re.match(r"^[a-zA-Z0-9_\-]+$", filename):
        return f"Ungültiger Task-Slug: {slug!r}"
    try:
        path = safe_path(f"10_Life/tasks/{filename}.md")
    except ValueError:
        return f"Ungültiger Task-Pfad: {slug!r}"
    if not path.exists():
        return f"Task nicht gefunden: {slug}"
    post = frontmatter.load(path)
    today = today_iso()
    post["status"] = "done"
    post["updated"] = today
    recurrence = post.metadata.get("recurrence")
    if recurrence:
        post["last_completed"] = today
        log_line = f"\n- {today}: erledigt (recurring={recurrence}, kommt automatisch wieder)\n"
    else:
        log_line = f"\n- {today}: erledigt\n"
    body = (post.content or "").rstrip() + log_line
    post.content = body
    atomic_write(path, frontmatter.dumps(post) + "\n")
    if recurrence:
        return f"Task erledigt: [[t-{filename}]] — wiederholt sich ({recurrence})"
    return f"Task erledigt: [[t-{filename}]]"


# ─── Tagesplanung: Listings + Agenda ────────────────────────────────────────

# Priority-Sortierung + Symbole (zentral — vermeidet Drift zwischen
# _format_task_line und compute_briefing wo vorher unterschiedliche
# Symbol-Maps definiert waren)
_PRIO_ORDER = {"urgent": 0, "high": 1, "medium": 2, "low": 3}
# Symbole nach visueller Lautstärke: 🔴 > 🟠 > • > · — low ist quasi unsichtbar,
# medium dezenter Bullet, urgent/high deutlich
PRIO_SYMBOLS = {"urgent": "🔴", "high": "🟠", "medium": "•", "low": "·"}


def _read_open_tasks() -> list:
    """Walked TASKS_DIR, gibt Liste aller offenen Tasks mit Frontmatter zurück."""
    if not TASKS_DIR.exists():
        return []
    out = []
    for f, post in iter_vault_md(TASKS_DIR, recursive=False, skip_noise=False):
        meta = post.metadata
        status = meta.get("status", "")
        if status not in ("open", "in-progress", "blocked"):
            continue
        out.append({
            "slug": f.stem,
            "id": meta.get("id", f"t-{f.stem}"),
            "title": meta.get("title") or f.stem,
            "status": status,
            "priority": meta.get("priority", "medium"),
            "due": meta.get("due"),
            "project": meta.get("project"),
            "area": meta.get("area"),
            "context": meta.get("context"),
            "tags": meta.get("tags", []) or [],
            "recurrence": meta.get("recurrence"),
        })
    return out


def list_open_tasks(when: Optional[str] = None,
                    priority: Optional[str] = None,
                    project: Optional[str] = None,
                    area: Optional[str] = None,
                    context: Optional[str] = None) -> str:
    """Listet offene Tasks mit smartem Filter, gruppiert + sortiert für Telegram.

    when: 'today' / 'tomorrow' / 'week' (nächste 7 Tage) / 'overdue' / 'nodate' / None=alle
    priority/project/area/context: optionale weitere Filter
    """
    tasks = _read_open_tasks()
    if not tasks:
        return "Keine offenen Tasks."

    today = datetime.now(TIMEZONE).date()

    # Nutze zentrale _due_to_date — strptime-Variante hatte TypeError-Falle
    # bei YAML-date-Objekten (Tasks mit unquoted `due:` aus Obsidian-Reload).
    def _due_date(t):
        return _due_to_date(t.get("due"))

    def _matches(t):
        # when
        if when:
            d = _due_date(t)
            w = when.strip().lower()
            if w == "today":
                if d != today:
                    return False
            elif w == "tomorrow":
                if d != today + timedelta(days=1):
                    return False
            elif w == "week":
                if d is None or not (today <= d <= today + timedelta(days=7)):
                    return False
            elif w == "overdue":
                if d is None or d >= today:
                    return False
            elif w == "nodate":
                if d is not None:
                    return False
            # unbekanntes when wird ignoriert (kein Filter)
        if priority:
            if t.get("priority", "").lower() != priority.strip().lower():
                return False
        if project:
            if (t.get("project") or "").strip().lower() != project.strip().lower():
                return False
        if area:
            if (t.get("area") or "").strip().lower() != area.strip().lower():
                return False
        if context:
            if (t.get("context") or "").strip().lower() != context.strip().lower():
                return False
        return True

    matched = [t for t in tasks if _matches(t)]
    if not matched:
        filter_desc = []
        if when: filter_desc.append(f"when={when}")
        if priority: filter_desc.append(f"prio={priority}")
        if project: filter_desc.append(f"projekt={project}")
        if area: filter_desc.append(f"area={area}")
        if context: filter_desc.append(f"context={context}")
        f_str = f" (Filter: {', '.join(filter_desc)})" if filter_desc else ""
        return f"Keine offenen Tasks{f_str}."

    # Wenn kein when-Filter: gruppieren in 4 Buckets nach Fälligkeit
    if not when:
        overdue, today_t, week_t, later_t, nodate_t = [], [], [], [], []
        for t in matched:
            d = _due_date(t)
            if d is None:
                nodate_t.append(t)
            elif d < today:
                overdue.append(t)
            elif d == today:
                today_t.append(t)
            elif d <= today + timedelta(days=7):
                week_t.append(t)
            else:
                later_t.append(t)

        def _sort(lst):
            return sorted(lst, key=lambda t: (_PRIO_ORDER.get(t.get("priority"), 9), _due_date(t) or date.max))

        sections = []
        for label, lst in [
            ("⚠️ Überfällig", _sort(overdue)),
            ("📅 Heute", _sort(today_t)),
            ("📆 Diese Woche", _sort(week_t)),
            ("🗓 Später", _sort(later_t)),
            ("∞ Ohne Datum", _sort(nodate_t)),
        ]:
            if not lst:
                continue
            sections.append(f"**{label}** ({len(lst)})")
            for t in lst[:15]:  # Cap pro Sektion
                sections.append(_format_task_line(t, today))
            if len(lst) > 15:
                sections.append(f"  _… {len(lst) - 15} weitere_")
            sections.append("")
        # Trailing-Empty entfernen
        while sections and not sections[-1]:
            sections.pop()
        return "\n".join(sections)

    # Mit when-Filter: flache Liste, nach Prio sortiert
    matched_sorted = sorted(matched, key=lambda t: (_PRIO_ORDER.get(t.get("priority"), 9), _due_date(t) or date.max))
    lines = [f"Offene Tasks (Filter: {when}, {len(matched_sorted)}):"]
    for t in matched_sorted[:30]:
        lines.append(_format_task_line(t, today))
    if len(matched_sorted) > 30:
        lines.append(f"  _… {len(matched_sorted) - 30} weitere_")
    return "\n".join(lines)


def _format_task_line(t: dict, today: date) -> str:
    """Eine Task als Bullet-Zeile mit Symbolen.

    Robust gegen YAML-Date-Objekte: PyYAML parsed unquoted `due: 2026-04-30`
    zu `datetime.date`, nicht zu str → strptime würde TypeError werfen.
    """
    prio_sym = PRIO_SYMBOLS.get(t.get("priority"), PRIO_SYMBOLS["medium"])
    rec_sym = " 🔁" if t.get("recurrence") else ""
    proj = f" [{t['project']}]" if t.get("project") else ""
    due_str = ""
    d = t.get("due")
    if d:
        # Normalisieren: date-Objekt oder String beide akzeptieren
        dd = None
        if isinstance(d, date) and not isinstance(d, datetime):
            dd = d
        elif isinstance(d, datetime):
            dd = d.date()
        elif isinstance(d, str):
            try:
                dd = datetime.strptime(d, "%Y-%m-%d").date()
            except ValueError:
                pass  # weiter unten als raw-String anzeigen

        if dd is not None:
            delta = (dd - today).days
            if delta == 0:
                due_str = " · heute"
            elif delta == 1:
                due_str = " · morgen"
            elif delta < 0:
                due_str = f" · ⚠️ vor {-delta}d"
            elif delta <= 7:
                due_str = f" · in {delta}d"
            else:
                due_str = f" · {dd.isoformat()}"
        else:
            due_str = f" · {d}"
    # Wikilink bewusst NICHT mehr in der Bullet-Zeile — der LLM kopiert ihn
    # sonst stumpf in die Telegram-Antwort und User sieht den Slug-Lärm.
    # ID steht weiter im Frontmatter und ist via search_vault auffindbar,
    # falls der LLM auf einen Task per ID referenzieren muss.
    return f"  {prio_sym} {t['title']}{proj}{due_str}{rec_sym}"


def _due_to_date(d) -> Optional[date]:
    """Normalisiert Frontmatter-due (str/date/datetime) zu date oder None."""
    if not d:
        return None
    if isinstance(d, date) and not isinstance(d, datetime):
        return d
    if isinstance(d, datetime):
        return d.date()
    if isinstance(d, str):
        try:
            return datetime.strptime(d, "%Y-%m-%d").date()
        except ValueError:
            return None
    return None


def _sort_tasks_by_prio_due(lst: list) -> list:
    """Sort: Prio (urgent first), dann Due (älteste first, None last)."""
    return sorted(lst, key=lambda t: (_PRIO_ORDER.get(t.get("priority"), 9),
                                      _due_to_date(t.get("due")) or date.max))


# Cache für collect_today_data — vermeidet Vault-Walk bei jedem
# /today, /tasks, get_today_agenda-Aufruf. 30s reicht für interaktive
# Konversation, ist kurz genug für unmittelbare Reaktion auf Edits.
_TODAY_DATA_CACHE: tuple = (0.0, None)
_TODAY_DATA_TTL_SEC = 30


def invalidate_today_data_cache() -> None:
    """Cache-Reset nach Mutations die die Tagesansicht verändern."""
    global _TODAY_DATA_CACHE
    _TODAY_DATA_CACHE = (0.0, None)


def collect_today_data() -> dict:
    """Single Source of Truth: aggregiert ALLE "heute"-Daten in einem Walk.

    Genutzt von compute_briefing (HTML-Render) UND get_today_agenda (Markdown).
    Vorher waren beide Renderer mit eigenem Aggregations-Code → Drift-Risiko.

    Cached 30s — bei jedem create_task/mark_task_done/create_reminder
    via invalidate_today_data_cache() invalidiert.
    """
    global _TODAY_DATA_CACHE
    ts, cached = _TODAY_DATA_CACHE
    if cached is not None and (time.time() - ts) <= _TODAY_DATA_TTL_SEC:
        return cached
    data = _collect_today_data_uncached()
    _TODAY_DATA_CACHE = (time.time(), data)
    return data


def _collect_today_data_uncached() -> dict:
    """Tatsächliche Berechnung — nur aus collect_today_data oder Tests rufen.

    Returns dict mit Keys:
      today: date · today_str: ISO-string · now: datetime
      reminders: [(fire_at_dt, reminder_dict), ...]  — heute & noch nicht gefeuert, sortiert
      meetings: [{id, title, status, attendees}, ...]  — heute, status != cancelled
      overdue_tasks: [task_dict, ...]   — open/in-progress, due < today, prio-sortiert
      today_tasks: [task_dict, ...]      — open/in-progress, due == today, prio-sortiert
      high_nodate_tasks: [task_dict, ...]  — open/in-progress, due is None, prio in (urgent, high)
    """
    today = datetime.now(TIMEZONE).date()
    today_str = today.strftime("%Y-%m-%d")
    now = datetime.now(TIMEZONE)

    # Reminders heute (noch nicht gefeuert), sortiert
    reminders = []
    for r in _load_reminders():
        try:
            fire_at = datetime.fromisoformat(r["fire_at"])
            if fire_at.tzinfo is None:
                fire_at = fire_at.replace(tzinfo=TIMEZONE)
            if fire_at.date() == today and fire_at >= now:
                reminders.append((fire_at, r))
        except (ValueError, KeyError, TypeError):
            continue
    reminders.sort(key=lambda x: x[0])

    # Meetings heute (nicht-cancelled)
    meetings = []
    if MEETINGS_DIR.exists():
        for f in MEETINGS_DIR.glob(f"{today_str}_*.md"):
            try:
                post = frontmatter.load(f)
                if post.metadata.get("status", "") == "cancelled":
                    continue
                meetings.append({
                    "id": post.metadata.get("id", f.stem),
                    "title": post.metadata.get("title", f.stem),
                    "status": post.metadata.get("status", ""),
                    "attendees": post.metadata.get("attendees", []),
                })
            except Exception:
                continue

    # Tasks: drei Buckets
    all_open = _read_open_tasks()
    overdue_tasks, today_tasks = [], []
    for t in all_open:
        dd = _due_to_date(t.get("due"))
        if dd is None:
            continue
        if dd < today:
            overdue_tasks.append(t)
        elif dd == today:
            today_tasks.append(t)
    high_nodate_tasks = [t for t in all_open
                         if t.get("due") is None
                         and t.get("priority") in ("urgent", "high")]

    return {
        "today": today,
        "today_str": today_str,
        "now": now,
        "reminders": reminders,
        "meetings": meetings,
        "overdue_tasks": _sort_tasks_by_prio_due(overdue_tasks),
        "today_tasks": _sort_tasks_by_prio_due(today_tasks),
        "high_nodate_tasks": _sort_tasks_by_prio_due(high_nodate_tasks),
    }


def get_today_agenda() -> str:
    """Markdown-Render des collect_today_data-Snapshots für Telegram-Tool-Call."""
    data = collect_today_data()
    today = data["today"]
    parts = [f"📋 Agenda für heute ({today.strftime('%a %d.%m.%Y')})"]

    if data["reminders"]:
        parts.append(f"\n⏰ **Erinnerungen heute** ({len(data['reminders'])})")
        for fire_at, r in data["reminders"]:
            rec = " 🔁" if r.get("recurrence") else ""
            parts.append(f"  • {fire_at.strftime('%H:%M')} — {r['message'][:80]}{rec}")

    if data["meetings"]:
        parts.append(f"\n🤝 **Meetings heute** ({len(data['meetings'])})")
        for m in data["meetings"]:
            att = ""
            if m["attendees"]:
                att_list = ", ".join(str(a) for a in m["attendees"][:3])
                if len(m["attendees"]) > 3:
                    att_list += f" +{len(m['attendees'])-3}"
                att = f" mit {att_list}"
            parts.append(f"  • [[{m['id']}]] {m['title']}{att}")

    if data["overdue_tasks"]:
        parts.append(f"\n⚠️ **Überfällige Tasks** ({len(data['overdue_tasks'])})")
        for t in data["overdue_tasks"][:10]:
            parts.append(_format_task_line(t, today))
        if len(data["overdue_tasks"]) > 10:
            parts.append(f"  _… {len(data['overdue_tasks'])-10} weitere_")

    if data["today_tasks"]:
        parts.append(f"\n📅 **Heute fällige Tasks** ({len(data['today_tasks'])})")
        for t in data["today_tasks"][:15]:
            parts.append(_format_task_line(t, today))
        if len(data["today_tasks"]) > 15:
            parts.append(f"  _… {len(data['today_tasks'])-15} weitere_")

    if data["high_nodate_tasks"]:
        parts.append(f"\n🔥 **Hohe Prio, kein Datum** ({len(data['high_nodate_tasks'])})")
        for t in data["high_nodate_tasks"][:10]:
            parts.append(_format_task_line(t, today))

    if len(parts) == 1:
        parts.append("\n_Heute steht aktuell nichts im Vault. Schreib mir was du heute vorhast — ich trag's ein._")

    return "\n".join(parts)


def create_meeting(title: str, attendees: Optional[list] = None,
                   meeting_date: Optional[str] = None,
                   tags: Optional[list] = None) -> str:
    """Create meeting protocol — Body dynamisch (kein Template-Cruft).

    meeting_date: NUR ISO YYYY-MM-DD akzeptiert. Bei Garbage/Path-Traversal
    Fallback auf heute (LLM darf nicht beliebige Strings in den Pfad pumpen).
    """
    if not title or not title.strip():
        return "Fehler: Meeting-Titel darf nicht leer sein."
    # Datum strikt validieren — sonst Path-Traversal über meeting_date möglich
    if meeting_date and re.match(r"^\d{4}-\d{2}-\d{2}$", str(meeting_date).strip()):
        try:
            datetime.strptime(meeting_date.strip(), "%Y-%m-%d")
            today = meeting_date.strip()
        except ValueError:
            log.warning(f"create_meeting: ungültiges Datum '{meeting_date}', nutze heute")
            today = today_iso()
    else:
        if meeting_date:
            log.warning(f"create_meeting: kein ISO-Datum '{meeting_date}', nutze heute")
        today = today_iso()
    slug = slugify(title)
    # safe_path defensiv anwenden — slug ist via slugify safe, today ist validiert.
    try:
        path = safe_path(f"10_Life/meetings/{today}_{slug}.md")
    except ValueError:
        return f"Ungültiger Meeting-Pfad ({today}_{slug})"
    n = 2
    while path.exists():
        try:
            path = safe_path(f"10_Life/meetings/{today}_{slug}-{n}.md")
        except ValueError:
            return f"Ungültiger Meeting-Pfad-Variante (n={n})"
        n += 1

    # Tags filtern
    clean_tags = []
    if tags and isinstance(tags, list):
        seen = set()
        for t in tags:
            if isinstance(t, str) and t.strip() and t.strip() not in seen:
                clean_tags.append(t.strip().lower())
                seen.add(t.strip())

    attendees_list = attendees or []
    attendees_str = ", ".join(f"[[{a}]]" for a in attendees_list) if attendees_list else "—"
    status = "done" if today <= today_iso() else "planned"

    body = (
        f"# {title}\n\n"
        f"**Datum**: {today} · **Teilnehmer**: {attendees_str} · **Status**: {status}\n\n"
        f"## Agenda\n- \n\n"
        f"## Diskussion\n\n"
        f"## Entscheidungen\n- \n\n"
        f"## Action Items\n- [ ] \n"
    )

    fm_data = {
        "id": f"meeting-{today}-{slug}",
        "title": title,
        "type": "meeting",
        "date": today,
        "created": today_iso(),
        "updated": today_iso(),
        "attendees": attendees_list,
        "status": status,
        "tags": clean_tags,
    }
    post = frontmatter.Post(body, **fm_data)
    atomic_write(path, frontmatter.dumps(post) + "\n")
    invalidate_link_index()
    return f"Meeting angelegt: [[meeting-{today}-{slug}]]"


def create_note(title: str, body: str, tags: Optional[list] = None) -> str:
    """Create a free note in 10_Life/notes/. Body wird dynamisch gebaut, kein Template-Platzhalter."""
    if not title or not title.strip():
        return "Fehler: Note-Titel darf nicht leer sein."

    today = today_iso()
    slug = slugify(title)
    path = NOTES_DIR / f"{today}_{slug}.md"
    n = 2
    while path.exists():
        path = NOTES_DIR / f"{today}_{slug}-{n}.md"
        n += 1

    # Tags filtern (nur strings, nicht-leer, deduplizieren, lowercase)
    clean_tags = []
    if tags and isinstance(tags, list):
        seen = set()
        for t in tags:
            if isinstance(t, str) and t.strip() and t.strip() not in seen:
                clean_tags.append(t.strip().lower())
                seen.add(t.strip())

    # Auto-Link bekannte Vault-IDs/Titles — exclude die Note selbst
    linked_body = auto_link(body.strip(), exclude_ids={slug})

    # Body dynamisch — direkt H1 + User-Body, kein Template-Comment
    note_body = f"# {title}\n\n{linked_body}\n"

    fm_data = {
        "id": slug,
        "title": title,
        "type": "note",
        "created": today,
        "updated": today,
        "tags": clean_tags,
        "status": "draft",
        "quelle": "telegram",
    }
    post = frontmatter.Post(note_body, **fm_data)
    atomic_write(path, frontmatter.dumps(post) + "\n")
    invalidate_link_index()
    return f"Notiz angelegt: [[{slug}]]"


def search_vault(query: str, limit: int = 5) -> str:
    """Volltext-Suche via vault_search.py (subprocess)."""
    script = VAULT / "07_Tools" / "search" / "vault_search.py"
    if not script.exists():
        return "vault_search.py nicht gefunden."
    try:
        # Minimales env: vault_search braucht keine Tokens.
        # Sonst würden TG_TOKEN/LLM_API_KEY/GITHUB_BACKUP_TOKEN an subprocess geleakt.
        minimal_env = {k: v for k, v in os.environ.items()
                       if k in ("PATH", "PYTHONPATH", "PYTHONIOENCODING", "LANG", "LC_ALL", "HOME")}
        result = subprocess.run(
            ["python3", str(script), "--json", query],
            capture_output=True, text=True, timeout=30, cwd=str(VAULT),
            env=minimal_env,
        )
        if result.returncode != 0:
            return f"Suche fehlgeschlagen: {result.stderr.strip()[:300]}"
        data = json.loads(result.stdout) if result.stdout.strip() else []
        if not data:
            return f"Keine Treffer für: {query}"
        lines = [f"Suche '{query}' — {len(data)} Treffer (top {limit}):"]
        for hit in data[:limit]:
            hid = hit.get("id", "?")
            htype = hit.get("type", "?")
            score = hit.get("score", "?")
            lines.append(f"- [[{hid}]] · `{htype}` · score {score}")
        return "\n".join(lines)
    except Exception as e:
        return f"Suche-Fehler: {e}"


def read_file(rel_path: str, strip_frontmatter: bool = True) -> str:
    """Read a file (relative to vault root). Capped at 8KB.

    strip_frontmatter=True (default): YAML-Frontmatter wird entfernt für
    saubere Anzeige. Auf False setzen wenn du Metadaten brauchst.
    """
    try:
        path = safe_path(rel_path)
        if not path.exists():
            return f"Datei nicht gefunden: {rel_path}"
        content = path.read_text(encoding="utf-8")
        if strip_frontmatter:
            # Frontmatter zwischen --- ... --- am Anfang entfernen
            content = re.sub(r"^---\n.*?\n---\n+", "", content, count=1, flags=re.DOTALL)
        return content[:8000]
    except Exception as e:
        return f"Lese-Fehler: {e}"


def move_path(src_rel: str, dst_rel: str, overwrite: bool = False) -> str:
    """Verschiebt/Renamet eine Datei oder einen Ordner innerhalb des Vaults.

    src_rel + dst_rel sind relativ zum Vault-Root. Path-Traversal-geschützt
    via safe_path. Wenn dst ein bestehender Ordner ist, wird src reingelegt.
    Wenn dst nicht existiert, wird src nach dst umbenannt/verschoben.

    overwrite=False (default): bricht ab wenn Ziel existiert.
    overwrite=True: überschreibt — nur nutzen wenn explizit gewollt.
    """
    if not src_rel or not src_rel.strip():
        return "Fehler: src darf nicht leer sein."
    if not dst_rel or not dst_rel.strip():
        return "Fehler: dst darf nicht leer sein."
    try:
        src = safe_path(src_rel)
        dst = safe_path(dst_rel)
    except ValueError as e:
        return f"Pfad-Fehler: {e}"

    if not src.exists():
        return f"Quelle nicht gefunden: {src_rel}"
    if src == VAULT.resolve():
        return "Fehler: Vault-Root selbst kann nicht verschoben werden."

    # Falls dst ein existierender Ordner ist → src darunter legen
    if dst.is_dir():
        final = dst / src.name
    else:
        final = dst

    if final.exists():
        if not overwrite:
            return f"Ziel existiert bereits: {final.relative_to(VAULT).as_posix()} (overwrite=True um zu überschreiben)"
        # Overwrite: erst alt löschen
        if final.is_dir():
            shutil.rmtree(final)
        else:
            final.unlink()

    final.parent.mkdir(parents=True, exist_ok=True)
    try:
        shutil.move(str(src), str(final))
    except Exception as e:
        return f"Move fehlgeschlagen: {e}"

    src_rel_clean = src.relative_to(VAULT).as_posix()
    final_rel = final.relative_to(VAULT).as_posix()
    kind = "Ordner" if final.is_dir() else "Datei"
    # Move kann .md-File in/aus Skip-Dirs (z.B. 99_Archive) bewegen → Index neu bauen
    invalidate_link_index()
    return f"✓ {kind} verschoben: `{src_rel_clean}` → `{final_rel}`"


def move_paths(srcs: list, dst_dir: str, overwrite: bool = False) -> str:
    """Bulk-Move: verschiebt mehrere Dateien/Ordner auf einmal in dst_dir.

    Spart massiv Tool-Calls bei Multi-File-Operationen (z.B. 6 Uploads in
    Projekt verschieben). Pro Item wird Erfolg/Fehler einzeln gemeldet,
    aber alle Items teilen sich denselben Tool-Call → Loop-Iterationen sparen.

    srcs: Liste von vault-relativen Pfaden (Dateien oder Ordner)
    dst_dir: vault-relatives Ziel-Verzeichnis (wird angelegt falls nicht da)
    """
    if not srcs or not isinstance(srcs, list):
        return "Fehler: srcs muss eine nicht-leere Liste sein."
    if not dst_dir or not dst_dir.strip():
        return "Fehler: dst_dir darf nicht leer sein."
    try:
        dst = safe_path(dst_dir)
    except ValueError as e:
        return f"Pfad-Fehler dst_dir: {e}"

    # dst_dir muss Ordner sein/werden — falls Datei mit gleichem Namen, ablehnen
    if dst.exists() and not dst.is_dir():
        return f"Ziel ist eine Datei, kein Ordner: {dst_dir}"
    dst.mkdir(parents=True, exist_ok=True)

    successes, failures = [], []
    for src_rel in srcs:
        if not isinstance(src_rel, str) or not src_rel.strip():
            failures.append(f"(leerer Eintrag)")
            continue
        try:
            src = safe_path(src_rel)
        except ValueError as e:
            failures.append(f"{src_rel}: Pfad-Fehler {e}")
            continue
        if not src.exists():
            failures.append(f"{src_rel}: nicht gefunden")
            continue
        if src == VAULT.resolve():
            failures.append(f"{src_rel}: Vault-Root unbeweglich")
            continue
        final = dst / src.name
        if final.exists():
            if not overwrite:
                failures.append(f"{src.name}: Ziel existiert (overwrite=False)")
                continue
            try:
                if final.is_dir():
                    shutil.rmtree(final)
                else:
                    final.unlink()
            except Exception as e:
                failures.append(f"{src.name}: Overwrite-Cleanup fehlschlug: {e}")
                continue
        try:
            shutil.move(str(src), str(final))
            successes.append(src.name)
        except Exception as e:
            failures.append(f"{src.name}: {e}")

    if successes:
        invalidate_link_index()  # Files könnten in/aus Skip-Dirs verschoben sein

    dst_rel_out = dst.relative_to(VAULT).as_posix()
    parts = [f"✓ {len(successes)} verschoben → `{dst_rel_out}/`"]
    if successes:
        parts.append("  " + ", ".join(successes[:8]) + (f" (+{len(successes)-8})" if len(successes) > 8 else ""))
    if failures:
        parts.append(f"✗ {len(failures)} fehlgeschlagen:")
        for f in failures[:5]:
            parts.append(f"  • {f}")
        if len(failures) > 5:
            parts.append(f"  • (+{len(failures)-5} weitere)")
    return "\n".join(parts)


def delete_path(rel_path: str, recursive: bool = False) -> str:
    """Löscht eine Datei oder einen leeren Ordner. recursive=True für Ordner mit Inhalt.

    GEFÄHRLICH bei recursive=True — ruft erst nach expliziter User-Bestätigung.
    Verschiebt nach 99_Archive/ statt zu löschen wäre sicherer; das ist
    aber Sache des Callers (z.B. via move_path nach 99_Archive).
    """
    if not rel_path or not rel_path.strip():
        return "Fehler: Pfad darf nicht leer sein."
    try:
        p = safe_path(rel_path)
    except ValueError as e:
        return f"Pfad-Fehler: {e}"
    if not p.exists():
        return f"Nicht gefunden: {rel_path}"
    if p == VAULT.resolve():
        return "Fehler: Vault-Root kann nicht gelöscht werden."
    try:
        was_dir = p.is_dir()
        if was_dir:
            if recursive:
                shutil.rmtree(p)
                msg = f"✓ Ordner rekursiv gelöscht: `{rel_path}`"
            else:
                p.rmdir()  # nur wenn leer
                msg = f"✓ Leerer Ordner gelöscht: `{rel_path}`"
        else:
            p.unlink()
            msg = f"✓ Datei gelöscht: `{rel_path}`"
        invalidate_link_index()  # gelöschte ID darf nicht mehr verlinkt werden
        return msg
    except OSError as e:
        return f"Löschen fehlgeschlagen: {e}"


def move_project(slug: str, parent: Optional[str] = None) -> str:
    """Verschiebt ein bestehendes Projekt — entweder als Subprojekt unter `parent`,
    oder zurück auf Top-Level wenn parent=None oder parent='' angegeben wird.

    Nutzt rekursive Suche → findet Projekte überall unter 05_Projects/.
    """
    if not slug or not slug.strip():
        return "Slug fehlt."
    slug = slug.strip().lower()
    if slug.startswith("project-"):
        slug = slug[len("project-"):]
    src = find_project_dir(slug)
    if src is None:
        return f"Projekt nicht gefunden (oder mehrdeutig): {slug}"

    # Ziel bestimmen
    if parent and parent.strip():
        parent = parent.strip().lower()
        if parent.startswith("project-"):
            parent = parent[len("project-"):]
        if parent == slug:
            return "Fehler: Projekt kann nicht sich selbst als Parent haben."
        parent_dir = find_project_dir(parent)
        if parent_dir is None:
            return f"Parent-Projekt nicht gefunden: {parent}"
        # Prevent moving a project into its own subtree
        try:
            parent_dir.relative_to(src)
            return f"Fehler: Parent `{parent}` liegt bereits unter `{slug}` — würde Schleife erzeugen."
        except ValueError:
            pass
        dst = parent_dir / slug
    else:
        dst = PROJECTS_DIR / slug

    if dst.resolve() == src.resolve():
        return f"Projekt liegt bereits an Zielposition: `{src.relative_to(VAULT).as_posix()}/`"
    if dst.exists():
        return f"Ziel existiert bereits: `{dst.relative_to(VAULT).as_posix()}/`"

    dst.parent.mkdir(parents=True, exist_ok=True)
    try:
        shutil.move(str(src), str(dst))
    except Exception as e:
        return f"Move fehlgeschlagen: {e}"

    src_rel = src.relative_to(VAULT).as_posix()
    dst_rel = dst.relative_to(VAULT).as_posix()
    parent_info = f" (jetzt Subprojekt von `{parent}`)" if parent else " (jetzt Top-Level)"
    invalidate_link_index()
    return f"✓ Projekt verschoben: `{src_rel}/` → `{dst_rel}/`{parent_info}"


def move(src: Optional[str] = None, srcs: Optional[list] = None,
         dst: Optional[str] = None,
         project_slug: Optional[str] = None,
         parent: Optional[str] = None,
         overwrite: bool = False) -> str:
    """Vereinheitlichtes Move-Tool — drei Use-Cases mit getrennten Feldern:

    a) Einzeln: move(src='foo.md', dst='bar.md')
    b) Bulk:    move(srcs=['a.md','b.md'], dst='ordner/')
    c) Projekt: move(project_slug='matura', parent='dachboden-umbau')

    Getrennte src/srcs-Felder statt anyOf — manche LLM-Provider
    (OpenAI strict, Ollama-Cloud) lehnen anyOf in Tool-Schemas ab.

    Defensive: src kann auch eine Liste sein (Tippfehler vom LLM) →
    behandelt wie srcs.
    """
    if project_slug:
        return move_project(project_slug, parent)
    # Defensive: LLM könnte src statt srcs senden mit Liste
    if isinstance(src, list) and not srcs:
        srcs = src
        src = None
    if srcs:
        if not dst:
            return "Fehler: dst (Ziel-Ordner) nötig für Bulk-Move."
        if not isinstance(srcs, list):
            return "Fehler: srcs muss eine Liste sein."
        return move_paths(srcs, dst, overwrite)
    if src:
        if not dst:
            return "Fehler: dst nötig für Einzel-Move."
        return move_path(src, dst, overwrite)
    return "Fehler: keiner der drei Modi erkannt — gib src+dst ODER srcs+dst ODER project_slug an."


EDIT_FILE_MAX_BYTES = 5 * 1024 * 1024     # 5 MB — keine Massenfile-Edits
EDIT_FILE_MAX_REGEX_LEN = 500              # Regex >500 Zeichen → wahrscheinlich Halluzination
# Pathological-Regex-Patterns die ReDoS triggern können (nested quantifiers
# auf demselben Pattern). Konservativ — fängt die häufigsten Fälle.
_REDOS_PATTERNS = [
    re.compile(r"\([^)]*[+*]\)[+*]"),     # (a+)+ / (a*)*
    re.compile(r"\([^)]*\|[^)]*\)[+*]"),  # (a|a)+ / (a|b)*
]


def edit_file(rel_path: str, find: str, replace: str, regex: bool = False) -> str:
    """Find/replace in a file.

    SECURITY: rel_path via safe_path. Bei regex=True: ReDoS-Schutz via
    Pattern-Heuristik + Längen-Cap. File-Cap 5MB gegen accidental DoS.
    """
    if not isinstance(find, str) or not find:
        return "Edit-Fehler: 'find' muss nicht-leerer String sein."
    if not isinstance(replace, str):
        return "Edit-Fehler: 'replace' muss String sein."
    try:
        path = safe_path(rel_path)
    except ValueError as e:
        return f"Pfad-Fehler: {e}"
    if not path.exists():
        return f"Datei nicht gefunden: {rel_path}"
    try:
        # File-Size-Check VOR dem Lesen — verhindert OOM bei Riesen-Files
        if path.stat().st_size > EDIT_FILE_MAX_BYTES:
            return f"Datei zu groß für edit_file ({path.stat().st_size} > {EDIT_FILE_MAX_BYTES}B). Manuell editieren."
        content = path.read_text(encoding="utf-8")
        if regex:
            # ReDoS-Schutz: Pattern-Länge + bekannte pathologische Patterns ablehnen
            if len(find) > EDIT_FILE_MAX_REGEX_LEN:
                return f"Regex zu lang ({len(find)} > {EDIT_FILE_MAX_REGEX_LEN}) — vereinfachen oder regex=false."
            for redos_pat in _REDOS_PATTERNS:
                if redos_pat.search(find):
                    return (f"Regex-Pattern '{find[:60]}' enthält pathologisches "
                            "Konstrukt (nested quantifier) — ReDoS-Risiko, abgelehnt. "
                            "Vereinfache das Pattern oder nutze regex=false.")
            try:
                new, n = re.subn(find, replace, content)
            except re.error as e:
                return f"Regex-Syntax-Fehler in '{find[:50]}': {e}"
        else:
            n = content.count(find)
            new = content.replace(find, replace)
        if n == 0:
            return f"Kein Treffer für '{find[:50]}' in {rel_path}"
        atomic_write(path, new)
        # Edit könnte Frontmatter (id/title/aliases) verändert haben → defensiv neu indexieren
        if path.suffix == ".md":
            invalidate_link_index()
        return f"{n}× ersetzt in {rel_path}"
    except Exception as e:
        return f"Edit-Fehler: {e}"


CLIP_URL_TIMEOUT = 15  # Sekunden — verhindert dass slowloris-Server den Bot hängen lassen


def _fetch_url_with_timeout(url: str, timeout: int = CLIP_URL_TIMEOUT) -> Optional[str]:
    """Wie trafilatura.fetch_url, aber mit hartem Socket-Timeout.

    trafilatura's eigener fetch_url akzeptiert kein timeout-Parameter →
    LLM kann clip_url('http://attacker/slowloris') aufrufen und Bot hängt.
    Stdlib urllib hat ein sauberes timeout das durchgereicht wird.
    """
    import urllib.request
    import urllib.error
    req = urllib.request.Request(
        url,
        headers={"User-Agent": "Mozilla/5.0 (compatible; KI-WIKI-Bot/1.0)"},
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.read().decode("utf-8", errors="replace")
    except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError, ValueError) as e:
        log.warning(f"clip_url fetch failed for {url[:80]}: {e}")
        return None


def clip_url(url: str) -> str:
    """Fetch URL (mit Timeout!), save as raw article."""
    if not url or not isinstance(url, str) or not url.startswith(("http://", "https://")):
        return f"Ungültige URL: {url!r}"
    try:
        downloaded = _fetch_url_with_timeout(url)
        if not downloaded:
            return f"Konnte URL nicht laden (Timeout {CLIP_URL_TIMEOUT}s oder Fehler): {url}"
        text = trafilatura.extract(
            downloaded, include_comments=False, include_tables=True,
            include_links=True,
        )
        meta = trafilatura.extract_metadata(downloaded)
        if not text:
            return f"Kein Inhalt extrahierbar: {url}"
        title = (meta.title if meta else None) or "untitled"
        author = (meta.author if meta else None) or "unknown"
        domain = re.sub(r"^www\.", "", re.sub(r"^https?://", "", url).split("/")[0])
        slug = slugify(title)
        today = today_iso()
        filename = f"{today}_{slugify(domain)}_{slug}.md"
        path = RAW_ARTICLES_DIR / filename
        post = frontmatter.Post(
            text,
            id=slugify(f"{domain}-{slug}"),
            title=title,
            type="article",
            source=url,
            author=author,
            captured=today,
            tags=[],
        )
        atomic_write(path, frontmatter.dumps(post) + "\n")
        try:
            append_to_daily(
                "Offen / Einsortieren",
                f"- [[{post['id']}]] *(geclipt: {title})*",
            )
        except Exception:
            pass
        return f"Artikel geclipt: [[{post['id']}]] ({filename})"
    except Exception as e:
        log.exception("clip_url failed")
        return f"Clip-Fehler: {e}"


# ─── Pending Deletions (Multi-File-fähig + Soft/Hard) ───────────────────────
# Two-Step-Delete: request_delete() merkt sich Pfad(e) + Modus,
# confirm_delete() führt aus.
# Modi: 'archive' (Default, sicher) oder 'permanent' (echtes rm).
PENDING_DELETIONS: dict[int, tuple[list[str], float, str]] = {}  # uid → (paths, ts, mode)
DELETE_CONFIRM_TIMEOUT = 300  # Sekunden


def request_delete(rel_path=None, rel_paths=None, permanent: bool = False) -> str:
    """Merkt sich eine Lösch-Anfrage. Akkumuliert.

    rel_path: einzelner Pfad (string) ODER
    rel_paths: Liste von Pfaden (Bulk).
    permanent: False (Default) = ins 99_Archive/ verschieben (reversibel).
               True = WIRKLICH LÖSCHEN (irreversibel via rm).

    Beide Parameter sind optional damit das Tool-Schema kein oneOf braucht
    (Gemini/Ollama lehnen oneOf/anyOf in Tool-Params ab).
    """
    # Input normalisieren — beide Parameter zusammenführen
    paths_in: list[str] = []
    if rel_path is not None:
        if isinstance(rel_path, str):
            paths_in.append(rel_path)
        elif isinstance(rel_path, list):
            # LLM hat trotz Schema doch Liste in rel_path gepackt → tolerieren
            paths_in.extend(rel_path)
        else:
            return f"request_delete: rel_path ungültiger Typ {type(rel_path)}"
    if rel_paths is not None:
        if isinstance(rel_paths, list):
            paths_in.extend(rel_paths)
        elif isinstance(rel_paths, str):
            # LLM hat trotz Schema String in rel_paths gepackt → tolerieren
            paths_in.append(rel_paths)
        else:
            return f"request_delete: rel_paths ungültiger Typ {type(rel_paths)}"
    if not paths_in:
        return "request_delete: keine Pfade angegeben (rel_path oder rel_paths erforderlich)"

    # Pfade validieren + sammeln
    valid = []
    errors = []
    for rp in paths_in:
        try:
            p = safe_path(rp)
            if not p.exists():
                errors.append(f"nicht gefunden: {rp}")
            else:
                valid.append(rp)
        except Exception as e:
            errors.append(f"{rp} ({e})")

    if not valid:
        return "Keine gültigen Pfade. " + "; ".join(errors)

    mode = "permanent" if permanent else "archive"

    # Akkumulieren — Mode darf nicht silent gewechselt werden
    existing = PENDING_DELETIONS.get(ALLOWED_USER_ID)
    if existing and existing[2] != mode:
        return (f"Konflikt: pending Löschung läuft im Modus '{existing[2]}', "
                f"aber neue Anfrage will '{mode}'. Erst confirm_delete oder cancel_delete.")

    existing_paths = existing[0] if existing else []
    combined = list(dict.fromkeys(existing_paths + valid))
    PENDING_DELETIONS[ALLOWED_USER_ID] = (combined, time.time(), mode)
    log.info(f"Pending delete ({mode}): {combined}")

    if mode == "permanent":
        msg = f"⚠️⚠️ <b>ENDGÜLTIG LÖSCHEN</b> — soll(en) {len(combined)} Datei(en) "
        msg += "<b>UNWIDERRUFLICH gelöscht</b> werden? (kein Archiv, kein Restore!)\n\n"
    else:
        msg = f"⚠️ Bestätigung: soll(en) {len(combined)} Datei(en) ins Archiv verschoben werden? (reversibel)\n\n"
    msg += "\n".join(f"• <code>{p}</code>" for p in combined)
    if errors:
        msg += "\n\n<i>(übersprungen: " + ", ".join(errors) + ")</i>"
    msg += f"\n\nAntworte mit 'ja' / 'bestätigt' / 'machs' (innerhalb {DELETE_CONFIRM_TIMEOUT//60} Min)."
    return msg


def confirm_delete(action: str = "confirm") -> str:
    """Führt pending Löschungen aus oder bricht ab je nach action.

    action='confirm' (default): Löschen ausführen (Archiv oder permanent
                                je nach Mode aus request_delete)
    action='cancel': Pending verwerfen ohne zu löschen.
    """
    pending = PENDING_DELETIONS.get(ALLOWED_USER_ID)
    if not pending or not pending[0]:
        return "Keine Löschung pending — gibts nichts zu bestätigen."
    paths, ts, mode = pending
    # Cancel-Pfad: einfach pending leeren
    if action == "cancel":
        PENDING_DELETIONS.pop(ALLOWED_USER_ID, None)
        return f"Löschanfrage für {len(paths)} Datei(en) abgebrochen ({mode}-Modus)."
    age = time.time() - ts
    if age > DELETE_CONFIRM_TIMEOUT:
        PENDING_DELETIONS.pop(ALLOWED_USER_ID, None)
        return f"Bestätigung zu spät ({int(age)}s > {DELETE_CONFIRM_TIMEOUT}s). Bitte nochmal anfordern."

    results = []

    if mode == "permanent":
        # Hart löschen via rm
        for rel_path in paths:
            try:
                src = safe_path(rel_path)
                if not src.exists():
                    results.append(f"✗ {rel_path} schon weg")
                    continue
                if src.is_dir():
                    shutil.rmtree(src)
                else:
                    src.unlink()
                results.append(f"💀 {rel_path}")
                log.info(f"Hard-deleted: {rel_path}")
            except Exception as e:
                log.exception(f"hard-delete {rel_path} failed")
                results.append(f"✗ {rel_path} ({e})")
        header = f"💀 ENDGÜLTIG GELÖSCHT ({len(paths)} Datei(en)):"
    else:
        # Soft: nach 99_Archive/ verschieben
        archive_root = VAULT / "99_Archive"
        for rel_path in paths:
            try:
                src = safe_path(rel_path)
                if not src.exists():
                    results.append(f"✗ {rel_path} verschwunden")
                    continue
                dst = archive_root / src.relative_to(VAULT)
                dst.parent.mkdir(parents=True, exist_ok=True)
                if dst.exists():
                    ts_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    dst = dst.with_name(f"{dst.stem}_{ts_str}{dst.suffix}")
                os.rename(src, dst)
                results.append(f"✓ {rel_path}")
                log.info(f"Archived: {rel_path} → {dst.relative_to(VAULT)}")
            except Exception as e:
                log.exception(f"archive {rel_path} failed")
                results.append(f"✗ {rel_path} ({e})")
        header = f"Verschoben nach 99_Archive/ ({len(paths)} Datei(en)):"

    PENDING_DELETIONS.pop(ALLOWED_USER_ID, None)
    invalidate_link_index()  # gelöschte/archivierte IDs nicht mehr verlinkbar
    return f"{header}\n" + "\n".join(results)


# ─── Reminders (persistent + JobQueue) ──────────────────────────────────────
# Reminders überleben Bot-Restart: JSON in 06_Meta/reminders.json.
# Bei Startup werden alle aktiven Reminders neu in die JobQueue eingehängt.
REMINDERS_FILE = VAULT / "06_Meta" / "reminders.json"
ACTIVE_REMINDER_JOBS: dict = {}  # id → telegram.ext.Job
BOT_APP = None  # wird in main() gesetzt — brauchen Zugriff auf job_queue von Tools aus


def _load_reminders() -> list:
    if not REMINDERS_FILE.exists():
        return []
    try:
        return json.loads(REMINDERS_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning(f"reminders.json kaputt: {e}")
        return []


def _save_reminders(reminders: list) -> None:
    REMINDERS_FILE.parent.mkdir(parents=True, exist_ok=True)
    REMINDERS_FILE.write_text(
        json.dumps(reminders, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


# Reminder-Recurrence (kein monthly — Reminders sind zeitpunkt-basiert,
# monthly würde unklar bei "31. + Februar"-Edge-Cases). Tasks haben eigene Konstante.
VALID_REMINDER_RECURRENCE = {None, "", "daily", "weekly", "weekdays"}


def create_reminder(when_iso: str, message: str, recurrence: Optional[str] = None) -> str:
    """Setzt eine Erinnerung. Wird zur angegebenen Zeit als Telegram-Nachricht geschickt.

    when_iso: ISO-Datetime YYYY-MM-DDTHH:MM:SS (Lokalzeit Europe/Vienna)
    recurrence: null/leer = einmalig, "daily" = täglich, "weekdays" = Mo-Fr,
                "weekly" = einmal pro Woche (gleicher Wochentag wie der erste Trigger)
    """
    if not message or not message.strip():
        return "Erinnerung-Text fehlt."
    try:
        when_dt = datetime.fromisoformat(when_iso)
    except ValueError:
        return f"Ungültiges Datum/Zeit: {when_iso}. Format: 2026-04-26T15:00:00"

    if when_dt.tzinfo is None:
        when_dt = when_dt.replace(tzinfo=TIMEZONE)

    rec = (recurrence or None) if recurrence else None
    if rec not in VALID_REMINDER_RECURRENCE:
        return f"Ungültige Recurrence '{recurrence}'. Erlaubt: leer / daily / weekly / weekdays"

    # Ein-shot Reminder in der Vergangenheit ablehnen (außer recurring)
    if not rec and when_dt < datetime.now(TIMEZONE):
        return f"Zeitpunkt liegt in der Vergangenheit: {when_dt.isoformat(timespec='minutes')}"

    rid = "rem-" + datetime.now().strftime("%Y%m%d-%H%M%S-%f")[:-3]
    reminder = {
        "id": rid,
        "fire_at": when_dt.isoformat(timespec="seconds"),
        "message": message.strip(),
        "recurrence": rec,
        "created": datetime.now(TIMEZONE).isoformat(timespec="seconds"),
    }
    reminders = _load_reminders()
    reminders.append(reminder)
    _save_reminders(reminders)

    if BOT_APP is not None:
        _schedule_reminder(BOT_APP, reminder)

    invalidate_today_data_cache()  # neuer Reminder kann heute feuern → Agenda neu

    rec_str = f", wiederholt {rec}" if rec else ""
    when_str = when_dt.strftime("%a %d.%m. %H:%M")
    return f"⏰ Erinnerung gesetzt: {when_str}{rec_str}\n→ \"{message.strip()[:80]}\""


def _schedule_reminder(app, reminder: dict) -> None:
    """Hängt einen Reminder in die JobQueue ein."""
    when_dt = datetime.fromisoformat(reminder["fire_at"])
    if when_dt.tzinfo is None:
        when_dt = when_dt.replace(tzinfo=TIMEZONE)

    rid = reminder["id"]
    rec = reminder.get("recurrence")
    job_data = {"id": rid, "message": reminder["message"]}

    try:
        if rec == "daily":
            job = app.job_queue.run_daily(
                reminder_callback,
                time=when_dt.timetz(),
                data=job_data,
                name=f"reminder-{rid}",
            )
        elif rec == "weekly":
            job = app.job_queue.run_daily(
                reminder_callback,
                time=when_dt.timetz(),
                days=(when_dt.weekday(),),
                data=job_data,
                name=f"reminder-{rid}",
            )
        elif rec == "weekdays":
            job = app.job_queue.run_daily(
                reminder_callback,
                time=when_dt.timetz(),
                days=(0, 1, 2, 3, 4),
                data=job_data,
                name=f"reminder-{rid}",
            )
        else:
            # Einmalig
            if when_dt <= datetime.now(TIMEZONE):
                log.info(f"Reminder {rid} liegt in Vergangenheit, skip.")
                # aus JSON entfernen
                _remove_reminder_from_json(rid)
                return
            job = app.job_queue.run_once(
                reminder_callback,
                when=when_dt,
                data=job_data,
                name=f"reminder-{rid}",
            )
        ACTIVE_REMINDER_JOBS[rid] = job
        log.info(f"Reminder scheduled: {rid} @ {when_dt.isoformat()} rec={rec}")
    except Exception as e:
        log.exception(f"Reminder-Schedule fehlgeschlagen für {rid}")


def _remove_reminder_from_json(rid: str) -> None:
    reminders = _load_reminders()
    reminders = [r for r in reminders if r["id"] != rid]
    _save_reminders(reminders)


async def reminder_callback(ctx: ContextTypes.DEFAULT_TYPE):
    """Wird von JobQueue ausgelöst wenn ein Reminder fällig wird."""
    data = ctx.job.data
    rid = data["id"]
    message = data["message"]
    try:
        await ctx.bot.send_message(
            chat_id=ALLOWED_USER_ID,
            text=f"⏰ <b>Erinnerung</b>\n\n{_esc_html(message)}",
            parse_mode=constants.ParseMode.HTML,
        )
        log.info(f"Reminder fired: {rid}")
    except Exception as e:
        log.exception(f"Reminder-Send fehlgeschlagen für {rid}")

    # Einmalige Reminder aus JSON + ACTIVE_JOBS entfernen
    reminders = _load_reminders()
    reminder = next((r for r in reminders if r["id"] == rid), None)
    if reminder and not reminder.get("recurrence"):
        _remove_reminder_from_json(rid)
        ACTIVE_REMINDER_JOBS.pop(rid, None)


def list_reminders() -> str:
    """Liste aller aktiven Reminders."""
    reminders = _load_reminders()
    if not reminders:
        return "Keine aktiven Erinnerungen."
    lines = [f"⏰ {len(reminders)} aktive Erinnerung(en):"]
    for r in sorted(reminders, key=lambda x: x.get("fire_at", "")):
        when = datetime.fromisoformat(r["fire_at"])
        when_str = when.strftime("%a %d.%m. %H:%M")
        rec = f" 🔁 {r['recurrence']}" if r.get("recurrence") else ""
        lines.append(f"• `{r['id']}` — {when_str}{rec}\n  {r['message'][:80]}")
    return "\n".join(lines)


def cancel_reminder(reminder_id: str) -> str:
    """Bricht einen Reminder ab (per ID, z.B. 'rem-20260426-153000-123')."""
    reminders = _load_reminders()
    found = next((r for r in reminders if r["id"] == reminder_id), None)
    if not found:
        return f"Erinnerung nicht gefunden: {reminder_id}"
    _remove_reminder_from_json(reminder_id)
    job = ACTIVE_REMINDER_JOBS.pop(reminder_id, None)
    if job:
        try:
            job.schedule_removal()
        except Exception:
            pass
    invalidate_today_data_cache()
    return f"✓ Erinnerung gecancelt: {found['message'][:60]}"


def backup_vault() -> str:
    """Manuelles Backup: Vault in privates GitHub-Repo pushen.

    Voraussetzung: GITHUB_BACKUP_REPO (z.B. 'julasim/KI_WIKI_Vault_Backup') und
    GITHUB_BACKUP_TOKEN (PAT mit Repo-Schreibrechten) in .env.

    Workflow: clone-on-first-use → rsync vault-content → commit → push
    """
    repo = os.environ.get("GITHUB_BACKUP_REPO", "").strip()
    token = os.environ.get("GITHUB_BACKUP_TOKEN", "").strip()
    if not repo or not token:
        return ("Backup nicht konfiguriert. Setze GITHUB_BACKUP_REPO und "
                "GITHUB_BACKUP_TOKEN in /opt/bot/.env, dann docker compose restart.")

    backup_dir = Path("/vault-backup")
    # Token URL-encodieren falls Sonderzeichen drin (defensive — GitHub-PATs sind
    # zwar normalerweise ASCII-alphanumerisch, aber sicher ist sicher)
    from urllib.parse import quote
    safe_token = quote(token, safe="")
    repo_url = f"https://x-access-token:{safe_token}@github.com/{repo}.git"

    def _run(cmd, cwd=None, env_extra=None):
        # Minimal-env: git/rsync brauchen nur PATH + lokale Vars. Vermeidet
        # dass TG_TOKEN/LLM_API_KEY/etc an Subprozesse vererbt werden (die
        # könnten in error-output landen oder von kompromittiertem git-binary
        # exfiltriert werden).
        env = {k: v for k, v in os.environ.items()
               if k in ("PATH", "HOME", "USER", "LANG", "LC_ALL", "TERM",
                        "GIT_AUTHOR_NAME", "GIT_AUTHOR_EMAIL",
                        "GIT_COMMITTER_NAME", "GIT_COMMITTER_EMAIL")}
        if env_extra:
            env.update(env_extra)
        return subprocess.run(cmd, cwd=cwd, env=env, capture_output=True, text=True, timeout=60)

    try:
        # Clone-on-first-use oder Pull
        if not (backup_dir / ".git").exists():
            backup_dir.mkdir(parents=True, exist_ok=True)
            log.info(f"Cloning backup repo: {repo}")
            r = _run(["git", "clone", repo_url, str(backup_dir)])
            if r.returncode != 0:
                # Sanitize: Token aus Fehler entfernen falls drin
                err = r.stderr.replace(token, "***")
                return f"Clone-Fehler: {err}"
        else:
            r = _run(["git", "pull", "--rebase"], cwd=backup_dir)
            if r.returncode != 0 and "no upstream" not in r.stderr.lower():
                err = r.stderr.replace(token, "***")
                # Pull-Fehler nicht fatal — wir pushen trotzdem
                log.warning(f"Pull warning: {err}")

        # Git-Identity setzen (idempotent, lokal zum Repo)
        _run(["git", "config", "user.email", "bot@ki-wiki.local"], cwd=backup_dir)
        _run(["git", "config", "user.name", "KI Wiki Bot"], cwd=backup_dir)

        # Vault-Content rüber — Pre-Flight: Source muss existieren + Files haben,
        # sonst würde rsync --delete uns das Backup-Repo leeren
        if not VAULT.exists() or not VAULT.is_dir():
            return f"Backup-Abbruch: Vault-Source {VAULT} nicht erreichbar (Mount kaputt?)"
        if not any(VAULT.iterdir()):
            return f"Backup-Abbruch: Vault-Source {VAULT} ist leer — nichts zu sichern, schütze gegen Daten-Verlust im Backup."

        target = backup_dir / "vault"
        target.mkdir(exist_ok=True)
        r = _run([
            "rsync", "-a", "--delete",
            "--exclude=.obsidian/workspace*",
            "--exclude=.obsidian/cache",
            "--exclude=*.tmp",
            "--exclude=__pycache__",
            f"{VAULT}/", f"{target}/",
        ])
        if r.returncode != 0:
            return f"Rsync-Fehler: {r.stderr}"

        # Stage + Commit (nur wenn Änderungen)
        _run(["git", "add", "-A"], cwd=backup_dir)
        diff = _run(["git", "diff", "--cached", "--quiet"], cwd=backup_dir)
        if diff.returncode == 0:
            return "✓ Keine Änderungen seit letztem Backup."

        # Commit
        commit_msg = f"Backup {datetime.now().isoformat(timespec='seconds')}"
        r = _run(["git", "commit", "-m", commit_msg], cwd=backup_dir)
        if r.returncode != 0:
            return f"Commit-Fehler: {r.stderr}"

        # Push
        r = _run(["git", "push"], cwd=backup_dir)
        if r.returncode != 0:
            err = r.stderr.replace(token, "***")
            return f"Push-Fehler: {err}"

        # Stats
        hash_r = _run(["git", "rev-parse", "--short", "HEAD"], cwd=backup_dir)
        commit_hash = hash_r.stdout.strip()
        # Datei-Anzahl im Backup
        count_r = _run(["bash", "-c", f"find '{target}' -type f -name '*.md' | wc -l"])
        file_count = count_r.stdout.strip() or "?"

        return (f"✓ Backup gepusht\n"
                f"Repo: {repo}@{commit_hash}\n"
                f"Files: {file_count} .md\n"
                f"Zeit: {datetime.now().strftime('%H:%M:%S')}")

    except subprocess.TimeoutExpired:
        return "Backup-Fehler: Timeout (>60s). Repo zu groß oder Netz langsam?"
    except Exception as e:
        log.exception("backup_vault failed")
        return f"Backup-Fehler: {e}"


def list_existing_tags(top_n: int = 30) -> str:
    """Liste der existierenden Tags im Vault, sortiert nach Häufigkeit.

    Hilft dem LLM beim Auto-Tagging konsistent zu bleiben statt jeden
    Eintrag neu zu erfinden ('arbeit' vs 'work' vs 'job' Drift vermeiden).
    """
    from collections import Counter
    counter: Counter = Counter()
    for _, post in iter_vault_md():
        tags = post.get("tags") or []
        if isinstance(tags, list):
            for t in tags:
                if isinstance(t, str) and t.strip():
                    counter[t.strip().lower()] += 1
    if not counter:
        return "Keine Tags im Vault — du legst die Konvention fest."
    top = counter.most_common(top_n)
    lines = [f"Top {len(top)} bestehende Tags ({sum(counter.values())} Verwendungen gesamt):"]
    for tag, count in top:
        lines.append(f"• `{tag}` ({count}×)")
    return "\n".join(lines)


def create_project(name: str, description: str = "", area: Optional[str] = None,
                   parent: Optional[str] = None) -> str:
    """Legt einen neuen Projekt-Ordner unter 05_Projects/<slug>/ an.

    Erzeugt Ordner + README.md mit Dataview-Queries die alle Tasks/Notes mit
    'project: <slug>' im Frontmatter sammeln. So können Files weiterhin in den
    Standard-Ordnern (10_Life/tasks/, /notes/) liegen, sind aber im Projekt-
    Container automatisch gelistet.

    parent: optional Slug eines existierenden Projekts → neues Projekt wird als
    Subprojekt unter dem Parent angelegt (05_Projects/<parent>/<slug>/).
    """
    if not name or not name.strip():
        return "Fehler: Projekt-Name darf nicht leer sein."
    slug = slugify(name)
    if not slug or slug == "untitled":
        return f"Fehler: Slug aus '{name}' nicht ableitbar — bitte aussagekräftigeren Namen wählen."

    PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

    # Parent auflösen falls Subprojekt
    if parent:
        parent_dir = find_project_dir(parent)
        if parent_dir is None:
            return f"Fehler: Parent-Projekt '{parent}' nicht gefunden (oder mehrdeutig)."
        proj_dir = parent_dir / slug
    else:
        proj_dir = PROJECTS_DIR / slug

    # Existenz-Check rekursiv (verhindert Doppel-Slugs an verschiedenen Orten)
    existing = find_project_dir(slug)
    if existing is not None:
        rel = existing.relative_to(VAULT).as_posix()
        return f"Projekt existiert bereits: [[project-{slug}]] (`{rel}/`)"

    proj_dir.mkdir(parents=True, exist_ok=False)
    readme_path = proj_dir / "README.md"
    today = today_iso()

    desc_block = description.strip() if description else f"Projekt-Container für **{name}**. Tasks/Notes mit `project: {slug}` im Frontmatter werden hier automatisch gelistet."

    body = f"""# {name}

{desc_block}

## Status
- **Status**: active
- **Gestartet**: {today}{chr(10) + f"- **Area**: [[{area}]]" if area else ""}

## Offene Tasks
```dataview
TABLE WITHOUT ID file.link AS Task, priority AS Prio, due AS Fällig
FROM "10_Life/tasks"
WHERE project = "{slug}" AND status != "done" AND status != "cancelled"
SORT priority DESC, due ASC
```

## Notizen
```dataview
LIST
FROM "10_Life/notes"
WHERE project = "{slug}"
SORT file.ctime DESC
```

## Meetings
```dataview
TABLE WITHOUT ID file.link AS Meeting, date AS Datum
FROM "10_Life/meetings"
WHERE project = "{slug}"
SORT date DESC
```

## Log
- {today}: Projekt angelegt
"""

    post_data = {
        "id": f"project-{slug}",
        "title": name,
        "type": "project",
        "started": today,
        "status": "active",
        "tags": [],
    }
    if area:
        post_data["area"] = area
    post = frontmatter.Post(body, **post_data)
    atomic_write(readme_path, frontmatter.dumps(post) + "\n")

    # CONTEXT.md leer initialisieren — kann später via update_project_context befüllt werden
    context_path = proj_dir / "CONTEXT.md"
    context_path.write_text(
        f"# Kontext: {slug}\n\n"
        "_Projekt-spezifische Regeln/Infos (Auftraggeber, Tech-Stack, Frist, Budget). "
        f"Wird automatisch in den Bot-Prompt geladen wenn `activate_project({slug})` aktiv._\n\n"
        "_(noch leer — fülle via Bot-Tool `update_project_context` oder direkt in Obsidian)_\n",
        encoding="utf-8",
    )

    invalidate_link_index()

    rel = proj_dir.relative_to(VAULT).as_posix()
    parent_info = f" (Subprojekt von `{parent}`)" if parent else ""
    return (f"✓ Projekt angelegt: [[project-{slug}]]{parent_info}\n"
            f"Ordner: `{rel}/` (README + CONTEXT.md)\n"
            f"Tipp: `activate_project {slug}` schaltet projektspez. Kontext im Bot scharf.")


# Noise-Verzeichnisse die in list_files-Output rausgefiltert werden
# (System-Internals, Templates, Meta-Reports, Tooling — nicht User-Content)
LIST_FILES_NOISE_DIRS = {
    ".obsidian", ".trash", "99_Archive",
    "08_Templates", "06_Meta", "07_Tools",
}
LIST_FILES_NOISE_FILES = {
    "README.md", "_index.md",
    "CLAUDE.md", "COMMANDS.md", "MOC.md",
    "PIPELINES.md", "SCHEMA.md",
}


def list_files(rel_dir: str = "", include_system: bool = False) -> str:
    """Liste alle .md-Files in einem Vault-Unterordner.

    Standardmäßig werden System-Verzeichnisse (Templates, Meta, Tools, Trash, Archive)
    und System-Docs (CLAUDE.md etc.) ausgefiltert — User sieht nur eigenen Content.
    Mit include_system=True wird alles gezeigt (für Debug).
    """
    try:
        base = safe_path(rel_dir) if rel_dir else VAULT
        if not base.exists() or not base.is_dir():
            return f"Verzeichnis nicht gefunden: {rel_dir}"

        def is_visible(p: Path) -> bool:
            if include_system:
                return True
            if p.name in LIST_FILES_NOISE_FILES:
                return False
            return not any(part in LIST_FILES_NOISE_DIRS for part in p.parts)

        files = sorted(p for p in base.rglob("*.md") if is_visible(p))
        if not files:
            return f"Keine User-Files in {rel_dir or 'Vault-Root'} (System-Files via include_system=true sichtbar)."

        rels = [str(f.relative_to(VAULT)).replace("\\", "/") for f in files]

        # Bei vielen Files: nach Top-Level-Ordner gruppieren für lesbare Ausgabe
        if len(rels) > 12:
            from collections import defaultdict
            grouped = defaultdict(list)
            for r in rels:
                top = r.split("/", 1)[0]
                grouped[top].append(r)
            lines = [f"{len(rels)} Files in {rel_dir or 'Vault-Root'}, gruppiert:"]
            for top in sorted(grouped):
                lines.append(f"\n**{top}/** ({len(grouped[top])})")
                for r in grouped[top][:8]:
                    lines.append(f"• `{r}`")
                if len(grouped[top]) > 8:
                    lines.append(f"  _… {len(grouped[top])-8} weitere_")
            return "\n".join(lines)

        return f"{len(rels)} Files in {rel_dir or 'Vault-Root'}:\n" + "\n".join(f"• `{r}`" for r in rels)
    except Exception as e:
        return f"List-Fehler: {e}"


# ─── Conversation Memory (3-Tier) ────────────────────────────────────────────
# Tier 1: RAM-Cache (letzte HISTORY_MAX_MESSAGES, schneller Zugriff)
# Tier 2: Persistent JSONL (überlebt Restart, lazy-loaded)
# Tier 3: Facts-File (long-term Fakten über User, always im System-Prompt)

CONVERSATION_HISTORY: dict[int, list] = {}
CONVERSATION_TIMESTAMPS: dict[int, float] = {}
HISTORY_MAX_MESSAGES = 60       # ca. 30 User+Assistant-Turns
HISTORY_TIMEOUT = 60 * 60       # 1h Inaktivität → RAM-Cache leeren, beim nächsten Zugriff von Disk lazy-laden
HISTORY_PERSIST_LIMIT = 1000    # max Lines im JSONL bevor compaction
HISTORY_COMPACT_KEEP = 200       # nach compact: behalte letzte N Lines

# Lock gegen Race-Conditions zwischen User-Messages und nightly_suggestion_job
# (asyncio-cooperative — kein echtes Threading, aber sauber)
_HISTORY_LOCK = asyncio.Lock()

BOT_MEMORY_DIR = VAULT / "06_Meta" / "bot-memory"
FACTS_FILE = BOT_MEMORY_DIR / "facts.md"
PREFERENCES_FILE = BOT_MEMORY_DIR / "preferences.md"
ACTIVE_PROJECT_FILE = BOT_MEMORY_DIR / "active-project.txt"
HISTORY_FILE = BOT_MEMORY_DIR / "conversation-history.jsonl"
CORRECTIONS_FILE = BOT_MEMORY_DIR / "corrections.jsonl"
PENDING_SUGGESTIONS_FILE = BOT_MEMORY_DIR / "pending-suggestions.json"


def _ensure_memory_dir():
    BOT_MEMORY_DIR.mkdir(parents=True, exist_ok=True)


def get_facts() -> str:
    """Long-term Fakten lesen (werden in System-Prompt eingespeist)."""
    if not FACTS_FILE.exists():
        return ""
    try:
        content = FACTS_FILE.read_text(encoding="utf-8").strip()
        # Frontmatter überspringen falls vorhanden
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                content = parts[2].strip()
        return content
    except Exception as e:
        log.warning(f"facts file read failed: {e}")
        return ""


def remember(fact: str) -> str:
    """Fügt einen persistenten Fakt zur Memory-Datei hinzu."""
    if not fact or not fact.strip():
        return "Fakt-Text fehlt."
    fact = fact.strip()
    _ensure_memory_dir()

    # Initialize file with header if needed
    if not FACTS_FILE.exists():
        FACTS_FILE.write_text(
            "# Bot-Memory: persistente Fakten\n\n"
            "_Hier sammelt der Bot Fakten die er sich dauerhaft merken soll. "
            "Du kannst manuell editieren — Änderungen sind beim nächsten LLM-Call wirksam._\n\n",
            encoding="utf-8",
        )

    today = today_iso()
    line = f"- ({today}) {fact}\n"
    with FACTS_FILE.open("a", encoding="utf-8") as f:
        f.write(line)
    log.info(f"Remembered fact: {fact[:80]}")
    return f"✓ Gemerkt: {fact[:120]}"


def list_facts() -> str:
    """Zeigt alle gemerkten Fakten."""
    facts = get_facts()
    if not facts:
        return "Keine persistenten Fakten gespeichert."
    return f"Persistente Fakten:\n\n{facts[:3000]}"


def forget_fact(pattern: str) -> str:
    """Entfernt Fakten die `pattern` enthalten (case-insensitive)."""
    if not FACTS_FILE.exists():
        return "Keine Fakten-Datei."
    if not pattern or not pattern.strip():
        return "Such-Text fehlt."
    needle = pattern.strip().lower()
    content = FACTS_FILE.read_text(encoding="utf-8")
    lines = content.splitlines()
    kept = []
    removed = []
    for line in lines:
        if line.startswith("- ") and needle in line.lower():
            removed.append(line)
        else:
            kept.append(line)
    if not removed:
        return f"Kein Fakt gefunden mit '{pattern}'."
    atomic_write(FACTS_FILE, "\n".join(kept) + "\n")
    return f"Entfernt ({len(removed)}):\n" + "\n".join(removed[:5])


# ─── Präferenzen (Stil/Tonalität — wie der Bot reden soll) ──────────────────

def _strip_md_intro(content: str) -> str:
    """Entfernt erste H1-Überschrift + Italic-Erklärung am Datei-Anfang."""
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            content = parts[2].strip()
    lines = content.splitlines()
    # H1 raus
    while lines and (lines[0].startswith("# ") or lines[0].strip() == ""):
        lines.pop(0)
    # Italic-Doku (z.B. _Hier sammelt..._) raus
    while lines and lines[0].strip().startswith("_") and lines[0].strip().endswith("_"):
        lines.pop(0)
        # Leerzeile danach
        while lines and lines[0].strip() == "":
            lines.pop(0)
    return "\n".join(lines).strip()


def get_preferences() -> str:
    """Lese Präferenzen-Inhalt (für System-Prompt-Injection)."""
    if not PREFERENCES_FILE.exists():
        return ""
    try:
        return _strip_md_intro(PREFERENCES_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning(f"preferences read failed: {e}")
        return ""


def set_preference(text: str) -> str:
    """Fügt Präferenz hinzu (z.B. 'Antworte direkt ohne Floskeln')."""
    if not text or not text.strip():
        return "Präferenz-Text fehlt."
    text = text.strip()
    _ensure_memory_dir()
    if not PREFERENCES_FILE.exists():
        PREFERENCES_FILE.write_text(
            "# Präferenzen — wie der Bot mir antworten soll\n\n"
            "_Stil, Tonalität, Format. Werden bei jedem LLM-Call automatisch in den System-Prompt eingespeist._\n\n",
            encoding="utf-8",
        )
    today = today_iso()
    line = f"- ({today}) {text}\n"
    with PREFERENCES_FILE.open("a", encoding="utf-8") as f:
        f.write(line)
    return f"✓ Präferenz gemerkt: {text[:120]}"


def list_preferences() -> str:
    prefs = get_preferences()
    if not prefs:
        return "Keine Präferenzen gespeichert."
    return f"Präferenzen:\n\n{prefs[:2000]}"


def forget_preference(pattern: str) -> str:
    if not PREFERENCES_FILE.exists():
        return "Keine Präferenzen-Datei."
    if not pattern or not pattern.strip():
        return "Such-Text fehlt."
    needle = pattern.strip().lower()
    content = PREFERENCES_FILE.read_text(encoding="utf-8")
    lines = content.splitlines()
    kept, removed = [], []
    for line in lines:
        if line.startswith("- ") and needle in line.lower():
            removed.append(line)
        else:
            kept.append(line)
    if not removed:
        return f"Keine Präferenz mit '{pattern}'."
    atomic_write(PREFERENCES_FILE, "\n".join(kept) + "\n")
    return f"Entfernt ({len(removed)}):\n" + "\n".join(removed[:5])


def forget(kind: str, pattern: str) -> str:
    """Vereinheitlichtes Forget für Memory: kind='fact' oder 'preference'.

    Konsolidiert forget_fact + forget_preference zu einem Tool.
    Der LLM-Agent muss nur noch entscheiden welches Memory-Tier.
    """
    k = (kind or "").strip().lower()
    if k in ("fact", "facts", "f"):
        return forget_fact(pattern)
    if k in ("preference", "preferences", "pref", "p"):
        return forget_preference(pattern)
    return f"Unbekanntes kind '{kind}'. Erlaubt: 'fact' oder 'preference'."


# ─── Projekt-Kontext (per-Projekt CONTEXT.md, on-demand) ────────────────────

def get_active_project() -> Optional[str]:
    """Slug des aktuell-aktiven Projekts (oder None)."""
    if not ACTIVE_PROJECT_FILE.exists():
        return None
    try:
        slug = ACTIVE_PROJECT_FILE.read_text(encoding="utf-8").strip()
        return slug or None
    except Exception:
        return None


def get_project_context(slug: str) -> str:
    """Liest CONTEXT.md eines Projekts (rekursive Suche unter 05_Projects/)."""
    proj_dir = find_project_dir(slug)
    if proj_dir is None:
        return ""
    context_file = proj_dir / "CONTEXT.md"
    if not context_file.exists():
        return ""
    try:
        return _strip_md_intro(context_file.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning(f"project context read failed for {slug}: {e}")
        return ""


def activate_project(slug: str) -> str:
    """Setzt ein Projekt als aktiv — sein CONTEXT.md wird automatisch im Prompt geladen."""
    if not slug or not slug.strip():
        return "Projekt-Slug fehlt."
    slug = slug.strip().lower()
    if slug.startswith("project-"):
        slug = slug[len("project-"):]
    proj_dir = find_project_dir(slug)
    if proj_dir is None:
        return f"Projekt nicht gefunden: {slug}"
    _ensure_memory_dir()
    ACTIVE_PROJECT_FILE.write_text(slug, encoding="utf-8")
    ctx = get_project_context(slug)
    ctx_info = f" (CONTEXT.md: {len(ctx)} Zeichen)" if ctx else " (noch keine CONTEXT.md)"
    return f"✓ Projekt aktiviert: [[project-{slug}]]{ctx_info}"


def deactivate_project() -> str:
    """Bricht aktives Projekt ab — CONTEXT.md wird nicht mehr geladen."""
    was = get_active_project()
    if ACTIVE_PROJECT_FILE.exists():
        ACTIVE_PROJECT_FILE.unlink()
    return f"Aktives Projekt zurückgesetzt (war: {was or 'keins'})."


def update_project_context(slug: str, text: str, mode: str = "append") -> str:
    """Update CONTEXT.md eines Projekts. mode: 'append' (default) oder 'replace'."""
    if not slug or not slug.strip():
        return "Slug fehlt."
    slug = slug.strip().lower()
    if slug.startswith("project-"):
        slug = slug[len("project-"):]
    proj_dir = find_project_dir(slug)
    if proj_dir is None:
        return f"Projekt nicht gefunden: {slug}"
    if mode not in ("append", "replace"):
        mode = "append"

    # Auto-Link bekannte Vault-IDs/Titles im Kontext-Text — exclude das Projekt selbst
    linked_text = auto_link(text.strip(), exclude_ids={slug, f"project-{slug}"})

    context_file = proj_dir / "CONTEXT.md"
    header = (
        f"# Kontext: {slug}\n\n"
        "_Projekt-spezifischer Kontext für den Bot. Nur aktiv wenn Projekt via "
        "`activate_project` aktiviert ist._\n\n"
    )

    if mode == "replace" or not context_file.exists():
        context_file.write_text(header + linked_text + "\n", encoding="utf-8")
        return f"✓ CONTEXT.md für {slug} {'ersetzt' if mode == 'replace' else 'angelegt'}"
    else:
        with context_file.open("a", encoding="utf-8") as f:
            f.write(f"\n{linked_text}\n")
        return f"✓ CONTEXT.md für {slug} erweitert"


def project_context(action: str, slug: Optional[str] = None,
                    text: Optional[str] = None, mode: str = "append") -> str:
    """Vereinheitlichtes Projekt-Kontext-Tool.

    action='activate' (slug nötig)   → setzt Projekt aktiv, lädt CONTEXT.md
    action='deactivate' (slug egal)  → bricht aktives Projekt ab
    action='update' (slug+text nötig)→ schreibt in CONTEXT.md (mode=append/replace)

    Konsolidiert die 3 Einzel-Tools (activate_project, deactivate_project,
    update_project_context) zu einem.
    """
    a = (action or "").strip().lower()
    if a in ("activate", "aktivier", "an"):
        if not slug:
            return "Slug fehlt für activate."
        return activate_project(slug)
    if a in ("deactivate", "deaktivier", "aus", "stop"):
        return deactivate_project()
    if a in ("update", "edit", "set"):
        if not slug or not text:
            return "Slug und text nötig für update."
        return update_project_context(slug, text, mode)
    return f"Unbekannte action '{action}'. Erlaubt: activate / deactivate / update."


# ─── Korrektur-Log (Trainings-Material für späteres Fine-Tuning) ─────────────

# ─── Nightly Memory-Vorschläge (Hybrid: LLM extrahiert, User approved) ─────

def _load_pending_suggestions() -> list:
    if not PENDING_SUGGESTIONS_FILE.exists():
        return []
    try:
        return json.loads(PENDING_SUGGESTIONS_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []


def _save_pending_suggestions(suggestions: list) -> None:
    _ensure_memory_dir()
    PENDING_SUGGESTIONS_FILE.write_text(
        json.dumps(suggestions, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _clear_pending_suggestions() -> None:
    if PENDING_SUGGESTIONS_FILE.exists():
        PENDING_SUGGESTIONS_FILE.unlink()


def _read_recent_history(hours: int = 24) -> list:
    """Letzte N Stunden Konversation aus JSONL lesen."""
    if not HISTORY_FILE.exists():
        return []
    cutoff = time.time() - hours * 3600
    out = []
    try:
        with HISTORY_FILE.open("r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    if rec.get("ts", 0) > cutoff:
                        out.append(rec)
                except Exception:
                    continue
    except Exception:
        pass
    return out


def _read_recent_corrections(hours: int = 24) -> list:
    """Letzte N Stunden Korrekturen."""
    if not CORRECTIONS_FILE.exists():
        return []
    cutoff_dt = datetime.now(TIMEZONE) - timedelta(hours=hours)
    out = []
    try:
        with CORRECTIONS_FILE.open("r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    ts_str = rec.get("ts", "")
                    if ts_str:
                        rec_dt = datetime.fromisoformat(ts_str)
                        if rec_dt > cutoff_dt:
                            out.append(rec)
                except Exception:
                    continue
    except Exception:
        pass
    return out


async def compute_memory_suggestions() -> list:
    """LLM analysiert letzte 24h, schlägt Memory-Updates vor.

    Output: list of {type: 'preference'/'fact'/'project_context',
                     text: str, evidence: str, project_slug?: str}
    """
    history = _read_recent_history(24)
    corrections = _read_recent_corrections(24)

    if not history and not corrections:
        return []

    # History aufs Wesentliche kondensieren
    history_text_parts = []
    for rec in history[-100:]:  # cap
        msg = rec.get("msg", {})
        role = msg.get("role", "?")
        content = msg.get("content", "")
        if isinstance(content, list):
            content = " ".join(str(c.get("text", "")) for c in content if isinstance(c, dict))
        content_str = str(content)[:300]
        history_text_parts.append(f"[{role}] {content_str}")
    history_text = "\n".join(history_text_parts)

    corrections_text = "\n".join(
        f"- FALSCH: {c.get('was_falsch','')} | RICHTIG: {c.get('was_richtig','')} | KONTEXT: {c.get('kontext','')}"
        for c in corrections
    ) or "(keine Korrekturen)"

    extraction_prompt = f"""Analysiere die folgende 24h-Konversation + Korrekturen und schlage konservativ neue Memory-Einträge vor.

KORREKTUREN:
{corrections_text[:2000]}

KONVERSATION:
{history_text[:6000]}

Drei Memory-Typen die du vorschlagen kannst:
1. preference — Stil/Antwort-Regel die User mehrfach erwähnt hat (z.B. "antworte kürzer")
2. fact — Bio/Setup-Fakt über User der mehrfach auftaucht (z.B. "Schule: HTL Villach")
3. project_context — projektspezifische Regel (mit project_slug Feld)

REGELN:
- KONSERVATIV: nur vorschlagen wenn klar/mehrfach/explizit. KEINE Halluzinationen.
- Max 5 Vorschläge insgesamt.
- Wenn nichts klar ist: leeres Array zurückgeben.
- Nicht das vorschlagen was bereits in PRÄFERENZEN/FAKTEN-Sektion oben steht.

Antworte AUSSCHLIESSLICH mit JSON-Array (kein Markdown, kein Erklärungs-Text):
[
  {{"type": "preference", "text": "...", "evidence": "Grund (kurz, 1 Satz)"}},
  {{"type": "fact", "text": "...", "evidence": "..."}},
  {{"type": "project_context", "project_slug": "...", "text": "...", "evidence": "..."}}
]
"""

    try:
        resp = await _llm_call_with_retry(
            llm,
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": "Du bist Memory-Extractor. Antworte nur mit valid JSON."},
                {"role": "user", "content": extraction_prompt},
            ],
            max_tokens=1500,
        )
        try:
            usage = getattr(resp, "usage", None)
            if usage is not None:
                _track_usage(LLM_MODEL,
                             getattr(usage, "prompt_tokens", 0),
                             getattr(usage, "completion_tokens", 0),
                             kind="memory-extract")
        except Exception:
            pass
        content = resp.choices[0].message.content or "[]"
        # Strip markdown fences if any
        content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content.strip())
        suggestions = json.loads(content)
        if not isinstance(suggestions, list):
            return []
        # Validate
        clean = []
        for s in suggestions[:5]:
            if not isinstance(s, dict):
                continue
            stype = s.get("type")
            text = s.get("text", "").strip()
            if stype not in ("preference", "fact", "project_context") or not text:
                continue
            clean.append(s)
        return clean
    except Exception as e:
        log.exception(f"compute_memory_suggestions failed: {e}")
        return []


def _format_suggestion_briefing(suggestions: list) -> str:
    """Formatiert Suggestions als Telegram-HTML mit Antwort-Anleitung."""
    if not suggestions:
        return ""
    type_labels = {
        "preference": "✨ <b>PRÄFERENZEN</b> (Stil/Antwort-Regeln)",
        "fact": "📌 <b>FAKTEN</b> (Bio/Setup über dich)",
        "project_context": "🎯 <b>PROJEKT-KONTEXT</b> (Projekt-Regeln)",
    }
    by_type: dict = {}
    for i, s in enumerate(suggestions, 1):
        by_type.setdefault(s["type"], []).append((i, s))

    msg = "🌙 <b>Memory-Vorschläge</b> <i>(Analyse der letzten 24h)</i>\n"
    for ptype in ("preference", "fact", "project_context"):
        items = by_type.get(ptype, [])
        if not items:
            continue
        msg += f"\n{type_labels[ptype]}\n"
        for i, s in items:
            extra = ""
            if ptype == "project_context" and s.get("project_slug"):
                extra = f" <i>[{s['project_slug']}]</i>"
            msg += f"<b>{i}.</b> {_esc_html(s['text'])}{extra}\n"
            if s.get("evidence"):
                msg += f"   <i>Grund: {_esc_html(s['evidence'][:120])}</i>\n"

    msg += (
        "\n────────────\n"
        "<b>Wie du antwortest:</b>\n"
        "• <code>1 3</code> oder <code>1,3</code> → speichert nur 1 und 3\n"
        "• <code>alle</code> oder <code>ja</code> → alle übernehmen\n"
        "• <code>0</code> oder <code>nein</code> → alle verwerfen\n"
        "• <code>erkläre 2</code> → mehr Detail zu Vorschlag 2"
    )
    return msg


async def nightly_suggestion_job(ctx: ContextTypes.DEFAULT_TYPE):
    """JobQueue-Callback: nächtlich Memory-Vorschläge generieren + senden."""
    try:
        suggestions = await compute_memory_suggestions()
        if not suggestions:
            log.info("Nightly suggestions: nothing to suggest.")
            return
        _save_pending_suggestions(suggestions)
        msg = _format_suggestion_briefing(suggestions)
        await ctx.bot.send_message(
            chat_id=ALLOWED_USER_ID,
            text=msg,
            parse_mode=constants.ParseMode.HTML,
        )
        log.info(f"Nightly briefing sent: {len(suggestions)} suggestions pending")
    except Exception as e:
        log.exception(f"nightly_suggestion_job failed: {e}")


def apply_memory_suggestion(action: str) -> str:
    """Verarbeitet User-Antwort auf Memory-Briefing.

    Action-Formate: '1 3', '1,3', 'alle', 'ja', '0', 'nein', 'erkläre 2'
    """
    pending = _load_pending_suggestions()
    if not pending:
        return "Keine pending Memory-Vorschläge — vielleicht schon übernommen oder noch keine generiert."

    action = (action or "").strip().lower()

    # 'erkläre N'
    explain_match = re.search(r"(?:erklär|erklar)\w*\s+(\d+)", action)
    if explain_match:
        n = int(explain_match.group(1))
        if 1 <= n <= len(pending):
            s = pending[n-1]
            details = (f"Vorschlag {n}:\n"
                       f"Typ: {s['type']}\n"
                       f"Text: {s['text']}\n"
                       f"Grund: {s.get('evidence', '-')}")
            if s.get("project_slug"):
                details += f"\nProjekt-Slug: {s['project_slug']}"
            return details
        return f"Vorschlag {n} existiert nicht (1-{len(pending)})."

    # Decide which to apply
    if action in ("alle", "all", "ja", "yes", "y"):
        nums = list(range(1, len(pending) + 1))
    elif action in ("0", "nein", "no", "n", "skip", "verwerfen"):
        _clear_pending_suggestions()
        return f"Alle {len(pending)} Vorschläge verworfen."
    else:
        nums = [int(x) for x in re.findall(r"\d+", action)]
        nums = [n for n in nums if 1 <= n <= len(pending)]
        if not nums:
            return (f"'{action}' nicht verstanden. Erlaubt: '1 3' / 'alle' / 'nein' / 'erkläre N'")

    applied = []
    skipped = []
    for n in nums:
        s = pending[n-1]
        try:
            if s["type"] == "preference":
                set_preference(s["text"])
                applied.append(f"✓ Präf {n}: {s['text'][:60]}")
            elif s["type"] == "fact":
                remember(s["text"])
                applied.append(f"✓ Fakt {n}: {s['text'][:60]}")
            elif s["type"] == "project_context":
                slug = s.get("project_slug", "")
                if slug:
                    update_project_context(slug, s["text"])
                    applied.append(f"✓ Proj-Ctx {n} ({slug}): {s['text'][:50]}")
                else:
                    skipped.append(f"✗ {n}: project_slug fehlt")
            else:
                skipped.append(f"✗ {n}: unbekannter Typ {s['type']}")
        except Exception as e:
            skipped.append(f"✗ {n}: {e}")

    _clear_pending_suggestions()

    out = "Übernommen:\n" + "\n".join(applied)
    if skipped:
        out += "\n\nÜbersprungen:\n" + "\n".join(skipped)
    return out


def log_correction(was_falsch: str, was_richtig: str, kontext: str = "") -> str:
    """Speichert eine Korrektur als Lern-Datenpunkt.

    Wird ausgelöst wenn Bot etwas getan hat und User korrigiert
    ('nein anders', 'ich meinte X', 'verschieb das doch nach Y').
    Diese Records sind später Fine-Tuning-Gold (instruction-pair-Format).
    """
    if not was_falsch or not was_richtig:
        return "log_correction: was_falsch + was_richtig sind beide Pflicht."
    _ensure_memory_dir()
    record = {
        "ts": datetime.now(TIMEZONE).isoformat(timespec="seconds"),
        "was_falsch": was_falsch.strip(),
        "was_richtig": was_richtig.strip(),
        "kontext": (kontext or "").strip(),
    }
    try:
        with CORRECTIONS_FILE.open("a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
        log.info(f"Correction logged: {was_falsch[:60]}")
        return "✓ Korrektur gespeichert (für späteres Lernen)."
    except Exception as e:
        log.exception("log_correction failed")
        return f"Log-Fehler: {e}"


def _save_history_line(user_id: int, message: dict) -> None:
    """Append einzelne Message ans persistente JSONL."""
    _ensure_memory_dir()
    record = {"user_id": user_id, "ts": time.time(), "msg": message}
    try:
        with HISTORY_FILE.open("a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    except Exception as e:
        log.warning(f"history persist failed: {e}")


HISTORY_TAIL_READ_BYTES = 2 * 1024 * 1024   # 2MB Tail-Read-Cap (verhindert OOM bei riesigem JSONL)


def _read_tail_lines(path: Path, max_bytes: int) -> list[str]:
    """Liest die letzten max_bytes der Datei + returnt Lines (ohne erste evtl. partial).

    Verhindert dass eine 500MB-history.jsonl die Memory-grenze sprengt.
    Bei normaler Größe (<2MB) liest das ganze File. Bei großen Files: nur Tail.
    """
    file_size = path.stat().st_size
    if file_size <= max_bytes:
        return path.read_text(encoding="utf-8", errors="replace").splitlines()
    with path.open("rb") as f:
        f.seek(file_size - max_bytes)
        chunk = f.read()
    text = chunk.decode("utf-8", errors="replace")
    lines = text.splitlines()
    if lines:
        lines = lines[1:]  # Erste Zeile ist evtl angeschnitten
    return lines


def _sanitize_loaded_history(msgs: list) -> list:
    """Repariert Messages für strenge Provider (Gemini).

    1. tool-Messages ohne 'name' bekommen den Namen aus der vorherigen
       assistant-Message mit passender tool_call_id.
    2. Wenn keine Zuordnung möglich → tool-msg DROP (sonst Gemini 400).
    3. assistant-Messages mit tool_calls aber ohne content bekommen
       content="" (Gemini lehnt fehlendes content ab wenn tool_calls da sind).
    4. tool_calls werden auf {id, type, function:{name, arguments}} whitelisted
       (entfernt provider-spezifische Extra-Felder die anderswo abgelehnt werden).
    5. tool_call.function.arguments wird zu String normalisiert (manche Modelle
       liefern Dict statt JSON-String → Provider-Validation crasht).
    """
    # Map tool_call_id → tool_name aus assistant-Messages
    id_to_name: dict[str, str] = {}
    for m in msgs:
        if m.get("role") == "assistant":
            for tc in m.get("tool_calls") or []:
                tc_id = tc.get("id") if isinstance(tc, dict) else getattr(tc, "id", None)
                fn = (tc.get("function") if isinstance(tc, dict) else getattr(tc, "function", None)) or {}
                fn_name = fn.get("name") if isinstance(fn, dict) else getattr(fn, "name", None)
                if tc_id and fn_name:
                    id_to_name[tc_id] = fn_name

    out: list = []
    for m in msgs:
        role = m.get("role")
        if role == "tool":
            # Fix #1: name-Feld nachrüsten oder droppen
            if not m.get("name"):
                tc_id = m.get("tool_call_id")
                fn_name = id_to_name.get(tc_id) if tc_id else None
                if fn_name:
                    m = dict(m)
                    m["name"] = fn_name
                else:
                    continue  # Nicht reparierbar
            # Content muss String sein
            if not isinstance(m.get("content"), str):
                m = dict(m)
                m["content"] = str(m.get("content", ""))
        elif role == "assistant":
            tcs = m.get("tool_calls")
            if tcs:
                m = dict(m)
                # Fix #4: content="" wenn fehlt aber tool_calls da
                if m.get("content") is None or "content" not in m:
                    m["content"] = ""
                # Fix #5: tool_calls whitelisten
                clean_tcs = []
                for tc in tcs:
                    if not isinstance(tc, dict):
                        continue
                    tc_id = tc.get("id")
                    fn = tc.get("function") or {}
                    fn_name = fn.get("name") if isinstance(fn, dict) else None
                    fn_args = fn.get("arguments") if isinstance(fn, dict) else None
                    # Defensive: id + name müssen vorhanden sein
                    if not tc_id or not fn_name:
                        continue
                    # arguments zu String (kann dict sein bei manchen Modellen)
                    if isinstance(fn_args, (dict, list)):
                        try:
                            fn_args = json.dumps(fn_args, ensure_ascii=False)
                        except Exception:
                            fn_args = "{}"
                    elif fn_args is None:
                        fn_args = "{}"
                    elif not isinstance(fn_args, str):
                        fn_args = str(fn_args)
                    clean_tcs.append({
                        "id": tc_id,
                        "type": "function",
                        "function": {"name": fn_name, "arguments": fn_args},
                    })
                if clean_tcs:
                    m["tool_calls"] = clean_tcs
                else:
                    # Alle tool_calls invalid → drop tool_calls, behalt content
                    m.pop("tool_calls", None)
                    if not m.get("content"):
                        continue  # Komplett leer → skippen
        out.append(m)
    return out


def _load_persistent_history(user_id: int) -> list:
    """Letzte HISTORY_MAX_MESSAGES Messages für User aus JSONL laden.

    Memory-safe via _read_tail_lines: bei riesigen Files nur die letzten
    2MB lesen (immer noch 1000+ Records typisch).
    """
    if not HISTORY_FILE.exists():
        return []
    try:
        lines = _read_tail_lines(HISTORY_FILE, HISTORY_TAIL_READ_BYTES)
        records = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            try:
                rec = json.loads(line)
                if rec.get("user_id") == user_id:
                    records.append(rec)
            except json.JSONDecodeError:
                continue
        msgs = [r["msg"] for r in records[-HISTORY_MAX_MESSAGES:]]
        return _sanitize_loaded_history(msgs)
    except Exception as e:
        log.warning(f"history load failed: {e}")
        return []


def _maybe_compact_history() -> None:
    """JSONL trimmen wenn zu groß: nur letzte 200 Lines behalten.

    Memory-safe: nutzt _read_tail_lines statt full read_text wenn File riesig.
    """
    if not HISTORY_FILE.exists():
        return
    try:
        # Schnell-Check: File-Größe → wenn klein, direkt lesen, sonst tail-only
        file_size = HISTORY_FILE.stat().st_size
        if file_size <= HISTORY_TAIL_READ_BYTES:
            lines = HISTORY_FILE.read_text(encoding="utf-8").splitlines()
        else:
            lines = _read_tail_lines(HISTORY_FILE, HISTORY_TAIL_READ_BYTES)
        if len(lines) <= HISTORY_PERSIST_LIMIT:
            return
        keep = lines[-HISTORY_COMPACT_KEEP:]
        HISTORY_FILE.write_text("\n".join(keep) + "\n", encoding="utf-8")
        log.info(f"History compacted: {len(lines)} → {len(keep)}")
    except Exception as e:
        log.warning(f"history compact failed: {e}")


def _check_cache_fresh(user_id: int) -> tuple[Optional[list], bool]:
    """Pure Read, no IO. Returns (cached_list_or_None, is_fresh).

    Wird ohne Lock aus async-Context gerufen (RAM-Read ist atomic in CPython).
    Bei is_fresh=True kann der Caller direkt den Cache verwenden.
    """
    last = CONVERSATION_TIMESTAMPS.get(user_id, 0.0)
    cached = CONVERSATION_HISTORY.get(user_id)
    is_fresh = cached is not None and (time.time() - last) <= HISTORY_TIMEOUT
    return cached, is_fresh


async def _ensure_history_loaded(user_id: int) -> list:
    """Stellt sicher dass die History für user_id im Cache ist.

    Disk-IO läuft via to_thread → blockiert NICHT den event-loop. Cache-
    Update unter Lock damit konkurrierende Calls keine inkonsistenten
    Reads sehen. Returns die geladene Liste (nicht aus Cache, da andere
    Tasks zwischenzeitlich verändert haben könnten).
    """
    cached, is_fresh = _check_cache_fresh(user_id)
    if is_fresh:
        return list(cached)
    # Disk-IO AUSSERHALB Lock — blockiert event loop nicht
    loaded = await asyncio.to_thread(_load_persistent_history, user_id)
    async with _HISTORY_LOCK:
        # Recheck — ein anderer Task könnte schon geladen+mutiert haben
        cached, is_fresh = _check_cache_fresh(user_id)
        if is_fresh:
            return list(cached)
        CONVERSATION_HISTORY[user_id] = loaded
        CONVERSATION_TIMESTAMPS[user_id] = time.time()
    return list(loaded)


async def get_history(user_id: int) -> list:
    """History für User. Bei Cache-Miss/Timeout: async lazy-load aus JSONL.

    Disk-IO läuft via to_thread → blockiert event-loop nicht mehr.
    """
    return await _ensure_history_loaded(user_id)


async def update_history(user_id: int, new_messages: list) -> None:
    """History anhängen + persistieren + auf Max-Länge trimmen.

    Hot-Path: RAM-Mutation unter Lock (schnell). Disk-Append + Compaction
    via to_thread AUSSERHALB Lock — JSONL-Append ist append-only, Race
    zwischen mehreren Schreibern wäre nur Reihenfolge-Issue (für Single-
    User-Bot irrelevant).
    """
    # Sicherstellen dass Cache da ist (lädt wenn nötig, async)
    await _ensure_history_loaded(user_id)
    # RAM-Update unter Lock (sehr schnell — keine Disk-IO)
    async with _HISTORY_LOCK:
        history = list(CONVERSATION_HISTORY.get(user_id, []))
        history.extend(new_messages)
        if len(history) > HISTORY_MAX_MESSAGES:
            history = history[-HISTORY_MAX_MESSAGES:]
        CONVERSATION_HISTORY[user_id] = history
        CONVERSATION_TIMESTAMPS[user_id] = time.time()
    # Disk-Append + ggf Compaction via to_thread, kein Lock
    await asyncio.to_thread(_persist_history_changes, user_id, new_messages)


def _persist_history_changes(user_id: int, new_messages: list) -> None:
    """Sync helper für update_history — schreibt JSONL-Append + ggf compact."""
    for msg in new_messages:
        _save_history_line(user_id, msg)
    _maybe_compact_history()


async def reset_history(user_id: int) -> None:
    """RAM-Cache leeren. JSONL bleibt erhalten — Memory wird beim nächsten Mal neu geladen.

    Wenn du wirklich permanent löschen willst: HISTORY_FILE manuell entfernen.
    """
    async with _HISTORY_LOCK:
        CONVERSATION_HISTORY.pop(user_id, None)
        CONVERSATION_TIMESTAMPS.pop(user_id, None)


# ============================================================================
# Token-Usage-Tracking (für /usage Reports + Cost-Awareness)
# ============================================================================
# Pro Tag eine Zeile in 06_Meta/usage/YYYY-MM-DD.json — append-only.
# Records: {"ts": iso, "model": str, "in": int, "out": int, "kind": str}
# kind: "chat" (LLM-Call), "vision" (Photo-Caption), "whisper" (lokal, $0)

USAGE_DIR = VAULT / "06_Meta" / "usage"

# Approximate Pricing (USD per 1M tokens) — nur für Schätzung, nicht exakt.
# Für Ollama-Cloud / lokale Modelle: 0/0. Bot zeigt im /usage diese Caveat.
# Werte sind Stand 2025/2026 grob — User updated bei Bedarf in der Datei.
PRICING_USD_PER_M = {
    # OpenRouter / Anthropic
    "anthropic/claude-sonnet-4-5": (3.0, 15.0),
    "anthropic/claude-opus-4": (15.0, 75.0),
    "anthropic/claude-haiku-4": (0.8, 4.0),
    # OpenAI
    "openai/gpt-4o": (2.5, 10.0),
    "openai/gpt-4o-mini": (0.15, 0.60),
    # Google
    "google/gemini-2.5-flash": (0.075, 0.30),
    "google/gemini-2.5-pro": (1.25, 10.0),
    # Ollama Cloud / lokale → kostenlos
    "gpt-oss:120b-cloud": (0.0, 0.0),
    "qwen3:235b-cloud": (0.0, 0.0),
}


def _track_usage(model: str, prompt_tokens: int, completion_tokens: int, kind: str = "chat") -> None:
    """Speichert ein Token-Usage-Record in 06_Meta/usage/YYYY-MM-DD.jsonl.

    Wird sync aufgerufen aus dem llm_loop nach jedem LLM-Call. Failures
    werden geloggt aber nicht propagiert — Tracking darf den Bot nie crashen.
    """
    try:
        USAGE_DIR.mkdir(parents=True, exist_ok=True)
        today = today_iso()
        path = USAGE_DIR / f"{today}.jsonl"
        rec = {
            "ts": datetime.now(TIMEZONE).isoformat(timespec="seconds"),
            "model": model,
            "in": int(prompt_tokens or 0),
            "out": int(completion_tokens or 0),
            "kind": kind,
        }
        with path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception as e:
        log.debug(f"usage-track failed (silently): {e}")


def _estimate_cost_usd(model: str, in_tokens: int, out_tokens: int) -> float:
    """Schätzt Kosten in USD für Model+Tokens. 0 wenn Model unbekannt."""
    pricing = PRICING_USD_PER_M.get(model)
    if not pricing:
        return 0.0
    in_price, out_price = pricing
    return (in_tokens * in_price + out_tokens * out_price) / 1_000_000


def get_usage_summary(days: int = 7) -> str:
    """Liest letzte N Tage Usage und produziert Report.

    Tool-callable von User via /usage oder LLM-Tool.
    """
    if not USAGE_DIR.exists():
        return "Noch keine Usage-Daten."
    today = datetime.now(TIMEZONE).date()
    by_day: dict = {}  # day_iso → {"in": x, "out": y, "calls": z, "cost": $, "by_model": {model: ...}}
    for delta in range(days):
        d = today - timedelta(days=delta)
        d_iso = d.isoformat()
        path = USAGE_DIR / f"{d_iso}.jsonl"
        if not path.exists():
            continue
        day_data = {"in": 0, "out": 0, "calls": 0, "cost_usd": 0.0, "by_model": {}}
        try:
            for line in path.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                except json.JSONDecodeError:
                    continue
                model = rec.get("model", "unknown")
                ti = int(rec.get("in", 0) or 0)
                to = int(rec.get("out", 0) or 0)
                cost = _estimate_cost_usd(model, ti, to)
                day_data["in"] += ti
                day_data["out"] += to
                day_data["calls"] += 1
                day_data["cost_usd"] += cost
                m = day_data["by_model"].setdefault(model, {"in": 0, "out": 0, "calls": 0, "cost": 0.0})
                m["in"] += ti
                m["out"] += to
                m["calls"] += 1
                m["cost"] += cost
            by_day[d_iso] = day_data
        except Exception as e:
            log.warning(f"usage-summary day {d_iso} failed: {e}")

    if not by_day:
        return f"Keine Usage-Daten in den letzten {days} Tagen."

    # Aggregiert
    total_in = sum(d["in"] for d in by_day.values())
    total_out = sum(d["out"] for d in by_day.values())
    total_calls = sum(d["calls"] for d in by_day.values())
    total_cost = sum(d["cost_usd"] for d in by_day.values())

    parts = [f"📊 **Token-Usage** (letzte {days} Tage)"]
    parts.append(f"")
    parts.append(f"**Total:** {total_calls} Calls · {total_in:,} in · {total_out:,} out")
    if total_cost > 0:
        parts.append(f"**Geschätzte Kosten:** ${total_cost:.3f} USD")
    else:
        parts.append("_Aktuelle Modelle haben keinen Preis hinterlegt (z.B. Ollama-Cloud = gratis)._")

    parts.append("")
    parts.append("**Pro Tag:**")
    for d_iso in sorted(by_day.keys(), reverse=True):
        d = by_day[d_iso]
        cost_str = f" · ${d['cost_usd']:.3f}" if d["cost_usd"] > 0 else ""
        parts.append(f"  • {d_iso}: {d['calls']} calls · {d['in']:,} in · {d['out']:,} out{cost_str}")

    # Top-Modelle
    model_totals: dict = {}
    for d in by_day.values():
        for m, mdata in d["by_model"].items():
            mt = model_totals.setdefault(m, {"in": 0, "out": 0, "calls": 0, "cost": 0.0})
            mt["in"] += mdata["in"]
            mt["out"] += mdata["out"]
            mt["calls"] += mdata["calls"]
            mt["cost"] += mdata["cost"]
    if len(model_totals) > 1:
        parts.append("")
        parts.append("**Pro Modell:**")
        for m, mt in sorted(model_totals.items(), key=lambda x: -x[1]["calls"]):
            cost_str = f" · ${mt['cost']:.3f}" if mt["cost"] > 0 else ""
            parts.append(f"  • `{m}`: {mt['calls']} calls · {mt['in']:,}/{mt['out']:,} in/out{cost_str}")

    return "\n".join(parts)


# ============================================================================
# Tool definitions (OpenAI function-calling format)
# ============================================================================

TOOLS = [
    {"type": "function", "function": {
        "name": "append_to_daily",
        "description": "Hängt Text an die heutige Daily-Note unter passender Sektion an. Wähle Sektion: 'Heute' für Tasks/Termine, 'Notizen & Gedanken' für Gedanken/Notizen (default), 'Offen / Einsortieren' für offene Sachen/Links, 'Abends' für Reflexion/Bewertung des Tages.",
        "parameters": {
            "type": "object",
            "properties": {
                "section": {"type": "string", "enum": list(VALID_SECTIONS)},
                "text": {"type": "string", "description": "Markdown-Text zum Anhängen."},
            },
            "required": ["section", "text"],
        },
    }},
    {"type": "function", "function": {
        "name": "create_task",
        "description": (
            "Legt neuen Task in 10_Life/tasks/ an, verlinkt unter 'Heute' in heutiger Daily. "
            "WICHTIG: Optionale Parameter NUR setzen wenn explizit aus User-Input ableitbar — "
            "NIE raten/defaulten. Z.B. due nur wenn User wirklich ein Datum nennt; "
            "project nur wenn User Projekt nennt oder es klar aus Conversation-Memory hervorgeht. "
            "tags: 2-5 thematische Schlagworte aus dem Inhalt extrahieren (kebab-case, Deutsch). "
            "recurrence NUR setzen bei klaren Wiederholungs-Triggern: 'jeden Tag', 'täglich', "
            "'jede Woche', 'jeden Montag', 'jeden Monat' o.ä. — NICHT bei einmaligen Tasks!"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Vollständiger Task-Titel (deutsch ok)"},
                "priority": {"type": "string", "enum": ["low", "medium", "high", "urgent"], "description": "Default medium falls User keine Prio nennt"},
                "due": {"type": "string", "description": "ISO YYYY-MM-DD. NUR setzen wenn User explizit ein Datum nennt!"},
                "area": {"type": "string", "description": "Area-ID, NUR wenn User Area nennt"},
                "project": {"type": "string", "description": "Projekt-Slug (z.B. 'sanierung-kiosk'), NUR wenn aus User-Input oder Memory klar"},
                "context": {"type": "string", "enum": ["home", "work", "errand", "phone", "computer"]},
                "tags": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "2-5 thematische Tags (kebab-case, Deutsch). Topical, nicht 'task'/'todo'. Bei Bedarf vorher list_existing_tags für konsistente Vokabular.",
                },
                "recurrence": {
                    "type": "string",
                    "enum": ["daily", "weekdays", "weekly", "monthly"],
                    "description": (
                        "Wiederholungs-Pattern. NUR setzen bei klaren Wiederholungs-Wörtern: "
                        "'jeden Tag/täglich' → daily, 'werktags/Mo-Fr' → weekdays, "
                        "'jede Woche/jeden <Wochentag>' → weekly, 'jeden Monat/monatlich' → monthly. "
                        "Wenn User nur einmal etwas erwähnt: NICHT setzen."
                    ),
                },
            },
            "required": ["title"],
        },
    }},
    {"type": "function", "function": {
        "name": "mark_task_done",
        "description": "Setzt einen Task auf done. Slug ohne 't-'-Präfix möglich.",
        "parameters": {
            "type": "object",
            "properties": {"slug": {"type": "string"}},
            "required": ["slug"],
        },
    }},
    {"type": "function", "function": {
        "name": "get_today_agenda",
        "description": (
            "Aggregiert die heutige Lage: heute fällige Tasks, überfällige Tasks, "
            "heute feuernde Reminders, heute geplante Meetings, plus High-Prio-Tasks ohne Datum. "
            "PRIMÄRER Tool-Call wenn User fragt 'was steht heute an', 'was muss ich noch erledigen', "
            "'was ist heute zu tun', 'mein Plan für heute', 'agenda', 'tagesplan'. "
            "Liefert fertig formatiertes Telegram-HTML — direkt zurück an User."
        ),
        "parameters": {"type": "object", "properties": {}, "required": []},
    }},
    {"type": "function", "function": {
        "name": "list_open_tasks",
        "description": (
            "Listet offene Tasks gefiltert. Ohne Filter: gruppiert in Buckets "
            "Überfällig/Heute/Diese Woche/Später/Ohne Datum. "
            "Mit when-Filter: flache Liste. "
            "Beispiele: list_open_tasks() — alle, gruppiert. "
            "list_open_tasks(when='overdue') — nur überfällige. "
            "list_open_tasks(project='matura', priority='urgent') — nur urgent in Matura."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "when": {
                    "type": "string",
                    "enum": ["today", "tomorrow", "week", "overdue", "nodate"],
                    "description": "Zeit-Filter. Weglassen → alle gruppiert.",
                },
                "priority": {"type": "string", "enum": ["low", "medium", "high", "urgent"]},
                "project": {"type": "string", "description": "Projekt-Slug"},
                "area": {"type": "string", "description": "Area-Slug"},
                "context": {"type": "string", "enum": ["home", "work", "errand", "phone", "computer"]},
            },
            "required": [],
        },
    }},
    {"type": "function", "function": {
        "name": "create_meeting",
        "description": "Legt ein Meeting-Protokoll in 10_Life/meetings/ an. tags: thematische Schlagworte extrahieren (z.B. ['blitztext', 'kunde', 'kickoff']).",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string"},
                "attendees": {"type": "array", "items": {"type": "string"}},
                "meeting_date": {"type": "string", "description": "ISO-Datum YYYY-MM-DD, default heute"},
                "tags": {"type": "array", "items": {"type": "string"}, "description": "2-5 thematische Tags, kebab-case Deutsch"},
            },
            "required": ["title"],
        },
    }},
    {"type": "function", "function": {
        "name": "create_note",
        "description": "Legt eine freie Notiz in 10_Life/notes/ an. Für längere strukturierte Inhalte (>3 Sätze, eigenes Thema), die weder Task noch Tagesreflexion sind. tags: 2-5 thematische Schlagworte aus dem Inhalt.",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string"},
                "body": {"type": "string"},
                "tags": {"type": "array", "items": {"type": "string"}, "description": "2-5 thematische Tags, kebab-case Deutsch (z.B. ['gesundheit', 'ernaehrung'])"},
            },
            "required": ["title", "body"],
        },
    }},
    {"type": "function", "function": {
        "name": "search_vault",
        "description": "Volltextsuche im Vault. Nutze für 'Was weiß ich über X', Wiki-Lookups, Task-Suche.",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string"},
                "limit": {"type": "integer", "default": 5},
            },
            "required": ["query"],
        },
    }},
    {"type": "function", "function": {
        "name": "read_file",
        "description": (
            "Liest eine Vault-Datei (Pfad relativ zu Vault-Root). "
            "strip_frontmatter=true (default): YAML-Header wird entfernt — nimm das wenn du "
            "dem User den Inhalt zeigst. strip_frontmatter=false: kompletter Inhalt inkl. "
            "Metadaten — nimm das wenn du Frontmatter-Felder brauchst (status, due, etc.)."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "rel_path": {"type": "string"},
                "strip_frontmatter": {"type": "boolean", "default": True},
            },
            "required": ["rel_path"],
        },
    }},
    {"type": "function", "function": {
        "name": "edit_file",
        "description": "Find/Replace in einer Vault-Datei. Bei regex=true wird find als Regex interpretiert.",
        "parameters": {
            "type": "object",
            "properties": {
                "rel_path": {"type": "string"},
                "find": {"type": "string"},
                "replace": {"type": "string"},
                "regex": {"type": "boolean", "default": False},
            },
            "required": ["rel_path", "find", "replace"],
        },
    }},
    {"type": "function", "function": {
        "name": "move",
        "description": (
            "Vereinheitlichtes Move-Tool für 3 Use-Cases — nutze GENAU EINEN Modus:\n"
            "(1) EINE Datei/Ordner: src='pfad', dst='ziel'.\n"
            "(2) MEHRERE Files in Zielordner (bulk, spart Tool-Calls): srcs=['a','b'], dst='ordner/'.\n"
            "(3) Projekt verschieben: project_slug='X', parent='Y' (oder leer für Top-Level).\n"
            "Für Multi-File-Upload IMMER (2) statt N-mal (1)."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "src": {"type": "string", "description": "Einzel-Move: Quell-Pfad (vault-relativ)."},
                "srcs": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Bulk-Move: Liste von Quell-Pfaden.",
                },
                "dst": {"type": "string", "description": "Ziel-Pfad/Ordner. Bei project_slug irrelevant."},
                "project_slug": {"type": "string", "description": "Bei Projekt-Move: Slug."},
                "parent": {"type": "string", "description": "Bei project_slug: neuer Parent oder leer für Top-Level."},
                "overwrite": {"type": "boolean", "default": False},
            },
            "required": [],
        },
    }},
    {"type": "function", "function": {
        "name": "clip_url",
        "description": "Fetcht eine URL und speichert den Hauptinhalt als Markdown-Artikel in 01_Raw/articles/.",
        "parameters": {
            "type": "object",
            "properties": {"url": {"type": "string"}},
            "required": ["url"],
        },
    }},
    {"type": "function", "function": {
        "name": "request_delete",
        "description": (
            "Schritt 1 von 2 für Löschen: meldet Löschanfrage an User. "
            "Akzeptiert einen ODER mehrere Pfade. Bei 'lösche alle X' erst list_files. "
            "permanent=False (Default): nach 99_Archive/ verschieben (reversibel). "
            "permanent=True NUR wenn User explizite Worte nutzt wie "
            "'endgültig', 'wirklich weg', 'permanent', 'unwiderruflich', 'hart löschen', "
            "'für immer', 'aus dem archiv auch'. Default ist immer Archiv (sicher). "
            "NIEMALS direkt confirm_delete ohne vorherige User-Bestätigung."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "rel_path": {
                    "type": "string",
                    "description": "Einzelner Pfad relativ zu Vault-Root. Für mehrere → rel_paths nutzen.",
                },
                "rel_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Liste von Pfaden relativ zu Vault-Root (Bulk-Delete).",
                },
                "permanent": {
                    "type": "boolean",
                    "description": "True = unwiderruflich rm. Default false (Archiv).",
                    "default": False,
                },
            },
            # Kein required — Handler akzeptiert rel_path ODER rel_paths.
            # oneOf wäre semantisch sauberer aber Gemini/Ollama strikt lehnen
            # oneOf/anyOf in Tool-Schemas ab → 400.
        },
    }},
    {"type": "function", "function": {
        "name": "confirm_delete",
        "description": (
            "Schritt 2 von 2: bestätigt oder bricht pending Löschungen ab. "
            "action='confirm' (default) führt aus (→ ins 99_Archive verschoben oder "
            "permanent gelöscht je nach mode). action='cancel' verwirft pending. "
            "User-Trigger: 'ja/machs/bestätigt' → confirm. 'nein/abbrechen/doch nicht' → cancel."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["confirm", "cancel"],
                    "default": "confirm",
                },
            },
            "required": [],
        },
    }},
    {"type": "function", "function": {
        "name": "list_files",
        "description": (
            "Listet User-Content-Files (.md) im Vault. System-Krempel (Templates, Meta, "
            "Tools, Trash, Archive, CLAUDE.md etc.) wird automatisch rausgefiltert. "
            "Bei >12 Files wird nach Top-Level-Ordner gruppiert."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "rel_dir": {
                    "type": "string",
                    "description": "Verzeichnis relativ zu Vault-Root, z.B. '10_Life/tasks'. Leer = ganzer Vault.",
                },
                "include_system": {
                    "type": "boolean",
                    "description": "True = auch Templates/Meta/Tools/Archive zeigen (Debug). Default false.",
                    "default": False,
                },
            },
            "required": [],
        },
    }},
    # apply_memory_suggestion + apply_health_action sind KEINE LLM-Tools mehr —
    # sie werden direkt im handle_text via _detect_pending_reply_intent gerufen
    # (Reply-Parser für Memory/Health-Briefing-Antworten, kein LLM-Roundtrip nötig).
    {"type": "function", "function": {
        "name": "log_correction",
        "description": (
            "Speichert eine Korrektur in corrections.jsonl. RUFE DAS AUF wenn der "
            "User dich korrigiert ('nein, anders', 'das war falsch', 'ich meinte X', "
            "'verschieb das doch nach Y', 'mach das anders'). "
            "was_falsch = was du getan/gesagt hast. "
            "was_richtig = was der User wollte. "
            "kontext = kurze Beschreibung der Situation. "
            "Diese Daten werden für nightly Memory-Vorschläge + Fine-Tuning genutzt."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "was_falsch": {"type": "string"},
                "was_richtig": {"type": "string"},
                "kontext": {"type": "string"},
            },
            "required": ["was_falsch", "was_richtig"],
        },
    }},
    {"type": "function", "function": {
        "name": "set_preference",
        "description": (
            "Speichert eine Stil-/Antwort-Präferenz dauerhaft in preferences.md. "
            "Nutze wenn User sagt 'antworte immer kürzer', 'verwende kein gerne!', "
            "'gib Datum in DD.MM. statt ISO', etc. UNTERSCHIED zu remember(): "
            "Präferenzen = wie der Bot reagieren soll, Facts = was der Bot wissen muss."
        ),
        "parameters": {
            "type": "object",
            "properties": {"text": {"type": "string", "description": "Präferenz als knapper Satz"}},
            "required": ["text"],
        },
    }},
    {"type": "function", "function": {
        "name": "project_context",
        "description": (
            "Steuert Projekt-Kontext-Loading + CONTEXT.md-Editing in einem Tool. "
            "action='activate' (slug nötig) → setzt Projekt aktiv, CONTEXT.md wird ab jetzt "
            "in jeden System-Prompt eingespeist. Trigger: 'lass uns über X reden'. "
            "action='deactivate' → bricht aktives Projekt ab. Trigger: 'fertig mit X'. "
            "action='update' (slug+text nötig) → schreibt nach 05_Projects/<slug>/CONTEXT.md, "
            "mode=append (default) oder replace. Trigger: 'merk für Projekt X: Y'."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "action": {"type": "string", "enum": ["activate", "deactivate", "update"]},
                "slug": {"type": "string", "description": "Projekt-Slug ohne 'project-' Präfix. Bei deactivate optional."},
                "text": {"type": "string", "description": "Nur bei action='update': Kontext-Text."},
                "mode": {"type": "string", "enum": ["append", "replace"], "default": "append"},
            },
            "required": ["action"],
        },
    }},
    {"type": "function", "function": {
        "name": "remember",
        "description": (
            "Fügt einen persistenten Fakt zur Bot-Memory hinzu (06_Meta/bot-memory/facts.md). "
            "Nutze wenn User sagt 'merk dir dass...', 'speicher als Fakt', "
            "'das musst du wissen', oder wenn ein offensichtlich dauerhafter Fakt "
            "über User/Setup auftaucht (z.B. Schule, KV-Lohn, Lieblings-Kunde). "
            "Werden bei jedem LLM-Call automatisch in den System-Prompt eingespeist."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "fact": {"type": "string", "description": "Der Fakt als kurzer Satz (z.B. 'KV-Lohn 2026 ist 32,80 €/h')"},
            },
            "required": ["fact"],
        },
    }},
    {"type": "function", "function": {
        "name": "forget",
        "description": (
            "Entfernt Memory-Einträge die ein Pattern enthalten (case-insensitive). "
            "kind='fact' für Fakten in facts.md, kind='preference' für Stil-Regeln in preferences.md. "
            "Trigger: 'vergiss X' / 'lösch die Regel Y' / 'das stimmt nicht mehr'."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "kind": {"type": "string", "enum": ["fact", "preference"]},
                "pattern": {"type": "string", "description": "Text-Snippet zum matchen"},
            },
            "required": ["kind", "pattern"],
        },
    }},
    {"type": "function", "function": {
        "name": "list_existing_tags",
        "description": (
            "Liste der bereits im Vault verwendeten Tags, sortiert nach Häufigkeit. "
            "Optional vor create_task/note/meeting aufrufen wenn unklar welche Tags konsistent sind. "
            "Hilft Drift zu vermeiden ('arbeit' vs 'work' vs 'job')."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "top_n": {"type": "integer", "default": 30, "description": "Anzahl Top-Tags zurückgeben"},
            },
            "required": [],
        },
    }},
    {"type": "function", "function": {
        "name": "create_project",
        "description": (
            "Legt einen neuen Projekt-Ordner unter 05_Projects/<slug>/ an. "
            "Erzeugt Folder + README.md mit Dataview-Queries für zugehörige "
            "Tasks/Notes/Meetings (gefiltert via project=<slug> Frontmatter). "
            "Nutze wenn User 'lege ein Projekt an', 'neues Projekt: ...' o.ä. sagt. "
            "Bei Subprojekt: parent=<slug-des-eltern-projekts> setzen — wird unter "
            "05_Projects/<parent>/<slug>/ angelegt. Bestehende Projekte verschiebt "
            "man stattdessen mit move_project. "
            "WICHTIG: dieses Tool legt den Ordner UND die README-Datei direkt an — "
            "nicht erst fragen, dann ankündigen, dann nochmal fragen. Direkt machen."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Projekt-Name, wird zu Slug umgewandelt"},
                "description": {"type": "string", "description": "Optional: kurze Beschreibung fürs README"},
                "area": {"type": "string", "description": "Optional: zugehörige Area-ID"},
                "parent": {"type": "string", "description": "Optional: Parent-Projekt-Slug für Subprojekt-Anlage"},
            },
            "required": ["name"],
        },
    }},
    {"type": "function", "function": {
        "name": "create_reminder",
        "description": (
            "Setzt eine Erinnerung — Bot schickt zur angegebenen Zeit eine Telegram-Nachricht "
            "mit dem message-Text. Berechne when_iso ALS ABSOLUTES DATUM + ZEIT in Lokalzeit "
            "(Europe/Vienna). 'morgen 15:00' → 2026-04-27T15:00:00. 'in 30 Min' → "
            "aktuelle Zeit + 30 Min. Heutiges Datum kennst du aus dem System-Prompt. "
            "Wiederholungen via recurrence: daily, weekdays (Mo-Fr), weekly (jede Woche "
            "selber Wochentag). Leer/null = einmalig."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "when_iso": {"type": "string", "description": "ISO-Datetime YYYY-MM-DDTHH:MM:SS, lokal Europe/Vienna"},
                "message": {"type": "string", "description": "Text der gesendet wird"},
                "recurrence": {"type": "string", "enum": ["daily", "weekdays", "weekly"]},
            },
            "required": ["when_iso", "message"],
        },
    }},
    {"type": "function", "function": {
        "name": "list_reminders",
        "description": "Listet alle aktiven Erinnerungen (für 'was steht an?', 'welche Reminder hab ich').",
        "parameters": {"type": "object", "properties": {}, "required": []},
    }},
    {"type": "function", "function": {
        "name": "cancel_reminder",
        "description": "Bricht eine Erinnerung ab. Erst list_reminders aufrufen wenn der User keine ID nennt.",
        "parameters": {
            "type": "object",
            "properties": {"reminder_id": {"type": "string", "description": "Reminder-ID, z.B. 'rem-20260426-153000-123'"}},
            "required": ["reminder_id"],
        },
    }},
    {"type": "function", "function": {
        "name": "backup_vault",
        "description": (
            "Manuelles Backup: pusht das gesamte Vault in das konfigurierte private "
            "GitHub-Repo. Nutze wenn der User 'mach backup', 'sicher das vault', "
            "'git push' o.ä. sagt. Idempotent — wenn nichts geändert, kommt entsprechende "
            "Meldung. Hat keine Parameter."
        ),
        "parameters": {"type": "object", "properties": {}, "required": []},
    }},
]

TOOL_HANDLERS = {
    "append_to_daily": append_to_daily,
    "create_task": create_task,
    "mark_task_done": mark_task_done,
    "get_today_agenda": get_today_agenda,
    "list_open_tasks": list_open_tasks,
    "create_meeting": create_meeting,
    "create_note": create_note,
    "search_vault": search_vault,
    "read_file": read_file,
    "edit_file": edit_file,
    "move": move,
    "clip_url": clip_url,
    "request_delete": request_delete,
    "confirm_delete": confirm_delete,
    "list_files": list_files,
    "list_existing_tags": list_existing_tags,
    "remember": remember,
    "forget": forget,
    "set_preference": set_preference,
    "project_context": project_context,
    "log_correction": log_correction,
    # apply_memory_suggestion + apply_health_action sind KEINE LLM-Tools mehr —
    # werden via _detect_pending_reply_intent in handle_text direkt gerufen.
    "create_project": create_project,
    "create_reminder": create_reminder,
    "list_reminders": list_reminders,
    "cancel_reminder": cancel_reminder,
    "backup_vault": backup_vault,
}

# ============================================================================
# System prompt (cached)
# ============================================================================

SYSTEM_PROMPT = """Du bist Julius' Vault-Assistent über Telegram. Deutsch. Heute ist {today}, jetzt {now} ({tz}).

VAULT
- 10_Life/{daily,tasks,notes,meetings,areas}/   Persönliches
- 05_Projects/<slug>/                            Projekte (Subprojekte als Subordner)
- 02_Wiki/                                       Kompiliertes Wissen
- 01_Raw/                                        Externe Quellen (Articles, Uploads)

# VERHALTEN — Default: nur antworten, nichts speichern. Tool nur bei explizitem Speicher-Intent.

| Trigger | Tool |
|---|---|
| "speicher / merk dir / notiere / ins tagebuch" | `append_to_daily` oder `create_note` |
| "task: …" / "todo: …" / Imperativ+Frist | `create_task` |
| "jeden Tag / täglich / jede Woche / jeden <Wochentag> / monatlich" | `create_task` mit `recurrence` (daily/weekdays/weekly/monthly) |
| "was steht heute an / agenda / tagesplan / was muss ich noch tun" | `get_today_agenda` |
| "alle offenen tasks / todo-liste" | `list_open_tasks()` (gruppiert) |
| "was ist überfällig / morgen / diese woche / ohne datum" | `list_open_tasks(when=overdue/tomorrow/week/nodate)` |
| "tasks für <projekt>" | `list_open_tasks(project='<slug>')` |
| "meeting: … / war im Termin mit" | `create_meeting` |
| "X erledigt / fertig / done" | `mark_task_done` |
| URL allein | erst fragen, dann `clip_url` |
| "lösche X" | `request_delete` (Default: ins Archiv = reversibel) |
| "lösche endgültig / komplett / vollständig / ganz weg / wirklich weg / unwiderruflich / für immer / hart" | `request_delete(permanent=true)` |
| "lösche alle X / leere Y" | erst `list_files`, dann `request_delete` mit Liste |
| "ja/bestätigt" nach request_delete | `confirm_delete()`. "nein/abbrechen" → `confirm_delete(action='cancel')`. |
| "lege Projekt X an / neues Projekt" | `create_project` (Ordner+README direkt, nicht erst nachfragen) |
| "X als Subprojekt von Y / schiebe X unter Y" | `move(project_slug='X', parent='Y')` |
| "verschieb A nach B / ordne in Y ein" | `move(src='A', dst='B')` (Einzel) oder `move(src=['a','b','c'], dst='ordner/')` (Bulk bei ≥2 Files) |
| "erinner mich an X um Y / wecker für 14 / in 30 Min" | `create_reminder` (when_iso = absolute Lokalzeit!) |
| "jeden Montag 8 Uhr / täglich um 7" | `create_reminder` mit recurrence (daily/weekdays/weekly) |
| "welche Reminder / cancel reminder" | `list_reminders` / `cancel_reminder` |
| "mach backup / sicher das vault" | `backup_vault` |

**Nachfragen vs. Direkt-Machen**:
- Mini-Input ohne Kontext ("ja", "?") → nachfragen.
- Klarer Auftrag → direkt machen, kein "wo soll ich"-Loop.

**KORREKTUR-PATTERN** (kritisch — sonst legst du falsche Tasks an):
Wenn User mit "nein", "stopp", "falsch", "besser X statt Y", "tausch X zu Y aus",
"X durch Y ersetzen" reagiert UND vorher gerade etwas erstellt/geschrieben wurde:
→ Das ist eine KORREKTUR der LETZTEN Aktion, KEIN neuer Auftrag.
→ NIEMALS "X zu Y aus" als Task-Titel interpretieren — das ist Imperativ "[tausche] X zu Y aus".
→ Aktion: passende Datei mit `edit_file` korrigieren ODER `request_delete` der falschen Datei.
→ Plus: `log_correction(was_falsch, was_richtig, kontext)` aufrufen für Lern-Material.

Beispiele:
- Vorher Task `t-foo` angelegt → User: "nein, das sollte 'bar' heißen" → edit_file Title+id, NICHT neue Task.
- Vorher CONTEXT.md geschrieben "X" → User: "besser Y statt X" → edit_file CONTEXT.md, NICHT neuer Eintrag.
- Vorher Reminder gesetzt → User: "stopp, falsche Uhrzeit" → cancel_reminder + neu erstellen.

**FRAME-DISZIPLIN** (NIE verwechseln):
- Du bist der Bot, User ist Julius. Antworte NIE in Ich-Form mit "Ja, bitte bestätigen"
  oder "Ja, bitte löschen" — das wäre die User-Antwort, nicht deine.
- Wenn du eine Bestätigung brauchst, FRAGE in 2.-Person: "Soll ich endgültig löschen? (ja/nein)"
- Wenn du Tool-Output einbettest: präsentiere ihn, simuliere keine User-Antwort darauf.

**FRUST-ERKENNUNG**:
Bei klarer User-Frustration ("Bist du dumm?", "STOP", "nein!!!", "wieder falsch", mehrfache !!!,
sarkastische Klammer-Fragen) → KURZ entschuldigen + klärend nachfragen statt blind weiter zu agieren.
Beispiel: "Sorry, hab's verbockt. Was soll ich konkret tun?" — KEINE neue Aktion ohne Rückversicherung.
Plus: `log_correction` aufrufen mit dem was schief lief.

**MULTI-FILE-WORKFLOW** (kritisch — sonst läuft Tool-Loop voll):
Bei Upload + "lege als Projekt X an": (1) `create_project`, (2) EINEN `move(src=[alle Pfade], dst='05_Projects/<slug>/')`-Call (NIEMALS N×einzeln), (3) optional `project_context(action='activate', slug=...)`. Max 3-4 Tool-Calls für N Files.

**TAGESPLAN-EXTRAKTION** (User listet mehrere Items, oft als Antwort aufs Morgens-Briefing):
- Klare Uhrzeit + Aktivität → `create_reminder` (KEIN Reminder ohne Uhrzeit erfinden!)
- Verb ohne Uhrzeit ("X machen / Y anrufen") → `create_task` mit `due=heute`
- Vage Tageszeit ("morgens/mittags/abends X") → Reminder mit Default-Stunde (8/12/18)

Pro Item ein Tool-Call (parallel im selben Loop-Step OK). Eine Bestätigung am Ende: "Eingetragen: 2 Reminders, 2 Tasks."

# LESEN
- "Was weiß ich über X" → `search_vault`, antworten mit echten `[[id]]`
- File zeigen → `read_file` (strip_frontmatter=true). Leere Sektionen weglassen, kein Navigation-Footer.
- Lange Files (>2000 Zeichen) zusammenfassen, nicht raw dumpen.

# AUSGABE
- Deutsch, direkt, kein Geschwurbel. Sparsame Emojis (✓ ✗ ⚠️).
- Aktion-Bestätigung: 1 Satz. Wikilink `[[id]]` NUR bei NEU erstellten Items (create_task/note/meeting/project/area), damit Julius hinklicken kann. Bei `mark_task_done`, `request_delete`, `move`, `edit_file` & Status-Änderungen: NUR Klartext-Titel ohne Slug-Wikilink (User kennt den Task ja schon). Frage: so lang wie nötig.
- Format: Bullets/Code/`**bold**`/`*italic*`. Headings nur bei langen Antworten.
- **TELEGRAM-TABELLEN**: nur ≤2 Spalten + kurze Zellen. Sobald Pfade/lange Texte/≥3 Spalten → kein Tabellen-Format, stattdessen pro Item: `**Name**` + eingerückte `• Label: Wert`-Bullets.
- **Wikilinks**: nur echte IDs aus search_vault/read_file. Keine Filepaths `[[10_Life/…]]`, keine Platzhalter `[[t-example]]`. Wenn keine ID → Klartext, KEIN Link.
  - **Auto-Linking aktiv**: bei `append_to_daily`, `create_note`, `project_context(action='update', ...)` macht der Renderer Wikilinks aus Klartext (Matura → `[[project-matura|Matura]]`). Bei direkter Telegram-Antwort selbst setzen.
- NIE HTML-Tags. NIE Frontmatter ausgeben.

# DATEN
- Datum ISO `YYYY-MM-DD`. "morgen" = +1, "nächsten Montag" → berechnen.
- **Tasks: `due` NUR wenn explizit genannt!** "Schränke aufbauen" → kein due. "morgen X" → due=morgen.
- **Priorität aus Sprache** (statt default medium):
  - "dringend/urgent/asap/sofort" → urgent · "wichtig/high prio" → high · "irgendwann/low prio" → low
  - Wenn Prio-Wort am Titel-Anfang: aus Titel ENTFERNEN (z.B. "dringend Backup" → title="Backup", priority=urgent)

# MEMORY — Decision-Tree

Drei Memory-Typen (alle automatisch im System-Prompt eingespeist):

| User sagt … | Tool | Begründung |
|---|---|---|
| "antworte kürzer / kein 'gerne!' / DD.MM. statt ISO / weniger Tabellen" | `set_preference` | WIE der Bot reagiert (Stil/Format/Tonalität) |
| "merk dir / ich heiße Julius / KV-Lohn 32,80 / studiere HTL" | `remember` | WAS über den User wahr ist (Bio/Setup/Fakten) |
| "Sanierung-Kiosk: ÖNORM B 2061 / Auftraggeber Schiemer" (+ Projekt aktiv) | `project_context(action='update', slug=..., text=...)` | Projekt-spezifische Regel |
| "lass uns über X reden / arbeite jetzt an X" | `project_context(action='activate', slug='X')` | — |
| "fertig mit X / weg vom Projekt" | `project_context(action='deactivate')` | — |

**MULTI-FAKT-KOMPRESSION**: Mehrere Fakten in einem Satz → EIN `remember`-Call mit `\n`-Trennung, NICHT 3× einzeln. Analog für `set_preference`.

**PRÄFERENZ-INVERSION-SCHUTZ** (kritisch — Beschwerden ≠ Anweisungen!):
Wenn User negativ über dein Verhalten spricht ("du machst X", "du springst zu wild",
"du reißt aus dem Kontext", "du verstehst nicht"), ist das eine BESCHWERDE über
ungewolltes Verhalten — NIEMALS direkt als Regel speichern!

Vor JEDEM `set_preference`-Call diese 2 Schritte:
1. PARAPHRASIEREN als positive Verhaltens-Regel: "Beschwerde: 'du machst X' → Regel: 'Bot soll NICHT X tun' bzw. 'Bot soll Y tun statt X'"
2. RÜCKBESTÄTIGEN: "Verstehe ich richtig: du willst dass ich künftig [paraphrasierte positive Regel]? Speichere ich es so?"

Erst nach User-OK speichern. Beispiel-Falle:
- User: "du springst wild zwischen Themen, bleib im Kontext!"
- FALSCH: set_preference("Bot soll wild zwischen Themen springen.")
- RICHTIG: nachfragen "Verstehe richtig: du willst dass ich im Konversationsfaden bleibe und Themen nicht abrupt wechsle? OK speichern?"

# KORREKTUREN LERNEN
Klare Korrektur-Trigger ("nein, anders", "das war falsch", "ich meinte X", "mach das anders",
"stopp falsch", "besser X statt Y", "nein STOP", "Bist du dumm", mehrere "!!!") →
`log_correction(was_falsch, was_richtig, kontext)`. Landet in corrections.jsonl für
Fine-Tuning + Nightly-Vorschläge. RUFE DAS AUCH BEI FRUST-ERKENNUNG (siehe oben).

# NIGHTLY MEMORY/HEALTH-VORSCHLÄGE
Antworten auf Memory- oder Health-Listen ("1 3", "alle", "0", "erkläre 2", oder mit Präfix "memory 1" / "health alle") werden DIREKT vom Bot-Loop verarbeitet, NICHT vom LLM. Du bekommst diese Pattern-Messages gar nicht zu sehen — wenn doch eine kommt heißt das Disambig war unklar (beide pending), dann frag User welche Liste gemeint ist.

# AUTO-TAGGING
Bei `create_task/note/meeting`: 2-5 thematische Tags via `tags`-Parameter (kebab-case Deutsch, z.B. `gesundheit`, `kunde-mueller`). TOPISCH, nicht strukturell (NICHT `task`/`note`). Bei vage Kontext lieber `[]` als schlechte Tags. Optional vorher `list_existing_tags` für Vokabular-Konsistenz.

# KONVERSATIONS-KONTEXT NUTZEN
Wenn vorhin ein Projekt erstellt/aktiviert wurde und User danach Tasks anlegt die offensichtlich dazu gehören → `project=<slug>` AUTOMATISCH setzen, nicht nachfragen.
"""

# ============================================================================
# LLM tool-use loop
# ============================================================================

# Sensitive-Pattern-Maskierung für Error-Messages die in History/User landen
_SENSITIVE_PATTERNS = [
    (re.compile(r"sk-[A-Za-z0-9_\-]{20,}"), "[REDACTED-API-KEY]"),       # OpenAI/Anthropic
    (re.compile(r"Bearer\s+[A-Za-z0-9_\-\.]{20,}"), "Bearer [REDACTED]"),
    # Telegram-Bot-Token: <8-12 digits>:<base64-ish 35 chars>
    (re.compile(r"\b\d{7,15}:[A-Za-z0-9_\-]{30,}"), "[REDACTED-TG-TOKEN]"),
    (re.compile(r"ghp_[A-Za-z0-9]{30,}"), "[REDACTED-GH-PAT]"),           # GitHub PAT classic
    (re.compile(r"github_pat_[A-Za-z0-9_]{50,}"), "[REDACTED-GH-PAT]"),   # GitHub fine-grained
    (re.compile(r"https://[^@\s]+:[^@\s]+@"), "https://[REDACTED]@"),     # URL-mit-Credentials
]


def _sanitize_error(msg: str) -> str:
    """Maskiert Tokens/Keys/Credentials in Error-Strings.

    Wird auf Tool-Fehler-Messages angewandt bevor sie in History oder User-
    Reply landen. Verhindert dass Stack-Traces oder Lib-Errors versehentlich
    Bearer-Tokens, OpenAI-Keys, Telegram-Bot-Tokens leaken.
    """
    if not isinstance(msg, str):
        msg = str(msg)
    for pat, replacement in _SENSITIVE_PATTERNS:
        msg = pat.sub(replacement, msg)
    return msg


# ─── LLM-API-Retry mit exponential backoff ──────────────────────────────────
# Ollama/OpenRouter/OpenAI haben gelegentlich transiente Fehler (502, timeout,
# rate-limit). Vorher: Exception bubble → llm_loop crashed → User muss neu
# schreiben. Jetzt: 3 Versuche mit 1s/2s/4s Backoff, dann erst geben wir auf.

LLM_RETRY_ATTEMPTS = 3
LLM_RETRY_BASE_DELAY = 1.0  # 1s, 2s, 4s

# Tool-Hard-Timeout: egal was ein Tool tut, nach 90s abbrechen.
# Real-World-Anker: backup_vault dauert ~30s, extract_pdf_text ~10-20s
# bei großem PDF, search_vault timeout intern auf 30s. 90s ist großzügig.
TOOL_TIMEOUT_SEC = 90

# Welche Exception-Typen retry-würdig sind (transient). Andere werfen direkt.
def _is_retriable_llm_error(e: Exception) -> bool:
    name = type(e).__name__
    # OpenAI-SDK + httpx Standard-Errors die transient sein können
    transient_names = (
        "APIConnectionError", "APITimeoutError", "InternalServerError",
        "RateLimitError", "ConnectionError", "TimeoutException",
        "ReadTimeout", "ConnectTimeout", "RemoteProtocolError",
    )
    if name in transient_names:
        return True
    # HTTPStatusError: 408/429/500/502/503/504 → retry, alles andere nicht
    status = getattr(e, "status_code", None) or getattr(getattr(e, "response", None), "status_code", None)
    if status in (408, 429, 500, 502, 503, 504):
        return True
    # Generic check via String-Match (Fallback)
    msg = str(e).lower()
    if any(s in msg for s in ("timeout", "connection", "rate limit", "502", "503", "504")):
        return True
    return False


# History-Token-Budget — verhindert Context-Overflow bei langen Sessions.
# Konservativ: 80k von typisch 128k Context-Window, Rest für Tool-Schemas
# (~3k) + Response (~8k) + Safety-Buffer.
HISTORY_TOKEN_BUDGET = 80_000


def _estimate_tokens(msg_list: list) -> int:
    """Rough char-based estimate: ~4 chars/token. Schnell + ohne extra Dep."""
    total = 0
    for m in msg_list:
        content = m.get("content", "")
        if isinstance(content, list):
            for c in content:
                if isinstance(c, dict):
                    total += len(str(c.get("text", "")))
                elif isinstance(c, str):
                    total += len(c)
        else:
            total += len(str(content))
        # Tool-calls dranzählen (Args sind JSON-strings)
        for tc in m.get("tool_calls", []) or []:
            try:
                total += len(json.dumps(tc, default=str))
            except Exception:
                pass
    return total // 4


def _truncate_history_for_budget(messages: list, budget: int = HISTORY_TOKEN_BUDGET) -> list:
    """Trim älteste history-Messages bis Token-Budget passt.

    Invarianten:
      - messages[0] (system) immer behalten
      - messages[-1] (neue user-message) immer behalten
      - middle wird in 2er-Schritten getrimmt (pair-aware: assistant + tool)
      - keine orphaned tool-Messages am middle-start (würde LLM verwirren)
    """
    if not messages or len(messages) <= 2:
        return messages
    if _estimate_tokens(messages) <= budget:
        return messages

    system = messages[0]
    last = messages[-1]
    middle = list(messages[1:-1])
    dropped = 0

    while middle and _estimate_tokens([system] + middle + [last]) > budget:
        # 2 raus, plus orphaned tool-message am Anfang weg
        middle = middle[2:]
        dropped += 2
        # Wenn middle[0] orphaned tool-message → auch raus
        while middle and middle[0].get("role") == "tool":
            middle = middle[1:]
            dropped += 1

    if dropped:
        log.warning(f"history-truncate: {dropped} alte Messages gecuttet (Budget {budget} tokens)")
    return [system] + middle + [last]


async def _llm_call_with_retry(client, **kwargs):
    """Wrapper um client.chat.completions.create mit Retry+Backoff.

    Bei retry-würdigem Error: bis zu 3 Versuche, exponential backoff.
    Bei nicht-retry-würdigem Error: sofort raise.
    """
    last_exc = None
    for attempt in range(LLM_RETRY_ATTEMPTS):
        try:
            return await asyncio.to_thread(client.chat.completions.create, **kwargs)
        except Exception as e:
            last_exc = e
            if not _is_retriable_llm_error(e):
                raise
            if attempt < LLM_RETRY_ATTEMPTS - 1:
                delay = LLM_RETRY_BASE_DELAY * (2 ** attempt)
                log.warning(
                    f"LLM-Call fehlgeschlagen (attempt {attempt+1}/{LLM_RETRY_ATTEMPTS}, "
                    f"retry in {delay}s): {type(e).__name__}: {str(e)[:200]}"
                )
                await asyncio.sleep(delay)
    # Alle Versuche aufgebraucht
    raise last_exc


async def llm_loop(user_text: str, user_id: int) -> str:
    """Run tool-use loop until final answer or limit reached.

    Mit Conversation-Memory: letzte ~12 Turns werden als Context übergeben.
    """
    now_local = datetime.now(TIMEZONE)
    sys_text = (SYSTEM_PROMPT
                .replace("{today}", today_iso())
                .replace("{now}", now_local.strftime("%H:%M"))
                .replace("{tz}", TIMEZONE.key if hasattr(TIMEZONE, "key") else str(TIMEZONE)))

    # ─── Memory-Tiers in System-Prompt einspeisen ───
    prefs = get_preferences()
    if prefs:
        sys_text += f"\n\n# PRÄFERENZEN (Stil/Tonalität — befolge diese)\n\n{prefs}\n"
    facts = get_facts()
    if facts:
        sys_text += f"\n\n# PERSISTENTE FAKTEN (Hintergrund über Julius)\n\n{facts}\n"
    active_proj = get_active_project()
    if active_proj:
        proj_ctx = get_project_context(active_proj)
        if proj_ctx:
            sys_text += f"\n\n# AKTIVES PROJEKT: {active_proj}\n\n{proj_ctx}\n"
        else:
            sys_text += f"\n\n# AKTIVES PROJEKT: {active_proj} (keine CONTEXT.md gesetzt)\n"

    # System-Prompt + History + neue User-Message
    history = await get_history(user_id)
    new_user_msg = {"role": "user", "content": user_text}
    # Provider-aware System-Message-Format:
    # Anthropic akzeptiert content-as-list mit cache_control (Prompt-Caching).
    # Gemini/OpenAI/Ollama erwarten content als plain String — sonst 400.
    if USE_ANTHROPIC_CACHE:
        system_msg = {
            "role": "system",
            "content": [
                {
                    "type": "text",
                    "text": sys_text,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
        }
    else:
        system_msg = {"role": "system", "content": sys_text}
    messages = [system_msg] + history + [new_user_msg]

    # Token-Budget-Truncation: bei langen Sessions würde messages das
    # Context-Window von Ollama (128k) sprengen. Wir trimmen ältere
    # history-Messages BEFORE jedem LLM-Call (system + neueste user
    # bleiben immer drin).
    messages = _truncate_history_for_budget(messages)

    # Diese Messages werden am Ende zur History dazugefügt
    new_history_msgs = [new_user_msg]

    # Iterations-Limit: 25 Iterationen reichen auch für Multi-File-Workflows
    # (6 Uploads → Projekt anlegen → bulk-move → activate ≈ 4 Calls bei Bulk-Tools).
    # Wenn LLM single-shot statt bulk arbeitet, fängt's spätestens bei 15 mit
    # einem Hint zu Bulk-Tools auf, statt blind weiterzumachen.
    LOOP_LIMIT = 25
    BULK_HINT_AT = 15
    bulk_hint_sent = False
    completed_tool_calls = []  # für End-Of-Loop-Error-Message
    # Self-Healing: wenn dasselbe Tool 3× in Folge fehlschlägt, abbrechen
    # statt blind weiter zu loopen. Vermeidet Frust + Token-Verschwendung.
    SELF_HEAL_FAIL_THRESHOLD = 3
    consecutive_failures: dict = {}  # tool_name → fail-count in Folge

    for iteration in range(LOOP_LIMIT):
        # Innerhalb des Loops können messages durch Tool-Outputs wachsen
        # (z.B. read_file mit 8KB content). Vor jedem LLM-Call re-trim.
        messages = _truncate_history_for_budget(messages)
        # Sanitize-on-Send: Falls irgendwo eine kaputte tool-msg ohne
        # name oder eine assistant-msg mit tool_calls aber content=None
        # entstanden ist → reparieren bevor sie an den Provider geht.
        # Idempotent + günstig, schützt vor Provider-400ern (besonders Gemini).
        messages = _sanitize_loaded_history(messages)
        resp = await _llm_call_with_retry(
            llm,
            model=LLM_MODEL,
            messages=messages,
            tools=TOOLS,
            tool_choice="auto",
            max_tokens=2048,
        )
        # Token-Usage tracken (best-effort, nie crash-causing)
        try:
            usage = getattr(resp, "usage", None)
            if usage is not None:
                _track_usage(LLM_MODEL,
                             getattr(usage, "prompt_tokens", 0),
                             getattr(usage, "completion_tokens", 0),
                             kind="chat")
        except Exception:
            pass
        msg = resp.choices[0].message
        msg_dict = msg.model_dump(exclude_none=True)
        messages.append(msg_dict)
        new_history_msgs.append(msg_dict)

        if not msg.tool_calls:
            await update_history(user_id, new_history_msgs)
            return msg.content or "(keine Antwort)"

        for tc in msg.tool_calls:
            tool_failed = False  # für Self-Healing-Counter
            try:
                args = json.loads(tc.function.arguments) if tc.function.arguments else {}
                handler = TOOL_HANDLERS.get(tc.function.name)
                if not handler:
                    result = f"Tool nicht bekannt: {tc.function.name}"
                    tool_failed = True
                else:
                    log.info(f"tool[{iteration+1}/{LOOP_LIMIT}]: {tc.function.name}({args})")
                    # KRITISCH: handler in Threadpool — sonst blockiert
                    # backup_vault/clip_url/_build_link_index/extract_pdf_text
                    # den ganzen Telegram-Event-Loop minutenlang.
                    # PLUS: hard timeout pro Tool — egal was es tut, nach
                    # TOOL_TIMEOUT_SEC ist Schluss. Verhindert dass ein
                    # hängendes Tool den ganzen Bot lahmlegt.
                    tool_start = time.time()
                    try:
                        result = await asyncio.wait_for(
                            asyncio.to_thread(handler, **args),
                            timeout=TOOL_TIMEOUT_SEC,
                        )
                    except asyncio.TimeoutError:
                        result = (f"Tool-Timeout: `{tc.function.name}` lief länger "
                                  f"als {TOOL_TIMEOUT_SEC}s und wurde abgebrochen.")
                        tool_failed = True
                        log.warning(f"tool[{tc.function.name}] timeout after {TOOL_TIMEOUT_SEC}s")
                    else:
                        elapsed = time.time() - tool_start
                        if elapsed > 5:
                            log.info(f"tool[{tc.function.name}] dauerte {elapsed:.1f}s")
                        completed_tool_calls.append(tc.function.name)
                        # Heuristic: Tool-Result der mit "Fehler"/"failed" anfängt → fail
                        if isinstance(result, str) and re.match(
                            r"^(Fehler|❌|Pfad-Fehler|Tool nicht|Ungültig|Tool-Timeout)", result
                        ):
                            tool_failed = True
            except Exception as e:
                log.exception(f"Tool {tc.function.name} failed")
                # Token/Credential-Maskierung — Error landet in LLM-History
                # und ggf in User-Reply, darf keine API-Keys leaken
                result = _sanitize_error(f"Tool-Fehler: {e}")
                tool_failed = True
            # Self-Healing: track consecutive failures pro Tool
            tn = tc.function.name
            if tool_failed:
                consecutive_failures[tn] = consecutive_failures.get(tn, 0) + 1
            else:
                consecutive_failures[tn] = 0  # Erfolg → Counter reset
            tool_msg = {
                "role": "tool",
                "tool_call_id": tc.id,
                # name-Feld nötig für Gemini's OpenAI-Kompat-Endpoint —
                # OpenRouter/Ollama akzeptieren auch ohne, Gemini wirft 400.
                "name": tc.function.name,
                "content": str(result),
            }
            messages.append(tool_msg)
            new_history_msgs.append(tool_msg)
            # Hard-Break wenn ein Tool 3× in Folge failed
            if consecutive_failures.get(tn, 0) >= SELF_HEAL_FAIL_THRESHOLD:
                await update_history(user_id, new_history_msgs)
                return (
                    f"⚠️ Tool `{tn}` ist {SELF_HEAL_FAIL_THRESHOLD}× in Folge fehlgeschlagen — "
                    f"breche ab statt weiter zu loopen.\n\n"
                    f"Letzter Fehler: {str(result)[:300]}\n\n"
                    f"Bitte prüfe ob das Tool ein Problem hat oder formuliere die Anfrage anders."
                )

        # Hint einschleusen wenn LLM bei langen Loops noch single-shot arbeitet.
        # Seit Tool-Konsolidierung gibt's nur noch `move` — mehrfach-Single-Calls
        # erkennt man am gleichen Tool-Namen mit string-src statt list.
        if iteration + 1 == BULK_HINT_AT and not bulk_hint_sent:
            move_count = sum(1 for c in completed_tool_calls if c == "move")
            if move_count >= 4:
                hint = {
                    "role": "user",
                    "content": (
                        f"[System-Hint]: Du hast bereits {move_count}× move einzeln "
                        f"aufgerufen. Nutze `move(src=[liste], dst='ordner/')` für Bulk in "
                        f"EINEM Call. Schließe die Operation jetzt zügig ab, sonst reicht "
                        f"das Iterations-Limit nicht."
                    ),
                }
                messages.append(hint)
                new_history_msgs.append(hint)
                bulk_hint_sent = True

    # Loop-Limit erreicht — User erfährt was tatsächlich passiert ist
    await update_history(user_id, new_history_msgs)
    summary = ", ".join(f"{n}×{t}" for t, n in
                        sorted({c: completed_tool_calls.count(c)
                                for c in set(completed_tool_calls)}.items(),
                               key=lambda x: -x[1]))
    return (
        f"⚠️ Operation zu komplex für eine Iteration ({LOOP_LIMIT} Tool-Calls verbraucht).\n\n"
        f"Bisher gemacht: {summary or '(nichts erfolgreich)'}.\n\n"
        f"Der Auftrag ist möglicherweise nur teilweise erledigt — "
        f"bitte prüfen oder nochmal mit präziserem/kleinerem Auftrag wiederholen. "
        f"Tipp: `move(src=[liste], dst='ordner/')` für Bulk statt 6× einzeln."
    )


# ============================================================================
# Telegram handlers
# ============================================================================

def require_auth(handler):
    async def wrapper(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
        uid = update.effective_user.id if update.effective_user else None
        if uid is None:
            return
        # Setup-Modus: noch kein User gebunden → Bot meldet User-ID, sonst nichts
        if ALLOWED_USER_ID == 0:
            log.info(f"Setup mode: first contact from user_id={uid}")
            await update.message.reply_text(
                f"🔓 <b>Setup-Modus</b>\n\n"
                f"Bot ist noch nicht an einen User gebunden.\n\n"
                f"Deine Telegram-User-ID: <code>{uid}</code>\n\n"
                f"Auf VPS:\n"
                f"<pre><code>nano /opt/bot/.env\n"
                f"# ALLOWED_USER_ID={uid} setzen\n"
                f"docker compose restart</code></pre>\n"
                f"Danach bin ich nur noch für dich da.",
                parse_mode=constants.ParseMode.HTML,
            )
            return
        if uid != ALLOWED_USER_ID:
            log.warning(f"Unauthorized access attempt: user_id={uid}")
            return
        return await handler(update, ctx)
    return wrapper


def _esc_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# Telegram-Mobile rendert ~36-40 Zeichen pro Zeile in Pre-Blocks ohne Umbruch.
# Tabellen die breiter sind verlieren Spalten-Ausrichtung → unleserlich.
TELEGRAM_TABLE_MAX_WIDTH = 38


def _render_md_table_html(table_text: str) -> str:
    """Markdown-Tabelle → Telegram-passendes HTML.

    Strategie:
      - schmale 2-Spalten-Tabellen (≤TELEGRAM_TABLE_MAX_WIDTH gesamt) → "Key: Value"-Liste
      - schmale n-Spalten-Tabellen (≤TELEGRAM_TABLE_MAX_WIDTH) → monospaced <pre>
      - sonst → Sektions-Layout: pro Row ein Block mit Bold-Header + Bullets je Spalte
    """
    lines = [l for l in table_text.strip().split("\n") if l.strip()]
    if len(lines) < 3:
        return f"<pre><code>{_esc_html(table_text)}</code></pre>"

    def cells(line: str) -> list:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    header = cells(lines[0])
    rows = [cells(l) for l in lines[2:]]  # lines[1] ist Separator
    n = len(header)
    rows = [r[:n] + [""] * max(0, n - len(r)) for r in rows]

    widths = [len(c) for c in header]
    for r in rows:
        for i in range(n):
            widths[i] = max(widths[i], len(r[i]))

    total_width = sum(widths) + (n - 1) * 3  # " │ "-Separatoren

    # ─── A) Schmal genug für monospaced Pre? Behält Tabellen-Look. ───
    if total_width <= TELEGRAM_TABLE_MAX_WIDTH:
        def fmt_row(cells_):
            return " │ ".join(c.ljust(widths[i]) for i, c in enumerate(cells_))
        sep = "─┼─".join("─" * w for w in widths)
        out = [fmt_row(header), sep] + [fmt_row(r) for r in rows]
        return f"<pre><code>{_esc_html(chr(10).join(out))}</code></pre>"

    # ─── B) Genau 2 Spalten? Kompakte "Key: Value"-Liste. ───
    if n == 2:
        out_lines = []
        for r in rows:
            key, val = r[0], r[1]
            if not key and not val:
                continue
            out_lines.append(f"<b>{_esc_html(key)}:</b> {_esc_html(val)}")
        return "\n".join(out_lines)

    # ─── C) ≥3 Spalten & breit → Sektions-Layout. ───
    # Erste Spalte = Item-Header (bold), restliche = Bullet-Liste mit "Header: Wert".
    out_lines = []
    for r in rows:
        primary = r[0].strip()
        if not primary and not any(c.strip() for c in r[1:]):
            continue
        out_lines.append(f"<b>{_esc_html(primary or '—')}</b>")
        for i in range(1, n):
            val = r[i].strip()
            if not val:
                continue
            col_header = header[i].strip() if i < len(header) else ""
            if col_header:
                out_lines.append(f"  • <b>{_esc_html(col_header)}:</b> {_esc_html(val)}")
            else:
                out_lines.append(f"  • {_esc_html(val)}")
        out_lines.append("")  # Leerzeile zwischen Items
    # Trailing-Leerzeile abschneiden
    while out_lines and not out_lines[-1]:
        out_lines.pop()
    return "\n".join(out_lines)


# Markdown-Tabelle: Header-Zeile + Separator-Zeile (nur -:| und Spaces) + 1+ Datenzeilen
TABLE_RE = re.compile(
    r"(^\|[^\n]+\|[ \t]*\n"            # Header
    r"\|[ \t]*[-:][\-:| \t]*\|[ \t]*\n"  # Separator
    r"(?:\|[^\n]+\|[ \t]*\n?)+)",      # Daten (1 oder mehr)
    re.MULTILINE,
)


def md_to_telegram_html(text: str) -> str:
    """Konvertiere Markdown → Telegram-kompatibles HTML.

    Telegram-Subset: <b>, <i>, <u>, <s>, <a>, <code>, <pre>, <blockquote>.
    Block-Konstrukte (Tabellen, Listen, Headings) werden zu Inline-Formaten:
    - Tabellen → monospaced <pre> mit Box-Drawing-Chars
    - Headings → <b>
    - Bullets → Unicode •
    - Numbered Lists → bleiben "1. text"
    - Horizontale Linien → ━━━━━━━━━━━━
    """
    # Stash für bereits-fertiges HTML, das nicht weiter verarbeitet werden soll
    stash = []

    def add_stash(html_fragment: str) -> str:
        stash.append(html_fragment)
        return f"\x00S{len(stash)-1}\x00"

    # 0) LLM-HTML-Tags zu Newlines/Plain umwandeln, BEVOR escape rennt.
    #    gpt-oss & Co. bauen trotz "NIE HTML"-Prompt-Regel <br>, <p>, <li> etc.
    #    rein. Würde später als &lt;br&gt; im Telegram landen → unleserlich.
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</?p\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</?div\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<li\s*/?>", "\n• ", text, flags=re.IGNORECASE)
    text = re.sub(r"</li>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"</?[ou]l\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<h[1-6]\s*/?>", "\n**", text, flags=re.IGNORECASE)
    text = re.sub(r"</h[1-6]>", "**\n", text, flags=re.IGNORECASE)
    # Kollabiere ≥3 Newlines zu max 2 (sonst tonnen Leerzeilen)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 1) TABELLEN — Renderer entscheidet je nach Breite/Spaltenzahl:
    #    schmal → monospaced <pre>, 2-Spalten breit → "Key: Value"-Liste,
    #    ≥3 Spalten breit → Sektions-Layout. Renderer liefert fertiges HTML.
    def _table_repl(m):
        return add_stash(_render_md_table_html(m.group(1)))
    text = TABLE_RE.sub(_table_repl, text)

    # 2) FENCED CODE BLOCKS (```...```)
    def _fenced_repl(m):
        return add_stash(f"<pre><code>{_esc_html(m.group(2))}</code></pre>")
    text = re.sub(r"```(\w+)?\n?(.*?)```", _fenced_repl, text, flags=re.DOTALL)

    # 3) INLINE CODE (`...`)
    def _inline_repl(m):
        return add_stash(f"<code>{_esc_html(m.group(1))}</code>")
    text = re.sub(r"`([^`\n]+)`", _inline_repl, text)

    # 4) Restlichen Text HTML-escapen
    text = _esc_html(text)

    # 5) Headings (# bis ######) → <b> + Newline davor für visuelle Trennung
    text = re.sub(r"^#{1,6}\s+(.+?)$", r"<b>\1</b>", text, flags=re.MULTILINE)

    # 6) Bold/Italic/Strike
    text = re.sub(r"\*\*([^*\n]+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"__([^_\n]+?)__", r"<b>\1</b>", text)
    text = re.sub(r"(?<![*\w])\*([^*\n]+?)\*(?!\w)", r"<i>\1</i>", text)
    text = re.sub(r"~~([^~\n]+?)~~", r"<s>\1</s>", text)

    # 7) Blockquotes (> text)  — beachte: > wurde zu &gt; escaped
    text = re.sub(
        r"(?:^&gt;\s?.+(?:\n|$))+",
        lambda m: "<blockquote>" + re.sub(r"^&gt;\s?", "", m.group(0), flags=re.MULTILINE).rstrip() + "</blockquote>\n",
        text,
        flags=re.MULTILINE,
    )

    # 8) Horizontale Linie (--- oder *** allein auf Zeile)
    text = re.sub(r"^[-*_]{3,}\s*$", "━" * 24, text, flags=re.MULTILINE)

    # 9) Links — Telegram akzeptiert nur absolute URLs in <a href>.
    # Echte http/https-Links → klickbar. Relative Pfade (../foo.md) → nur Text kursiv,
    # damit kein doppelter Markdown-Salat in Telegram entsteht.
    text = re.sub(
        r"\[([^\]]+)\]\((https?://[^)]+)\)",
        r'<a href="\2">\1</a>',
        text,
    )
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"<i>\1</i>", text)

    # 10) Wikilinks [[id]] → in <code> einpacken damit Telegram nicht auto-linkt
    # (z.B. wenn id "2026-04-25.md" enthält, würde Telegram .md → URL machen)
    text = re.sub(r"\[\[([^\]]+)\]\]", r"<code>[[\1]]</code>", text)

    # 11) Bullet-Listen mit Indentation → Unicode • mit erhaltener Einrückung
    def _bullet_repl(m):
        indent = m.group(1)
        content = m.group(2)
        # Verschachtelte Bullets: 2 Leerzeichen pro Ebene → ◦ statt •
        depth = len(indent) // 2
        marker = "•" if depth == 0 else ("◦" if depth == 1 else "▪")
        return f"{indent}{marker} {content}"
    text = re.sub(r"^([ \t]*)[-*+]\s+(.+)$", _bullet_repl, text, flags=re.MULTILINE)

    # 12) Numbered Lists — bleiben als "1. text", aber Spaces normalisieren
    text = re.sub(r"^([ \t]*)(\d+)\.\s+(.+)$", r"\1\2. \3", text, flags=re.MULTILINE)

    # 13) Stashed HTML restoren
    def _restore(m):
        return stash[int(m.group(1))]
    text = re.sub(r"\x00S(\d+)\x00", _restore, text)

    return text


def _strip_html(text: str) -> str:
    """Fallback: entferne HTML-Tags für Plain-Text-Send."""
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return text


def _strip_paren_wikilinks(text: str) -> str:
    """Entfernt redundante Slug-Wikilinks in Klammern: 'Foo ([[t-foo]])' → 'Foo'.

    Hintergrund: Bei Status-Bestätigungen (mark_task_done/delete/move) hängt
    der LLM trotz Prompt-Regel oft noch '([[t-slug]])' an den Titel — purer
    Lärm, weil der User den Task gerade selbst genannt hat. Wir strippen das
    deterministisch raus. Bare Wikilinks (nicht in Klammern) bleiben erhalten,
    weil das echte Klick-Anker zu Neu-Erstellungen sind.
    """
    if not text:
        return text
    # `([[...]])` mit optional Whitespace davor — Multi-line safe
    return re.sub(r"[ \t]*\(\[\[[^\[\]\n]+\]\]\)", "", text)


async def safe_reply(update: Update, text: str) -> None:
    """Split + send. HTML-Format mit Plain-Fallback bei Parse-Fehler."""
    if not text:
        await update.message.reply_text("(leer)")
        return

    # Noise-Wikilinks rausstrippen, BEVOR Markdown→HTML konvertiert
    text = _strip_paren_wikilinks(text)

    # Erst Markdown→HTML, dann splitten
    html = md_to_telegram_html(text)

    # Splitten an Newlines wo möglich (Smart-Breaks)
    chunks = []
    remaining = html
    while remaining:
        if len(remaining) <= TG_MAX_MESSAGE:
            chunks.append(remaining)
            break
        cut = remaining.rfind("\n\n", 0, TG_MAX_MESSAGE)
        if cut < 1000:
            cut = remaining.rfind("\n", 0, TG_MAX_MESSAGE)
        if cut < 1000:
            cut = remaining.rfind(" ", 0, TG_MAX_MESSAGE)
        if cut < 1:
            cut = TG_MAX_MESSAGE
        chunks.append(remaining[:cut])
        remaining = remaining[cut:].lstrip()

    for i, chunk in enumerate(chunks, 1):
        prefix = f"({i}/{len(chunks)})\n" if len(chunks) > 1 else ""
        try:
            await update.message.reply_text(
                prefix + chunk,
                parse_mode=constants.ParseMode.HTML,
                disable_web_page_preview=True,
            )
        except Exception as e:
            log.warning(f"HTML-Send fehlgeschlagen, Fallback Plain: {e}")
            await update.message.reply_text(prefix + _strip_html(chunk))


def _detect_pending_reply_intent(text: str) -> Optional[str]:
    """Erkennt ob User-Message eine Antwort auf pending Memory/Health-Liste ist.

    Returns: 'memory' oder 'health' oder None.
    Heuristik:
    - "memory <antwort>" / "memory: <antwort>" → memory
    - "health <antwort>" → health
    - Sonst nur wenn EXAKT ein Pending-File existiert (Disambig)
      UND Text matches typisches Reply-Pattern (Zahlen, "alle", "nein", "0", "ja")
    """
    t = (text or "").strip().lower()
    if not t:
        return None
    # Explizite Präfixe
    if t.startswith(("memory ", "memory:")) or t == "memory":
        return "memory"
    if t.startswith(("health ", "health:")) or t == "health":
        return "health"
    # Typisches Reply-Pattern: nur Zahlen+Spaces+Komma, "alle", "ja", "nein", "0", "skip", "erkläre N"
    is_reply_shape = bool(
        re.fullmatch(r"\s*\d+(\s*[,\s]\s*\d+)*\s*", t)        # "1 2 3" / "1,3"
        or t in ("alle", "ja", "all", "yes", "y",
                 "nein", "no", "n", "0", "skip", "verwerfen")
        or re.match(r"^(erklär|erklar)", t)                    # "erkläre 2"
    )
    if not is_reply_shape:
        return None
    # Disambig: welches Pending-File existiert?
    has_mem = PENDING_SUGGESTIONS_FILE.exists()
    has_health = PENDING_HEALTH_ACTIONS_FILE.exists()
    if has_mem and not has_health:
        return "memory"
    if has_health and not has_mem:
        return "health"
    if has_mem and has_health:
        # Beide pending — ohne Präfix nicht eindeutig, lass LLM entscheiden
        return None
    return None


def _strip_intent_prefix(text: str) -> str:
    """Entfernt 'memory '/'health '-Präfix vom User-Text vor Action-Parser."""
    t = text.strip()
    for prefix in ("memory:", "memory ", "memory", "health:", "health ", "health"):
        if t.lower().startswith(prefix):
            return t[len(prefix):].strip()
    return t


@require_auth
async def handle_text(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = update.message.text or ""
    log.info(f"text: {text[:120]}")
    await update.message.chat.send_action(constants.ChatAction.TYPING)

    # Intent-Detection: ist das eine Antwort auf pending Memory/Health-Liste?
    # Wenn ja: direkt den entsprechenden Action-Parser rufen (kein LLM-Roundtrip).
    intent = _detect_pending_reply_intent(text)
    if intent == "memory":
        action = _strip_intent_prefix(text) or text.strip()
        try:
            reply = await asyncio.to_thread(apply_memory_suggestion, action)
        except Exception as e:
            log.exception("apply_memory_suggestion failed")
            reply = _sanitize_error(f"Fehler: {e}")
        await safe_reply(update, reply)
        return
    if intent == "health":
        action = _strip_intent_prefix(text) or text.strip()
        try:
            reply = await asyncio.to_thread(apply_health_action, action)
        except Exception as e:
            log.exception("apply_health_action failed")
            reply = _sanitize_error(f"Fehler: {e}")
        await safe_reply(update, reply)
        return

    # Normaler Pfad: ans LLM mit Tool-Loop
    try:
        reply = await llm_loop(text, update.effective_user.id)
    except Exception as e:
        log.exception("llm_loop failed")
        reply = f"Fehler: {e}"
    await safe_reply(update, reply)


@require_auth
async def handle_voice(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    voice = update.message.voice or update.message.audio
    if not voice:
        return
    log.info(f"voice: duration={voice.duration}s")
    await update.message.chat.send_action(constants.ChatAction.TYPING)
    try:
        file = await voice.get_file()
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
            tmp_path = tmp.name
        await file.download_to_drive(tmp_path)
        try:
            segments, _ = await asyncio.to_thread(
                whisper.transcribe, tmp_path, language=WHISPER_LANG
            )
            transcript = " ".join(seg.text.strip() for seg in segments).strip()
        finally:
            try: os.unlink(tmp_path)
            except Exception: pass
        if not transcript:
            await update.message.reply_text("(Sprachnachricht leer/unverständlich)")
            return
        log.info(f"transcript: {transcript[:120]}")
        await update.message.reply_text(f"📝 {transcript}")
        reply = await llm_loop(transcript, update.effective_user.id)
    except Exception as e:
        log.exception("voice handler failed")
        reply = f"Voice-Fehler: {e}"
    await safe_reply(update, reply)


@require_auth
async def handle_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    photo = update.message.photo[-1] if update.message.photo else None
    if not photo:
        return
    user_caption = update.message.caption or ""
    log.info(f"photo: file_id={photo.file_id[:20]}..., caption={user_caption[:60]}")
    await update.message.chat.send_action(constants.ChatAction.TYPING)
    try:
        file = await photo.get_file()
        ATTACHMENTS_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        ext = "jpg"
        if file.file_path and "." in file.file_path:
            ext = file.file_path.rsplit(".", 1)[-1].lower()
        filename = f"photo_{ts}.{ext}"
        save_path = ATTACHMENTS_DIR / filename
        await file.download_to_drive(str(save_path))
        # Vision-Caption — Base64-Encode in Threadpool (sync read+encode auf
        # 5MB-Photos kann Event-Loop spürbar blockieren)
        def _read_and_encode():
            with open(save_path, "rb") as f:
                return base64.b64encode(f.read()).decode()
        b64 = await asyncio.to_thread(_read_and_encode)
        prompt_text = (
            "Beschreibe in 1-2 Sätzen knapp, was auf dem Bild ist."
            f" Kontext vom User: \"{user_caption}\"" if user_caption else
            "Beschreibe in 1-2 Sätzen knapp, was auf dem Bild ist."
        )
        vision_resp = await _llm_call_with_retry(
            vision_llm,
            model=VISION_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}},
                ],
            }],
            max_tokens=300,
        )
        try:
            vusage = getattr(vision_resp, "usage", None)
            if vusage is not None:
                _track_usage(VISION_MODEL,
                             getattr(vusage, "prompt_tokens", 0),
                             getattr(vusage, "completion_tokens", 0),
                             kind="vision")
        except Exception:
            pass
        vision_caption = (vision_resp.choices[0].message.content or "(keine Beschreibung)").strip()

        # OCR (Tesseract) — parallel zur Vision-Caption, macht Bild-Text suchbar
        ocr_text = await asyncio.to_thread(ocr_image, save_path)
        ocr_text = ocr_text.strip()
        # Sehr kurze OCR-Ergebnisse (<10 Zeichen) sind meist Rauschen
        ocr_block = ""
        if len(ocr_text) >= 10:
            # Auf 800 Zeichen begrenzen für Daily-Eintrag, voller Text bleibt im File-Namen referenzierbar
            ocr_truncated = ocr_text[:800] + ("..." if len(ocr_text) > 800 else "")
            ocr_block = f"\n  *OCR*:\n  > {ocr_truncated}"

        # Link in Daily — mit Vision-Caption + OCR (falls was erkannt)
        if user_caption:
            link_block = f"![[{filename}]] — {user_caption}\n  *Vision*: {vision_caption}{ocr_block}"
        else:
            link_block = f"![[{filename}]] — {vision_caption}{ocr_block}"
        try:
            append_to_daily("Notizen & Gedanken", link_block)
        except Exception as e:
            log.warning(f"Daily-Link für Photo fehlgeschlagen: {e}")

        # Reply
        reply = f"🖼 {filename}\n\n<b>Vision</b>: {_esc_html(vision_caption)}"
        if ocr_text:
            reply += f"\n\n<b>OCR</b> ({len(ocr_text)} Zeichen):\n<pre><code>{_esc_html(ocr_text[:1500])}</code></pre>"
        await update.message.reply_text(reply, parse_mode=constants.ParseMode.HTML)

        # Upload-Event in LLM-History
        try:
            hist_msg = (
                f"[Upload-Event] User hat Foto \"{filename}\" hochgeladen, "
                f"gespeichert in `09_Attachments/{filename}`. "
                f"Vision-Beschreibung: {vision_caption[:200]}"
            )
            if ocr_text:
                hist_msg += f"\nOCR-Text:\n{ocr_text[:400]}"
            await update_history(update.effective_user.id, [
                {"role": "user", "content": hist_msg}
            ])
        except Exception as e:
            log.warning(f"Photo-Upload-Event nicht in History: {e}")

        # Caption als Instruction an LLM routen falls vorhanden
        if user_caption and user_caption.strip():
            await update.message.chat.send_action(constants.ChatAction.TYPING)
            try:
                llm_reply = await llm_loop(user_caption.strip(), update.effective_user.id)
                await safe_reply(update, llm_reply)
            except Exception as e:
                log.exception("Photo-Caption-LLM-Routing failed")
    except Exception as e:
        log.exception("photo handler failed")
        await update.message.reply_text(f"Photo-Fehler: {e}")


# ─── Document-Upload-Helpers ────────────────────────────────────────────────
# handle_document war 155 LOC, machte Download + Speichern + PDF-Wrap +
# Daily-Eintrag + Reply + History + Caption-Routing in einer Funktion.
# Aufgespalten in drei sync Helper + dünner async Orchestrator.

def _sanitize_filename(filename: str) -> str:
    """Filename-Sanitization gegen Path-Traversal + OS-reservierte Namen.

    Telegram-Client kann beliebige Filenames setzen (inkl. '../' oder
    Windows-reservierte 'con.txt'/'nul.md'). Path(...).name strippt
    Pfad-Komponenten, regex normalisiert verbleibende Sonderzeichen.
    """
    # Schritt 1: Pfad-Komponenten weg ('../foo' → 'foo')
    base = Path(filename).name or "upload"
    # Schritt 2: nur kebab-case Buchstaben/Ziffern/Punkt/Bindestrich/Unterstrich
    cleaned = re.sub(r"[^\w.\-]", "_", base)
    # Schritt 3: leading/trailing dots+spaces weg (Windows-Edge-Case)
    cleaned = cleaned.strip(". ") or "upload"
    # Schritt 4: max 200 Zeichen (Filesystem-Limit-Buffer)
    if len(cleaned) > 200:
        stem, _, ext = cleaned.rpartition(".")
        cleaned = stem[:195] + "." + ext if ext else cleaned[:200]
    return cleaned


def _save_uploaded_doc(tmp_path: str, filename: str) -> tuple[Path, str, str]:
    """Verschiebt tmp-Datei in passenden Vault-Ordner, liefert (dest, kind, preview).

    kind: 'Text/Markdown' / 'PDF' / 'Datei'
    preview: erste 1500 Zeichen bei .md/.txt, sonst leer.

    SECURITY: Filename wird VOR jeder Pfad-Konstruktion sanitisiert
    (gegen Telegram-injizierte Path-Traversal wie '../../etc/passwd').
    """
    filename = _sanitize_filename(filename)
    ext = Path(filename).suffix.lower()
    if ext in (".md", ".markdown", ".txt"):
        dest_dir = VAULT / "01_Raw" / "uploads"
        kind = "Text/Markdown"
    elif ext == ".pdf":
        dest_dir = VAULT / "01_Raw" / "papers"
        kind = "PDF"
    elif ext == ".docx":
        dest_dir = VAULT / "01_Raw" / "uploads"
        kind = "Word-Dokument"
    else:
        dest_dir = VAULT / "09_Attachments"
        kind = "Datei"
    dest_dir.mkdir(parents=True, exist_ok=True)

    # Konflikt-Handling: Suffix-Counter
    dest = dest_dir / filename
    counter = 1
    while dest.exists():
        dest = dest_dir / f"{Path(filename).stem}-{counter}{ext}"
        counter += 1
    shutil.move(tmp_path, dest)

    body_preview = ""
    if ext in (".md", ".markdown", ".txt"):
        try:
            body_preview = dest.read_text(encoding="utf-8", errors="replace")[:1500]
        except Exception:
            body_preview = "(Inhalt nicht lesbar)"
    return dest, kind, body_preview


def _create_pdf_wrapper(dest: Path, filename: str) -> tuple[Path, str, str, int]:
    """PDF-Sibling als .md anlegen (volltext-suchbar via vault_search).

    Returns: (md_path, md_id, pdf_text, total_pages).
    """
    pdf_text, pdf_meta, total_pages = extract_pdf_text(dest)
    pdf_title = pdf_meta.get("title") or Path(filename).stem
    pdf_author = pdf_meta.get("author") or "unknown"

    md_path = dest.with_suffix(".md")
    n = 1
    while md_path.exists():
        md_path = dest.with_name(f"{dest.stem}-{n}.md")
        n += 1

    md_id = slugify(f"paper-{pdf_title}")
    md_body = (
        f"# {pdf_title}\n\n"
        f"📎 [Original PDF: {dest.name}](./{dest.name})\n\n"
        f"**Autor**: {pdf_author} · **Seiten**: {total_pages}"
        + (f" · **Subject**: {pdf_meta['subject']}" if pdf_meta.get('subject') else "")
        + "\n\n---\n\n"
        + (pdf_text or "(Text-Extraktion ergab keinen Inhalt)")
    )
    md_post = frontmatter.Post(
        md_body,
        id=md_id,
        title=pdf_title,
        type="paper",
        source=str(dest.relative_to(VAULT)),
        author=pdf_author,
        captured=today_iso(),
        pages=total_pages,
        tags=[],
    )
    atomic_write(md_path, frontmatter.dumps(md_post) + "\n")
    return md_path, md_id, pdf_text or "", total_pages


def _create_docx_wrapper(dest: Path, filename: str) -> tuple[Path, str, str, int]:
    """DOCX-Sibling als .md anlegen (volltext-suchbar via vault_search).

    Analog zu _create_pdf_wrapper. Tabellen werden als Markdown
    gerendert (in extract_docx_text). type='document' im Frontmatter.

    Returns: (md_path, md_id, docx_text, paragraph_count).
    """
    docx_text, docx_meta, para_count = extract_docx_text(dest)
    docx_title = docx_meta.get("title") or Path(filename).stem
    docx_author = docx_meta.get("author") or "unknown"

    md_path = dest.with_suffix(".md")
    n = 1
    while md_path.exists():
        md_path = dest.with_name(f"{dest.stem}-{n}.md")
        n += 1

    md_id = slugify(f"doc-{docx_title}")
    md_body = (
        f"# {docx_title}\n\n"
        f"📎 [Original DOCX: {dest.name}](./{dest.name})\n\n"
        f"**Autor**: {docx_author} · **Absätze**: {para_count}"
        + (f" · **Subject**: {docx_meta['subject']}" if docx_meta.get('subject') else "")
        + "\n\n---\n\n"
        + (docx_text or "(Text-Extraktion ergab keinen Inhalt)")
    )
    md_post = frontmatter.Post(
        md_body,
        id=md_id,
        title=docx_title,
        type="document",
        source=str(dest.relative_to(VAULT)),
        author=docx_author,
        captured=today_iso(),
        paragraphs=para_count,
        tags=[],
    )
    atomic_write(md_path, frontmatter.dumps(md_post) + "\n")
    return md_path, md_id, docx_text or "", para_count


def _record_upload_in_daily(rel: Path, wrapper_link: Optional[Path],
                            md_id: Optional[str], user_caption: str,
                            kind: str = "Datei") -> None:
    """Eintrag in heutige Daily ('Notizen & Gedanken'-Sektion). Failures werden geloggt."""
    try:
        if wrapper_link and md_id:
            # Format-spezifisches Emoji
            emoji = {"PDF": "📑", "Word-Dokument": "📝"}.get(kind, "📄")
            link_text = f"{emoji} {kind} hochgeladen: <code>{rel}</code> → durchsuchbar als [[{md_id}]]"
        else:
            link_text = f"📄 Datei hochgeladen: <code>{rel}</code>"
        if user_caption:
            link_text += f" — {user_caption}"
        append_to_daily("Notizen & Gedanken", link_text)
    except Exception as e:
        log.warning(f"Daily-Link fuer Document fehlgeschlagen: {e}")


def _build_upload_event_msg(filename: str, kind: str, rel: Path,
                            wrapper_link: Optional[Path], md_id: Optional[str],
                            body_preview: str) -> str:
    """Baut die Upload-Event-Message für die LLM-History."""
    msg = f"[Upload-Event] User hat Datei \"{filename}\" hochgeladen ({kind}), gespeichert als `{rel}`."
    if wrapper_link and md_id:
        msg += f" Wrapper: `{wrapper_link}` (id: `{md_id}`)."
    if body_preview:
        msg += f"\nInhalts-Auszug:\n{body_preview[:400]}"
    return msg


@require_auth
async def handle_document(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Datei-Upload ins Vault speichern. Orchestriert Download → Save → ggf
    PDF-Wrap → Daily-Eintrag → Reply → History → Caption-an-LLM-Routing."""
    doc = update.message.document
    if not doc:
        return
    user_caption = update.message.caption or ""
    log.info(f"document: {doc.file_name}, size={doc.file_size}, caption={user_caption[:60]}")
    await update.message.chat.send_action(constants.ChatAction.TYPING)

    tmp_path = None
    try:
        # 1. Download nach tmp
        file = await doc.get_file()
        suffix = Path(doc.file_name or "upload").suffix
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp_path = tmp.name
        await file.download_to_drive(tmp_path)

        filename = doc.file_name or f"upload-{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # 2. In Vault verschieben (sync, schnell)
        dest, kind, body_preview = await asyncio.to_thread(
            _save_uploaded_doc, tmp_path, filename
        )
        tmp_path = None  # ownership transferred → kein cleanup mehr nötig
        rel = dest.relative_to(VAULT)

        # 3. Format-spezifischer Wrapper für Volltext-Suche
        wrapper_link, md_id, extra_html = None, None, ""
        ext_lower = dest.suffix.lower()
        if ext_lower == ".pdf":
            await update.message.chat.send_action(constants.ChatAction.TYPING)
            md_path, md_id, pdf_text, total_pages = await asyncio.to_thread(
                _create_pdf_wrapper, dest, filename
            )
            wrapper_link = md_path.relative_to(VAULT)
            body_preview = pdf_text[:1500] if pdf_text else ""
            extra_html = f"\n📑 Wrapper: <code>{wrapper_link}</code> ({total_pages} S., id <code>{md_id}</code>)"
        elif ext_lower == ".docx":
            await update.message.chat.send_action(constants.ChatAction.TYPING)
            md_path, md_id, docx_text, para_count = await asyncio.to_thread(
                _create_docx_wrapper, dest, filename
            )
            wrapper_link = md_path.relative_to(VAULT)
            body_preview = docx_text[:1500] if docx_text else ""
            extra_html = f"\n📝 Wrapper: <code>{wrapper_link}</code> ({para_count} Absätze, id <code>{md_id}</code>)"

        # 4. Daily-Eintrag (sync, schnell)
        _record_upload_in_daily(rel, wrapper_link, md_id, user_caption, kind=kind)

        # 5. Antwort an User (kompakter falls Caption — LLM antwortet danach detailliert)
        reply = f"📄 <b>{kind}</b> gespeichert: <code>{rel}</code>"
        if extra_html:
            reply += extra_html
        if body_preview and not user_caption:
            reply += f"\n\n<b>Inhalt (Vorschau):</b>\n<pre><code>{_esc_html(body_preview[:600])}</code></pre>"
        await update.message.reply_text(reply, parse_mode=constants.ParseMode.HTML)

        # 6. Upload-Event in LLM-History (für spätere "lege als Projekt an"-Anfragen)
        try:
            history_msg = _build_upload_event_msg(filename, kind, rel, wrapper_link, md_id, body_preview)
            await update_history(update.effective_user.id, [
                {"role": "user", "content": history_msg}
            ])
        except Exception as e:
            log.warning(f"Upload-Event nicht in History: {e}")

        # 7. Caption als LLM-Anweisung weiterreichen ("Lege als Projekt an" etc.)
        if user_caption.strip():
            await update.message.chat.send_action(constants.ChatAction.TYPING)
            try:
                llm_reply = await llm_loop(user_caption.strip(), update.effective_user.id)
                await safe_reply(update, llm_reply)
            except Exception as e:
                log.exception("Caption-LLM-Routing failed")
                await update.message.reply_text(f"Caption-Verarbeitung fehlgeschlagen: {e}")

    except Exception as e:
        log.exception("document handler failed")
        await update.message.reply_text(f"Document-Fehler: {e}")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


# ============================================================================
# Vault Health-Check (nightly 02:00)
# ============================================================================
# Dreischichtig:
#   1. CHECKS         — read-only, immer (10 Items, in collect_health_data)
#   2. AUTO-FIX       — silent ohne Approval (10 Items, ab Tag 4)
#   3. PROPOSALS      — User-Approval (3 Items: Inbox-Triage, Broken-Links,
#                       Recurring-stuck) — analog zu nightly_suggestion_job

HEALTH_REPORTS_DIR = VAULT / "06_Meta" / "health-reports"
TAG_ALIASES_FILE = VAULT / "06_Meta" / "tag-aliases.json"
PENDING_HEALTH_ACTIONS_FILE = VAULT / "06_Meta" / "pending-health-actions.json"

HEALTH_FIRST_RUN_GRACE_DAYS = 3       # 3 Tage Report-only, ab Tag 4 Auto-Fix
HEALTH_DONE_TASK_AGE_DAYS = 30        # Done-Tasks >30d → archivieren
HEALTH_STALE_INBOX_DAYS = 7           # Inbox-Files >7d → Proposal
HEALTH_RECURRING_STUCK_DAYS = 14      # Recurring nicht aktiviert >14d → Proposal
HEALTH_TAG_TYPO_LEVENSHTEIN = 1       # Schwelle für Tippfehler-Cluster
HEALTH_TAG_MAJORITY_RATIO = 5         # 5:1-Verhältnis für Stufe-C-Konsolidierung

# Auto-Fix-Codes für Reports + Approval-Identifizierung
FIX_TAGS_EMPTY = "tags-empty"
FIX_UPDATED_DATE = "updated-date"
FIX_MISSING_ID = "missing-id"
FIX_KEBAB_ID = "kebab-id"
FIX_DATE_FORMAT = "date-iso"
FIX_FRONTMATTER_ORDER = "fm-order"
FIX_EMPTY_DAILY = "empty-daily"
FIX_AUTO_LINK = "auto-link"
FIX_DONE_ARCHIVE = "done-archive"
FIX_TAG_CONSOLIDATE = "tag-consolidate"


def _levenshtein(a: str, b: str) -> int:
    """Standard Levenshtein-Distanz (case-insensitive Vergleich erfordert pre-lower)."""
    if not a:
        return len(b)
    if not b:
        return len(a)
    m, n = len(a) + 1, len(b) + 1
    cur = list(range(n))
    for i in range(1, m):
        prev, cur = cur, [i] + [0] * (n - 1)
        for j in range(1, n):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            cur[j] = min(prev[j] + 1, cur[j - 1] + 1, prev[j - 1] + cost)
    return cur[-1]


def _is_health_first_run() -> bool:
    """True während der ersten N Tage nach erstem Lauf — Auto-Fix dann disabled.

    Verhindert "alles auf einmal verschwindet"-Schock beim ersten Health-Lauf.
    """
    if not HEALTH_REPORTS_DIR.exists():
        return True
    reports = sorted(HEALTH_REPORTS_DIR.glob("*.md"))
    return len(reports) < HEALTH_FIRST_RUN_GRACE_DAYS


def _load_tag_aliases() -> dict:
    """Liest 06_Meta/tag-aliases.json (Format: {"alt": "neu"})."""
    if not TAG_ALIASES_FILE.exists():
        return {}
    try:
        return json.loads(TAG_ALIASES_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning(f"tag-aliases.json korrupt: {e}")
        return {}


def collect_health_data() -> dict:
    """Single Source: alle Vault-Health-Probleme erkennen (read-only).

    Walks alle .md-Files EINMAL und sammelt parallel alle 10 Check-Ergebnisse.
    Returns dict mit Issue-Listen je Kategorie.
    """
    today = datetime.now(TIMEZONE).date()
    today_str = today.strftime("%Y-%m-%d")

    # Sammelbuckets
    lint_issues = []           # (rel_path, problem)
    broken_links = []          # (rel_path, target)
    orphans = []               # rel_path (Wiki-Files ohne Backlinks)
    duplicate_ids = []         # (id, [rel_paths])
    daily_gaps = []            # ISO-Date-Strings (Werktage ohne Daily)
    stale_inbox = []           # (rel_path, age_days)
    stale_uploads = []         # rel_path (uploads ohne Wrapper-Bezug)
    done_old_tasks = []        # task-Dicts (status=done, >30d, nicht-recurring)
    recurring_stuck = []       # task-Dicts (recurring, last_completed > 2× Periode her)
    tag_clusters = []          # (canonical, [variants_with_counts])

    # Walk: alle non-noise .md mit Frontmatter
    all_notes = []  # [(path, post)]
    id_to_paths: dict = {}  # für duplicate-detection
    backlinks_count: dict = {}  # id → count
    all_tags: dict = {}  # tag → count

    for path, post in iter_vault_md():
        all_notes.append((path, post))
        meta = post.metadata
        rel = path.relative_to(VAULT).as_posix()

        # ── Check 1: Schema-Lint (Pflichtfelder, Type-Folder-Konformität) ──
        if not meta:
            lint_issues.append((rel, "kein Frontmatter"))
            continue
        t = meta.get("type")
        if not t:
            lint_issues.append((rel, "type fehlt"))
            continue

        # ── Check 7: Doppel-IDs sammeln ──
        fid = meta.get("id")
        if fid and isinstance(fid, str):
            id_to_paths.setdefault(fid, []).append(rel)

        # ── Check 6: Tag-Counter füllen ──
        tags = meta.get("tags") or []
        if isinstance(tags, list):
            for tg in tags:
                if isinstance(tg, str) and tg.strip():
                    all_tags[tg.strip().lower()] = all_tags.get(tg.strip().lower(), 0) + 1

        # ── Wikilinks für Broken-Link + Orphan-Check ──
        body = post.content or ""
        # Code-Spans + already-wikilinks rausfiltern
        body_clean = re.sub(r"```.*?```", "", body, flags=re.DOTALL)
        body_clean = re.sub(r"`[^`\n]+`", "", body_clean)
        for m in re.finditer(r"\[\[([^\]|#]+?)(?:\|[^\]]*)?\]\]", body_clean):
            target = m.group(1).strip()
            backlinks_count[target] = backlinks_count.get(target, 0) + 1

    # ── Check 2: Broken Wikilinks (Pass 2 nach komplettem Walk) ──
    # Wie Obsidian: Wikilink ist auflösbar wenn target == id ODER == filename-stem.
    # Vorher nur id-Match → User-Files mit `[[Panama – Budget]]` (kein id-Frontmatter,
    # aber File `Panama – Budget.md` existiert) wurden fälschlich als broken gemeldet.
    all_ids = {p.metadata.get("id") for _, p in all_notes if p.metadata.get("id")}
    all_stems = {p.stem for p, _ in all_notes}
    # Plus: System-File-Stems (CLAUDE/SCHEMA/PIPELINES/COMMANDS/MOC/README) —
    # die werden in iter_vault_md geskippt, aber valide Wikilink-Targets:
    # ein User-File darf `[[SCHEMA]]` schreiben ohne dass das als broken gilt.
    system_stems = set()
    for sysfile in _VAULT_ROOT_SYSTEM_DOCS:
        if (VAULT / sysfile).exists():
            system_stems.add(Path(sysfile).stem)
    all_known = all_ids | all_stems | system_stems
    # Plus: Wikilinks mit nested [[ am Ende des targets sind syntaktisch kaputt
    # (z.B. "[[t-foo-[[bar]]" statt "[[t-foo|bar]]"). Werden trotzdem gemeldet
    # weil das ein echter Bug im Source-File ist.
    for path, post in all_notes:
        body = post.content or ""
        body_clean = re.sub(r"```.*?```", "", body, flags=re.DOTALL)
        body_clean = re.sub(r"`[^`\n]+`", "", body_clean)
        rel = path.relative_to(VAULT).as_posix()
        for m in re.finditer(r"\[\[([^\]|#]+?)(?:\|[^\]]*)?\]\]", body_clean):
            target = m.group(1).strip()
            if target not in all_known:
                broken_links.append((rel, target))

    # ── Check 3: Orphans (Wiki-Files ohne incoming Backlinks) ──
    for path, post in all_notes:
        meta = post.metadata
        t = meta.get("type", "")
        # Nur Wiki-Typen: andere haben legitim oft keine Backlinks
        if t not in ("concept", "topic", "person", "organization", "tool", "method", "glossary"):
            continue
        fid = meta.get("id")
        if fid and backlinks_count.get(fid, 0) == 0:
            orphans.append(path.relative_to(VAULT).as_posix())

    # ── Check 7: Duplicate IDs ──
    for fid, paths in id_to_paths.items():
        if len(paths) > 1:
            duplicate_ids.append((fid, paths))

    # ── Check 4: Daily-Gaps (Werktage der letzten 7 Tage ohne Daily) ──
    for delta in range(1, 8):  # gestern bis vor 7 Tagen
        d = today - timedelta(days=delta)
        if d.weekday() >= 5:  # Samstag/Sonntag → kein Werktag, kein Gap
            continue
        daily_path = DAILY_DIR / f"{d.isoformat()}.md"
        if not daily_path.exists():
            daily_gaps.append(d.isoformat())

    # ── Check 5: Stale Inbox ──
    inbox_dir = VAULT / "00_Inbox"
    if inbox_dir.exists():
        for f in inbox_dir.iterdir():
            if not f.is_file() or f.name.startswith("."):
                continue
            try:
                mtime = datetime.fromtimestamp(f.stat().st_mtime, TIMEZONE).date()
                age_days = (today - mtime).days
                if age_days > HEALTH_STALE_INBOX_DAYS:
                    stale_inbox.append((f.relative_to(VAULT).as_posix(), age_days))
            except Exception:
                continue

    # ── Check: Stale Uploads (PDFs ohne .md-Wrapper-Sibling) ──
    uploads_dir = VAULT / "01_Raw" / "uploads"
    if uploads_dir.exists():
        for f in uploads_dir.glob("*.pdf"):
            md_sibling = f.with_suffix(".md")
            if not md_sibling.exists():
                stale_uploads.append(f.relative_to(VAULT).as_posix())

    # ── Check 8: Done-Tasks >30d (nicht recurring) ──
    cutoff_done = today - timedelta(days=HEALTH_DONE_TASK_AGE_DAYS)
    for path, post in all_notes:
        meta = post.metadata
        if meta.get("type") != "task":
            continue
        if meta.get("status") != "done":
            continue
        if meta.get("recurrence"):  # recurring NIE auto-archivieren
            continue
        # Nutze last_completed wenn vorhanden, sonst updated
        ref_date_raw = meta.get("last_completed") or meta.get("updated")
        ref_date = _due_to_date(ref_date_raw)
        if ref_date is None or ref_date > cutoff_done:
            continue
        done_old_tasks.append({
            "path": path.relative_to(VAULT).as_posix(),
            "id": meta.get("id", path.stem),
            "title": meta.get("title", path.stem),
            "ref_date": ref_date.isoformat(),
            "age_days": (today - ref_date).days,
        })

    # ── Check 9: Recurring stuck (Task hat recurrence aber wurde lange nicht aktiviert) ──
    for path, post in all_notes:
        meta = post.metadata
        if meta.get("type") != "task" or not meta.get("recurrence"):
            continue
        if meta.get("status") != "done":
            continue  # nur done-Tasks die auf Reset warten
        last = _due_to_date(meta.get("last_completed"))
        if last is None:
            continue  # noch nie done — unkritisch
        days_since = (today - last).days
        # Schwelle: 2× Pattern-Periode oder 14d minimum
        rec = meta.get("recurrence")
        threshold = {
            "daily": 3,         # >3d nach done = stuck
            "weekdays": 4,
            "weekly": 14,       # >2 Wochen
            "monthly": 60,      # >2 Monate
        }.get(rec, HEALTH_RECURRING_STUCK_DAYS)
        if days_since > threshold:
            recurring_stuck.append({
                "path": path.relative_to(VAULT).as_posix(),
                "id": meta.get("id", path.stem),
                "title": meta.get("title", path.stem),
                "recurrence": rec,
                "last_completed": last.isoformat(),
                "days_since": days_since,
            })

    # ── Check 10: Tag-Drift (Cluster-Vorschläge) ──
    aliases = _load_tag_aliases()
    # Alias-Map auflösen: wenn 'work' in Aliases → wird zu aliases['work']
    sorted_tags = sorted(all_tags.items(), key=lambda x: -x[1])  # häufigste zuerst
    seen = set()
    for tag, count in sorted_tags:
        if tag in seen:
            continue
        cluster = [(tag, count)]
        for other_tag, other_count in sorted_tags:
            if other_tag == tag or other_tag in seen:
                continue
            # Stufe A: User-Alias
            if aliases.get(other_tag) == tag:
                cluster.append((other_tag, other_count))
                continue
            # Stufe B: Tippfehler (Levenshtein ≤ 1, kleinerer ist seltener)
            if (_levenshtein(tag, other_tag) <= HEALTH_TAG_TYPO_LEVENSHTEIN
                    and other_count <= 2 and count >= 3):
                cluster.append((other_tag, other_count))
                continue
            # Stufe C: 5:1-Mehrheit (Plural/Singular oder DE/EN-Varianten)
            if other_count >= 1 and count >= other_count * HEALTH_TAG_MAJORITY_RATIO:
                # Heuristik: gemeinsame 4+ Anfangs-Buchstaben (Plural-Singular-Detektor)
                # ODER explizites Alias
                shared_prefix = sum(1 for x, y in zip(tag, other_tag) if x == y)
                if shared_prefix >= 4 or aliases.get(other_tag) == tag:
                    cluster.append((other_tag, other_count))
        if len(cluster) > 1:
            tag_clusters.append((tag, cluster))
            seen.update(t for t, _ in cluster)

    return {
        "today": today,
        "today_str": today_str,
        "first_run": _is_health_first_run(),
        "lint_issues": lint_issues,
        "broken_links": broken_links,
        "orphans": orphans,
        "duplicate_ids": duplicate_ids,
        "daily_gaps": daily_gaps,
        "stale_inbox": stale_inbox,
        "stale_uploads": stale_uploads,
        "done_old_tasks": done_old_tasks,
        "recurring_stuck": recurring_stuck,
        "tag_clusters": tag_clusters,
        "total_notes": len(all_notes),
        "all_tags": all_tags,
    }


# ─── Auto-Fix-Funktionen ────────────────────────────────────────────────────
# Jede gibt zurück: list[(fix_code, rel_path, beschreibung)]
# Bei dry_run=True wird nichts geschrieben, nur die Liste zurückgegeben.

# Welche Types haben tags als Pflichtfeld (siehe SCHEMA.md)
REQUIRED_TYPES_WITH_TAGS = {
    "concept", "topic", "person", "organization", "tool", "method",
    "glossary", "timeline", "article", "paper", "repo", "video", "book",
    "dataset", "summary", "report", "slides", "query", "project",
    "daily", "task", "note", "meeting", "area",
}


def _autofix_frontmatter_hygiene(dry_run: bool) -> list:
    """Sammelfix für Frontmatter-Kleinkram: tags=[] ergänzen, updated nachtragen,
    id aus Filename ableiten, kebab-case, ISO-Date. Pro File ein Save."""
    fixes = []
    today_str = datetime.now(TIMEZONE).date().isoformat()
    for path, post in iter_vault_md():
        meta = post.metadata
        if not meta:
            continue
        rel = path.relative_to(VAULT).as_posix()
        changed = False

        # Fix 1: tags=[] wenn Pflichtfeld fehlt
        if "tags" not in meta and meta.get("type") in REQUIRED_TYPES_WITH_TAGS:
            if not dry_run:
                post["tags"] = []
            fixes.append((FIX_TAGS_EMPTY, rel, "tags: [] ergänzt"))
            changed = True

        # Fix 2: id aus Filename ableiten
        if "id" not in meta or not meta.get("id"):
            derived = path.stem
            # Bei Daily-Notes: id = "daily-YYYY-MM-DD"
            if path.parent == DAILY_DIR:
                derived = f"daily-{path.stem}"
            if not dry_run:
                post["id"] = derived
            fixes.append((FIX_MISSING_ID, rel, f"id: {derived} (aus Filename)"))
            changed = True

        # Fix 3: ISO-Date-Format normalisieren (created/updated/date/captured)
        for date_field in ("created", "updated", "date", "captured", "started", "due"):
            v = meta.get(date_field)
            if not isinstance(v, str):
                continue
            # Versuche zu normalisieren wenn nicht-ISO
            if not re.match(r"^\d{4}-\d{2}-\d{2}$", v):
                # Häufige Varianten: 2026/04/27, 27.04.2026, 27-04-2026
                fixed = None
                for fmt in ("%Y/%m/%d", "%d.%m.%Y", "%d-%m-%Y", "%Y.%m.%d"):
                    try:
                        fixed = datetime.strptime(v, fmt).date().isoformat()
                        break
                    except ValueError:
                        continue
                if fixed:
                    if not dry_run:
                        post[date_field] = fixed
                    fixes.append((FIX_DATE_FORMAT, rel, f"{date_field}: {v} → {fixed}"))
                    changed = True

        # Fix 4: updated nachtragen wenn fehlt aber created vorhanden
        if "updated" not in meta and meta.get("created"):
            try:
                mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).date().isoformat()
                if not dry_run:
                    post["updated"] = mtime_iso
                fixes.append((FIX_UPDATED_DATE, rel, f"updated: {mtime_iso} (aus mtime)"))
                changed = True
            except Exception:
                pass

        if changed and not dry_run:
            try:
                atomic_write(path, frontmatter.dumps(post) + "\n")
            except Exception as e:
                log.warning(f"autofix write failed for {rel}: {e}")
    return fixes


def _autofix_empty_dailies(dry_run: bool) -> list:
    """Löscht Daily-Notes die nur Template-Platzhalter enthalten + älter als 7d sind."""
    fixes = []
    today = datetime.now(TIMEZONE).date()
    cutoff = today - timedelta(days=7)
    if not DAILY_DIR.exists():
        return fixes
    placeholder_lines = {
        "- [ ]", "- [x]", "-", "•",
        "- Was lief gut?", "- Was nehme ich mit?",
    }
    for path in DAILY_DIR.glob("*.md"):
        try:
            d = datetime.strptime(path.stem, "%Y-%m-%d").date()
        except ValueError:
            continue
        if d > cutoff:
            continue  # zu jung, nicht löschen
        try:
            post = frontmatter.load(path)
            body = post.content or ""
        except Exception:
            continue
        # Body ohne Headings + Whitespace prüfen
        non_meta_lines = [
            l.strip() for l in body.splitlines()
            if l.strip() and not l.strip().startswith("#")
        ]
        meaningful = [l for l in non_meta_lines if l not in placeholder_lines]
        if not meaningful:
            rel = path.relative_to(VAULT).as_posix()
            if not dry_run:
                try:
                    path.unlink()
                except Exception as e:
                    log.warning(f"empty-daily delete failed: {e}")
                    continue
            fixes.append((FIX_EMPTY_DAILY, rel, f"leere Daily ({d.isoformat()}) gelöscht"))
    return fixes


def _autofix_auto_link_existing(dry_run: bool) -> list:
    """Geht durch 10_Life/notes/ + 10_Life/daily/ und linkt Klartext-Erwähnungen
    zu mittlerweile existierenden IDs nach. Nur in 10_Life/ — Wiki + Projekte tabu.
    """
    fixes = []
    pmap = _get_link_index()
    if not pmap:
        return fixes
    for target_dir in (DAILY_DIR, NOTES_DIR):
        if not target_dir.exists():
            continue
        for path, post in iter_vault_md(target_dir, recursive=False, skip_noise=False):
            body = post.content or ""
            # Eigene ID excluden um Self-Links zu vermeiden
            self_id = post.metadata.get("id")
            exclude = {self_id} if self_id else set()
            new_body = auto_link(body, exclude_ids=exclude)
            if new_body != body:
                rel = path.relative_to(VAULT).as_posix()
                # Differenz zählen für Reporting
                added_links = new_body.count("[[") - body.count("[[")
                if not dry_run:
                    post.content = new_body
                    try:
                        atomic_write(path, frontmatter.dumps(post) + "\n")
                    except Exception as e:
                        log.warning(f"auto-link write failed: {e}")
                        continue
                fixes.append((FIX_AUTO_LINK, rel, f"{added_links} Wikilink(s) nachgetragen"))
    return fixes


def _autofix_archive_done_tasks(done_old_tasks: list, dry_run: bool) -> list:
    """Verschiebt Done-Tasks (>30d, nicht-recurring) nach 99_Archive/10_Life/tasks/."""
    fixes = []
    if not done_old_tasks:
        return fixes
    archive_dir = VAULT / "99_Archive" / "10_Life" / "tasks"
    today_str = datetime.now(TIMEZONE).date().isoformat()
    if not dry_run:
        archive_dir.mkdir(parents=True, exist_ok=True)
    for t in done_old_tasks:
        src = VAULT / t["path"]
        if not src.exists():
            continue
        dst = archive_dir / src.name
        n = 1
        while dst.exists():
            dst = archive_dir / f"{src.stem}-{n}{src.suffix}"
            n += 1
        if not dry_run:
            try:
                # Atomar: lade Frontmatter, schreibe mit archived-Datum DIREKT
                # zum dst-Pfad via atomic_write (tmp+rename), dann src löschen.
                # Vorher: src.write_text + shutil.move = 2 Operationen, Crash
                # dazwischen hinterließ src mit archived-Feld am Quellort.
                post = frontmatter.load(src)
                post["archived"] = today_str
                atomic_write(dst, frontmatter.dumps(post) + "\n")
                src.unlink()
            except Exception as e:
                log.warning(f"archive-done failed for {src.name}: {e}")
                # Cleanup: wenn dst schon geschrieben aber src noch da, dst entfernen
                # damit nächster Lauf nicht denkt der Task wäre schon archiviert
                if dst.exists() and src.exists():
                    try:
                        dst.unlink()
                    except Exception:
                        pass
                continue
        fixes.append((FIX_DONE_ARCHIVE, t["path"],
                      f"archiviert ({t['age_days']}d alt): {t['title']}"))
    if fixes and not dry_run:
        invalidate_link_index()
    return fixes


def _autofix_consolidate_tags(tag_clusters: list, dry_run: bool) -> list:
    """Vereinheitlicht Tag-Cluster im ganzen Vault.
    cluster: (canonical, [(variant, count), ...]) — variants werden zu canonical."""
    fixes = []
    if not tag_clusters:
        return fixes
    # Build replace-map: variant_lower → canonical
    replace_map: dict = {}
    cluster_summaries = []
    for canonical, variants in tag_clusters:
        for variant, _count in variants:
            if variant != canonical:
                replace_map[variant.lower()] = canonical
        cluster_summaries.append((canonical, [v for v, _ in variants if v != canonical]))
    if not replace_map:
        return fixes

    file_changes = 0
    for path, post in iter_vault_md():
        tags = post.metadata.get("tags") or []
        if not isinstance(tags, list) or not tags:
            continue
        new_tags = []
        seen_in_file = set()
        any_changed = False
        for tg in tags:
            if not isinstance(tg, str):
                new_tags.append(tg)
                continue
            tl = tg.strip().lower()
            target = replace_map.get(tl, tg)
            if target.lower() != tl:
                any_changed = True
            if target.lower() not in seen_in_file:
                new_tags.append(target)
                seen_in_file.add(target.lower())
        if any_changed:
            file_changes += 1
            if not dry_run:
                post["tags"] = new_tags
                try:
                    atomic_write(path, frontmatter.dumps(post) + "\n")
                except Exception as e:
                    log.warning(f"tag-consolidate write failed: {e}")
                    continue

    for canonical, variants in cluster_summaries:
        if variants:
            fixes.append((FIX_TAG_CONSOLIDATE, "*",
                          f"{', '.join(variants)} → {canonical} (in {file_changes} Files)"))
    return fixes


def run_health_autofixes(data: dict, dry_run: bool = False) -> list:
    """Führt alle 10 Auto-Fixes aus, gibt kombinierte Liste zurück.

    Wird vom nightly Job nur gerufen wenn NICHT first_run.
    """
    all_fixes = []
    all_fixes.extend(_autofix_frontmatter_hygiene(dry_run))
    all_fixes.extend(_autofix_empty_dailies(dry_run))
    all_fixes.extend(_autofix_auto_link_existing(dry_run))
    all_fixes.extend(_autofix_archive_done_tasks(data["done_old_tasks"], dry_run))
    all_fixes.extend(_autofix_consolidate_tags(data["tag_clusters"], dry_run))
    return all_fixes


# ─── Approval-Proposals (3 Items, brauchen User-Input) ──────────────────────

def _build_health_proposals(data: dict) -> list:
    """Baut Approval-Proposals aus den Check-Daten.

    Returns list von Dicts: {id, type, summary, items, action_options}
    """
    proposals = []

    # Proposal 1: Stale Inbox
    if data["stale_inbox"]:
        items = data["stale_inbox"][:5]  # max 5 zeigen
        proposals.append({
            "id": "p-inbox",
            "type": "stale_inbox",
            "summary": f"{len(data['stale_inbox'])} Files seit >7d in 00_Inbox",
            "items": [{"path": p, "age_days": d} for p, d in items],
            "options": ["zu Notes", "zu Tasks", "Archiv", "skip"],
        })

    # Proposal 2: Broken Wikilinks
    if data["broken_links"]:
        items = data["broken_links"][:8]
        proposals.append({
            "id": "p-broken",
            "type": "broken_links",
            "summary": f"{len(data['broken_links'])} Wikilink(s) ohne Ziel",
            "items": [{"path": p, "target": t} for p, t in items],
            "options": ["zeig Liste", "zu Klartext", "skip"],
        })

    # Proposal 3: Recurring stuck
    if data["recurring_stuck"]:
        proposals.append({
            "id": "p-stuck",
            "type": "recurring_stuck",
            "summary": f"{len(data['recurring_stuck'])} Recurring Task(s) hängen >Periode",
            "items": data["recurring_stuck"][:5],
            "options": ["jetzt aktivieren", "recurrence entfernen", "skip"],
        })

    return proposals


def _save_pending_health_actions(proposals: list) -> None:
    """Persistiert Proposals in 06_Meta/pending-health-actions.json."""
    if not proposals:
        if PENDING_HEALTH_ACTIONS_FILE.exists():
            try:
                PENDING_HEALTH_ACTIONS_FILE.unlink()
            except Exception:
                pass
        return
    PENDING_HEALTH_ACTIONS_FILE.parent.mkdir(parents=True, exist_ok=True)
    PENDING_HEALTH_ACTIONS_FILE.write_text(
        json.dumps({"created": datetime.now(TIMEZONE).isoformat(), "proposals": proposals},
                   indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def _load_pending_health_actions() -> list:
    if not PENDING_HEALTH_ACTIONS_FILE.exists():
        return []
    try:
        data = json.loads(PENDING_HEALTH_ACTIONS_FILE.read_text(encoding="utf-8"))
        return data.get("proposals", [])
    except Exception:
        return []


# ─── Health-Report-Writer ───────────────────────────────────────────────────

def write_health_report(data: dict, autofixes: list, proposals: list) -> Path:
    """Schreibt Detail-Report nach 06_Meta/health-reports/YYYY-MM-DD.md."""
    HEALTH_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    today_str = data["today_str"]
    report_path = HEALTH_REPORTS_DIR / f"{today_str}.md"

    lines = [
        "---",
        f"title: Vault-Health {today_str}",
        "type: meta",
        f"updated: {today_str}",
        "maintained_by: nightly_health_job",
        "---",
        "",
        f"# 🔧 Vault-Health-Check — {today_str}",
        "",
        f"**Notes gescannt:** {data['total_notes']}  ",
        f"**Tags total:** {len(data['all_tags'])}  ",
        f"**First-Run-Mode:** {'JA (Auto-Fix disabled)' if data['first_run'] else 'nein'}",
        "",
    ]

    # ── Auto-Fix-Sektion ──
    if autofixes:
        by_code: dict = {}
        for code, path, desc in autofixes:
            by_code.setdefault(code, []).append((path, desc))
        lines.append(f"## ✅ Auto-Fixed ({len(autofixes)})")
        lines.append("")
        for code, items in sorted(by_code.items()):
            lines.append(f"### {code} ({len(items)})")
            for path, desc in items[:20]:
                lines.append(f"- `{path}`: {desc}")
            if len(items) > 20:
                lines.append(f"- _… {len(items)-20} weitere_")
            lines.append("")
    elif not data["first_run"]:
        lines.append("## ✅ Auto-Fixed (0)")
        lines.append("")
        lines.append("Nichts zu fixen — Vault ist sauber.")
        lines.append("")

    # ── Issues (read-only) ──
    issues_blocks = []
    if data["lint_issues"]:
        b = [f"### Lint-Issues ({len(data['lint_issues'])})"]
        for path, prob in data["lint_issues"][:20]:
            b.append(f"- `{path}`: {prob}")
        if len(data["lint_issues"]) > 20:
            b.append(f"- _… {len(data['lint_issues'])-20} weitere_")
        issues_blocks.append("\n".join(b))
    if data["broken_links"]:
        b = [f"### Broken Wikilinks ({len(data['broken_links'])})"]
        for path, tgt in data["broken_links"][:10]:
            b.append(f"- `{path}` → `[[{tgt}]]`")
        issues_blocks.append("\n".join(b))
    if data["orphans"]:
        b = [f"### Orphans — Wiki-Files ohne Backlinks ({len(data['orphans'])})"]
        for p in data["orphans"][:10]:
            b.append(f"- `{p}`")
        issues_blocks.append("\n".join(b))
    if data["duplicate_ids"]:
        b = [f"### Doppelte IDs ({len(data['duplicate_ids'])})"]
        for fid, paths in data["duplicate_ids"]:
            b.append(f"- `{fid}` → {', '.join(f'`{p}`' for p in paths)}")
        issues_blocks.append("\n".join(b))
    if data["daily_gaps"]:
        b = [f"### Daily-Lücken (Werktage ohne Daily-Note, letzte 7 Tage)"]
        b.append(", ".join(data["daily_gaps"]))
        issues_blocks.append("\n".join(b))
    if data["stale_uploads"]:
        b = [f"### Stale Uploads — PDFs ohne .md-Wrapper ({len(data['stale_uploads'])})"]
        for p in data["stale_uploads"][:10]:
            b.append(f"- `{p}`")
        issues_blocks.append("\n".join(b))

    if issues_blocks:
        lines.append("## ⚠️ Issues (read-only)")
        lines.append("")
        for block in issues_blocks:
            lines.append(block)
            lines.append("")

    # ── Pending Approvals ──
    if proposals:
        lines.append(f"## 🔧 Pending Approvals ({len(proposals)})")
        lines.append("")
        for i, p in enumerate(proposals, 1):
            lines.append(f"### {i}. {p['summary']}")
            for item in p["items"][:3]:
                if isinstance(item, dict):
                    lines.append(f"- {item}")
            lines.append(f"  Optionen: {', '.join(p['options'])}")
            lines.append("")

    # ── Tag-Übersicht (immer am Ende, kompakt) ──
    if data["all_tags"]:
        top_tags = sorted(data["all_tags"].items(), key=lambda x: -x[1])[:15]
        lines.append("## 🏷️ Top-Tags")
        lines.append("")
        lines.append(", ".join(f"`{t}` ({c})" for t, c in top_tags))
        lines.append("")

    report_path.write_text("\n".join(lines), encoding="utf-8")
    return report_path


def cleanup_old_health_reports() -> int:
    """Löscht Health-Reports älter als 30 Tage. Returns Anzahl gelöscht."""
    if not HEALTH_REPORTS_DIR.exists():
        return 0
    cutoff = datetime.now(TIMEZONE).date() - timedelta(days=30)
    deleted = 0
    for path in HEALTH_REPORTS_DIR.glob("*.md"):
        try:
            d = datetime.strptime(path.stem, "%Y-%m-%d").date()
            if d < cutoff:
                path.unlink()
                deleted += 1
        except (ValueError, OSError):
            continue
    return deleted


# ─── Orchestrator + JobQueue-Callback ──────────────────────────────────────

def run_nightly_health() -> dict:
    """Synchroner Orchestrator — läuft im to_thread aus dem JobQueue-Callback.

    Returns dict mit Stats für Briefing-Integration.
    """
    if not VAULT.exists():
        log.error(f"health: VAULT existiert nicht: {VAULT}")
        return {"error": "vault_missing"}

    data = collect_health_data()
    autofixes = []
    if not data["first_run"]:
        try:
            autofixes = run_health_autofixes(data, dry_run=False)
        except Exception as e:
            log.exception(f"health autofix failed: {e}")
    proposals = _build_health_proposals(data)
    _save_pending_health_actions(proposals)
    report_path = write_health_report(data, autofixes, proposals)
    cleaned = cleanup_old_health_reports()
    log.info(
        f"health: report → {report_path.relative_to(VAULT)} | "
        f"autofixes={len(autofixes)}, proposals={len(proposals)}, "
        f"first_run={data['first_run']}, old-reports-cleaned={cleaned}"
    )
    return {
        "report_path": str(report_path.relative_to(VAULT)),
        "autofix_count": len(autofixes),
        "proposal_count": len(proposals),
        "issue_summary": {
            "lint": len(data["lint_issues"]),
            "broken_links": len(data["broken_links"]),
            "orphans": len(data["orphans"]),
            "duplicate_ids": len(data["duplicate_ids"]),
            "daily_gaps": len(data["daily_gaps"]),
            "stale_uploads": len(data["stale_uploads"]),
        },
        "first_run": data["first_run"],
    }


async def nightly_health_job(ctx: ContextTypes.DEFAULT_TYPE):
    """JobQueue-Callback — läuft täglich 02:00."""
    try:
        await asyncio.to_thread(run_nightly_health)
    except Exception as e:
        log.exception(f"nightly_health_job failed: {e}")


# ─── Tool: apply_health_action ──────────────────────────────────────────────

def apply_health_action(action: str) -> str:
    """Verarbeitet User-Antwort auf pending Health-Proposals.

    Action-Formate (wie bei apply_memory_suggestion):
      "1 2"       — proposals 1 und 2 mit Default-Action akzeptieren
      "1,2"       — wie oben
      "alle"/"ja" — alle akzeptieren
      "0"/"nein"  — alles skippen + pending leeren

    Default-Action pro Proposal-Typ:
      stale_inbox    → Files nach 99_Archive/
      broken_links   → Wikilink-Klammern entfernen (zu Klartext)
      recurring_stuck→ Task jetzt aktivieren (status=open + last_completed=heute)
    """
    proposals = _load_pending_health_actions()
    if not proposals:
        return "Keine pending Health-Aktionen."

    action = (action or "").strip().lower()
    if action in ("0", "nein", "skip", "no"):
        try:
            PENDING_HEALTH_ACTIONS_FILE.unlink()
        except Exception:
            pass
        return f"⊘ {len(proposals)} Health-Aktionen verworfen."

    if action in ("alle", "ja", "all", "yes"):
        indices = list(range(1, len(proposals) + 1))
    else:
        # Zahlen extrahieren
        nums = [int(x) for x in re.findall(r"\d+", action) if int(x) > 0]
        if not nums:
            return f"Konnte keine Auswahl erkennen aus '{action}'. Format: '1 2', 'alle', '0'."
        indices = nums

    results = []
    today = datetime.now(TIMEZONE).date()
    today_str = today.strftime("%Y-%m-%d")

    for idx in indices:
        if idx < 1 or idx > len(proposals):
            results.append(f"#{idx}: ungültig (1-{len(proposals)})")
            continue
        p = proposals[idx - 1]

        if p["type"] == "stale_inbox":
            archived = 0
            archive_dir = VAULT / "99_Archive" / "00_Inbox"
            archive_dir.mkdir(parents=True, exist_ok=True)
            for item in p["items"]:
                src = VAULT / item["path"]
                if not src.exists():
                    continue
                dst = archive_dir / src.name
                n = 1
                while dst.exists():
                    dst = archive_dir / f"{src.stem}-{n}{src.suffix}"
                    n += 1
                try:
                    shutil.move(str(src), str(dst))
                    archived += 1
                except Exception as e:
                    log.warning(f"inbox archive failed: {e}")
            results.append(f"#{idx}: ✓ {archived} Inbox-Files archiviert")

        elif p["type"] == "broken_links":
            cleaned_files = 0
            cleaned_links = 0
            for item in p["items"]:
                src = VAULT / item["path"]
                if not src.exists():
                    continue
                target = item["target"]
                try:
                    post = frontmatter.load(src)
                    body = post.content or ""
                    pat = re.compile(
                        r"\[\[" + re.escape(target) + r"(?:\|([^\]]*))?\]\]"
                    )
                    n = 0
                    def _strip(m):
                        nonlocal n
                        n += 1
                        return m.group(1) if m.group(1) else target
                    new_body = pat.sub(_strip, body)
                    if n > 0:
                        post.content = new_body
                        atomic_write(src, frontmatter.dumps(post) + "\n")
                        cleaned_files += 1
                        cleaned_links += n
                except Exception as e:
                    log.warning(f"broken-link cleanup failed: {e}")
            results.append(f"#{idx}: ✓ {cleaned_links} broken Wikilinks zu Klartext ({cleaned_files} Files)")

        elif p["type"] == "recurring_stuck":
            reactivated = 0
            for item in p["items"]:
                src = VAULT / item["path"]
                if not src.exists():
                    continue
                try:
                    post = frontmatter.load(src)
                    post["status"] = "open"
                    post["updated"] = today_str
                    body = (post.content or "").rstrip() + (
                        f"\n- {today_str}: manuell reaktiviert (war stuck recurring)\n"
                    )
                    post.content = body
                    atomic_write(src, frontmatter.dumps(post) + "\n")
                    reactivated += 1
                except Exception as e:
                    log.warning(f"recurring-stuck reactivate failed: {e}")
            results.append(f"#{idx}: ✓ {reactivated} stuck Recurring Tasks reaktiviert")
        else:
            results.append(f"#{idx}: unbekannter Proposal-Typ '{p['type']}'")

    # Akzeptierte aus pending entfernen
    remaining = [p for i, p in enumerate(proposals, 1) if i not in indices]
    if remaining:
        _save_pending_health_actions(remaining)
        tail = f" ({len(remaining)} Aktionen bleiben pending)"
    else:
        try:
            PENDING_HEALTH_ACTIONS_FILE.unlink()
        except Exception:
            pass
        tail = ""

    return "Health-Aktionen verarbeitet:\n" + "\n".join(results) + tail


# apply_health_action ist KEIN LLM-Tool mehr — wird via _detect_pending_reply_intent
# in handle_text direkt gerufen. Funktion selbst bleibt für diesen Aufruf erhalten.


def compute_briefing() -> str:
    """Generiere die morgendliche Zusammenfassung als HTML-String.

    Inhalt: Datum, überfällige Tasks, heute geplant (aus Daily 'Heute'),
    offene Tasks (sortiert nach Prio), gestern Abends-Reflexion.
    Bei fundamentalen Problemen (Vault nicht erreichbar) → klare Fehlermeldung.
    """
    # Vault-Reachability-Check
    if not VAULT.exists() or not VAULT.is_dir():
        return f"❌ <b>Briefing fehlgeschlagen</b>\nVault unter <code>{VAULT}</code> nicht erreichbar.\nMount oder Container-Volume prüfen."

    today = today_iso()
    yesterday = (datetime.now(TIMEZONE).date() - timedelta(days=1)).isoformat()

    # Single Source of Truth — gleiche Daten wie get_today_agenda
    data = collect_today_data()

    parts = [f"☀️ <b>Guten Morgen — {today}</b>"]

    # ─── Today's Daily: was steht für heute geplant? ───
    today_path = DAILY_DIR / f"{today}.md"
    if today_path.exists():
        try:
            post = frontmatter.load(today_path)
            body = post.content or ""
            m = re.search(r"## Heute\s*\n(.*?)(?=\n## |\Z)", body, re.DOTALL)
            if m and m.group(1).strip() and m.group(1).strip() != "- [ ]":
                heute_text = m.group(1).strip()[:600]
                parts.append(f"\n📋 <b>Heute geplant</b>\n<pre>{_esc_html(heute_text)}</pre>")
        except Exception as e:
            log.warning(f"briefing: today daily parse failed: {e}")

    # ─── Reminders heute ───
    if data["reminders"]:
        parts.append("\n⏰ <b>Erinnerungen heute</b>")
        for fire_at, r in data["reminders"][:6]:
            rec = " 🔁" if r.get("recurrence") else ""
            parts.append(f"• {fire_at.strftime('%H:%M')} — {_esc_html(r['message'][:80])}{rec}")

    # ─── Meetings heute ───
    if data["meetings"]:
        parts.append("\n🤝 <b>Meetings heute</b>")
        for m in data["meetings"][:5]:
            parts.append(f"• {_esc_html(str(m['title']))}")

    # ─── Überfällig + Heute fällig (kombiniert für Briefing-Übersicht) ───
    today_date = data["today"]
    if data["overdue_tasks"]:
        parts.append("\n⚠️ <b>Überfällig</b>")
        for t in data["overdue_tasks"][:8]:
            # Konsistente Date-Normalisierung — vorher konnten date-Objekte aus
            # YAML-Frontmatter als rohe Repr im Telegram landen.
            due_d = _due_to_date(t.get("due"))
            if due_d is not None:
                delta = (today_date - due_d).days
                due_disp = f"vor {delta}d" if delta > 0 else due_d.isoformat()
            else:
                due_disp = "—"
            parts.append(f"• {_esc_html(str(t['title']))} <i>({due_disp})</i>")

    # Andere offene Tasks (heute + nodate-high-prio + Rest, dedupliziert)
    seen_ids = {t["id"] for t in data["overdue_tasks"]}
    other_open = (
        [t for t in data["today_tasks"] if t["id"] not in seen_ids]
        + [t for t in data["high_nodate_tasks"] if t["id"] not in seen_ids]
    )
    # Resterest aus allen offenen die noch nicht in einem Bucket sind
    seen_ids.update(t["id"] for t in other_open)
    extra = [t for t in _read_open_tasks() if t["id"] not in seen_ids]
    other_open.extend(_sort_tasks_by_prio_due(extra))

    if other_open:
        parts.append("\n✏️ <b>Offene Tasks</b>")
        for t in other_open[:8]:
            due_d = _due_to_date(t.get("due"))
            if due_d is None:
                due_str = ""
            else:
                delta = (due_d - today_date).days
                if delta == 0:
                    due_str = " <i>(heute)</i>"
                elif delta == 1:
                    due_str = " <i>(morgen)</i>"
                elif 0 < delta <= 7:
                    due_str = f" <i>(in {delta}d)</i>"
                else:
                    due_str = f" <i>(bis {due_d.isoformat()})</i>"
            prio_emoji = PRIO_SYMBOLS.get(t.get("priority"), PRIO_SYMBOLS["medium"])
            parts.append(f"{prio_emoji} {_esc_html(str(t['title']))}{due_str}")
        if len(other_open) > 8:
            parts.append(f"<i>… und {len(other_open)-8} weitere</i>")

    # ─── Gestern: Abends-Reflexion ───
    yest_path = DAILY_DIR / f"{yesterday}.md"
    if yest_path.exists():
        try:
            ypost = frontmatter.load(yest_path)
            ybody = ypost.content or ""
            m = re.search(r"## Abends\s*\n(.*?)(?=\n## |\Z)", ybody, re.DOTALL)
            if m and m.group(1).strip():
                abends = m.group(1).strip()
                # Template-Standardtext rausfiltern
                template_lines = ("- Was lief gut?", "- Was nehme ich mit?")
                non_template = [l for l in abends.split("\n") if l.strip() and l.strip() not in template_lines]
                if non_template:
                    abends_clean = "\n".join(non_template)[:400]
                    parts.append(f"\n🌙 <b>Gestern Abends</b>\n<i>{_esc_html(abends_clean)}</i>")
        except Exception:
            pass

    # ─── Wenn nichts da: gentle morning ───
    has_anything = (data["overdue_tasks"] or other_open or data["reminders"]
                    or data["meetings"] or today_path.exists())
    if not has_anything:
        parts.append("\n<i>Heute steht noch nichts an.</i>")

    # ─── Vault-Health-Status (kompakt) ──
    # Liest den heutigen Health-Report (geschrieben um 02:00) für Briefing-Anhang.
    health_today = HEALTH_REPORTS_DIR / f"{today}.md"
    pending_actions = _load_pending_health_actions()
    if health_today.exists() or pending_actions:
        parts.append("\n🔧 <b>Vault-Pflege</b>")
        # Auto-Fix-Counter + Issue-Counter aus Report extrahieren
        # (regex schärft auf bekannte Issue-Headlines — sonst zählen wir
        # Auto-Fix-Subsections und Approval-Header doppelt)
        if health_today.exists():
            try:
                content = health_today.read_text(encoding="utf-8")
                m = re.search(r"##\s+✅\s+Auto-Fixed\s+\((\d+)\)", content)
                if m and int(m.group(1)) > 0:
                    parts.append(f"• {m.group(1)} Auto-Fixes über Nacht durchgeführt")
                # Issues-Counter: nur die offiziellen Issue-Headlines aus Schicht 1
                issue_headlines = (
                    "Lint-Issues", "Broken Wikilinks", "Orphans",
                    "Doppelte IDs", "Stale Uploads",
                )
                issue_count = 0
                for hl in issue_headlines:
                    m = re.search(rf"###\s+{re.escape(hl)}[^()]*\((\d+)\)", content)
                    if m:
                        issue_count += int(m.group(1))
                if issue_count > 0:
                    parts.append(f"• {issue_count} read-only Issues — siehe <code>{health_today.relative_to(VAULT)}</code>")
            except Exception:
                pass
        if pending_actions:
            parts.append(f"• {len(pending_actions)} Aktionen brauchen deine Approval:")
            for i, p in enumerate(pending_actions, 1):
                parts.append(f"  {i}. {_esc_html(p['summary'])}")
            parts.append("  <i>Antworte mit \"health 1\" / \"health 1 2\" / \"health 0\" (skip alle).</i>")

    # ─── Tagesplan-Frage am Ende — macht das Briefing zur Conversation ───
    parts.append(
        "\n━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "💬 <b>Was hast du sonst noch vor heute?</b>\n"
        "<i>Schreib einfach in Klartext: \"11 Uhr Bus, 14 Doc Müller, abends einkaufen\". "
        "Ich trage Termine als Reminders, To-Dos als Tasks ein.</i>"
    )

    return "\n".join(parts)


# ─── Recurring Tasks: Daily-Reset ──────────────────────────────────────────
# Tasks mit frontmatter.recurrence werden vom Reset-Job morgens reaktiviert,
# wenn sie done sind und das Pattern fällig ist. So lebt EINE Task-Datei für
# alle Wiederholungen, History sammelt sich im Log.

def _is_recurrence_due(pattern: str, last_completed: str, today: date) -> bool:
    """Bestimmt ob eine recurring Task heute reaktiviert werden soll.

    pattern: 'daily' | 'weekdays' | 'weekly' | 'monthly'
    last_completed: ISO-Datum (str) ODER date-Objekt (PyYAML kann unquoted
                    YYYY-MM-DD direkt zu date parsen). _due_to_date handlet beide.
    today: aktuelles Datum
    """
    last = _due_to_date(last_completed)
    if last is None:
        # Kein gültiges last_completed → nicht reaktivieren
        # (sicherer als raten — User soll erst einmal manuell done markieren)
        return False
    if last >= today:
        return False  # heute schon done oder in Zukunft (defensiv)

    if pattern == "daily":
        return True  # jeden Tag wieder
    if pattern == "weekdays":
        return today.weekday() < 5  # Mo-Fr (0-4)
    if pattern == "weekly":
        # Reaktivieren wenn ≥7 Tage seit letztem Done
        return (today - last).days >= 7
    if pattern == "monthly":
        # Reaktivieren wenn anderer Monat UND wir den last.day-Tag-of-Month
        # erreicht haben — bzw. am Monatsende falls last.day > Monatslänge
        # (sonst würde 31er-Task in Februar/April/... NIE reaktivieren).
        if today.month == last.month and today.year == last.year:
            return False
        # Letzter Tag des aktuellen Monats:
        if today.month == 12:
            next_first = date(today.year + 1, 1, 1)
        else:
            next_first = date(today.year, today.month + 1, 1)
        last_day_of_month = (next_first - timedelta(days=1)).day
        target_day = min(last.day, last_day_of_month)
        return today.day >= target_day
    return False


def reset_recurring_tasks() -> dict:
    """Walked alle Tasks, reaktiviert fällige recurring Tasks.

    Returns: dict mit Statistik {"checked": n, "reactivated": [slugs]}
    """
    if not TASKS_DIR.exists():
        return {"checked": 0, "reactivated": []}
    today = datetime.now(TIMEZONE).date()
    today_str = today.strftime("%Y-%m-%d")
    checked = 0
    reactivated = []

    for task_file, post in iter_vault_md(TASKS_DIR, recursive=False, skip_noise=False):
        checked += 1
        meta = post.metadata
        recurrence = meta.get("recurrence")
        if not recurrence or recurrence not in VALID_TASK_RECURRENCE:
            continue
        if meta.get("status") != "done":
            continue
        last_completed = meta.get("last_completed", "")
        if not _is_recurrence_due(recurrence, last_completed, today):
            continue

        # Reaktivieren
        post["status"] = "open"
        post["updated"] = today_str
        body = (post.content or "").rstrip() + f"\n- {today_str}: reaktiviert (recurring={recurrence})\n"
        post.content = body
        try:
            atomic_write(task_file, frontmatter.dumps(post) + "\n")
            slug = task_file.stem
            reactivated.append(slug)
            # Link in heutige Daily damit User es im Briefing sieht
            try:
                title = meta.get("title", slug)
                append_to_daily("Heute", f"- [ ] [[t-{slug}|{title}]] (wiederkehrend)")
            except Exception as e:
                log.warning(f"Daily-Link für reaktivierten Task fehlgeschlagen: {e}")
        except Exception as e:
            log.warning(f"recurring-reset: konnte {task_file.name} nicht schreiben: {e}")

    if reactivated:
        invalidate_today_data_cache()  # Tasks haben heute Status-Wechsel
    return {"checked": checked, "reactivated": reactivated}


async def recurring_task_reset_job(ctx: ContextTypes.DEFAULT_TYPE):
    """JobQueue-Callback — läuft täglich vor dem Briefing.

    Setzt fällige recurring Tasks zurück auf 'open'. Schreibt nichts an User
    (das Briefing-Job zeigt die reaktivierten Tasks ja in der Daily-Note).
    """
    try:
        stats = await asyncio.to_thread(reset_recurring_tasks)
        if stats["reactivated"]:
            log.info(f"recurring-reset: {len(stats['reactivated'])}/{stats['checked']} reaktiviert: {stats['reactivated']}")
        else:
            log.info(f"recurring-reset: {stats['checked']} Tasks geprüft, keine fällig")
    except Exception as e:
        log.exception(f"recurring_task_reset_job failed: {e}")


async def daily_briefing_job(ctx: ContextTypes.DEFAULT_TYPE):
    """JobQueue-Callback — wird täglich um BRIEFING_HOUR ausgeführt."""
    try:
        # Erst recurring Tasks reaktivieren, dann Briefing — so sieht User
        # die wiederkehrenden Tasks in der heutigen Daily.
        await asyncio.to_thread(reset_recurring_tasks)
        text = await asyncio.to_thread(compute_briefing)
        await ctx.bot.send_message(
            chat_id=ALLOWED_USER_ID,
            text=text,
            parse_mode=constants.ParseMode.HTML,
        )
        log.info(f"Daily briefing sent to {ALLOWED_USER_ID}")
    except Exception as e:
        log.exception(f"daily_briefing_job failed: {e}")
        # Versuche zumindest eine Fehler-Notification zu schicken
        try:
            await ctx.bot.send_message(
                chat_id=ALLOWED_USER_ID,
                text=f"⚠️ Daily-Briefing-Fehler: {type(e).__name__}\n<code>{_esc_html(str(e))[:300]}</code>",
                parse_mode=constants.ParseMode.HTML,
            )
        except Exception:
            pass


@require_auth
async def handle_briefing(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/briefing — manueller Trigger für das Morgens-Briefing."""
    text = await asyncio.to_thread(compute_briefing)
    await update.message.reply_text(text, parse_mode=constants.ParseMode.HTML)


@require_auth
async def handle_reminders(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/reminders — alle aktiven Erinnerungen listen."""
    text = await asyncio.to_thread(list_reminders)
    await safe_reply(update, text)


@require_auth
async def handle_backup(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/backup — manueller Push des Vaults ins konfigurierte GitHub-Repo."""
    await update.message.chat.send_action(constants.ChatAction.TYPING)
    await update.message.reply_text("⏳ Backup läuft...")
    result = await asyncio.to_thread(backup_vault)
    await safe_reply(update, result)


@require_auth
async def handle_reset(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """Setzt Conversation-Memory + pending Deletes zurück."""
    # Konsistent ALLOWED_USER_ID nutzen — ist single-user-bot
    await reset_history(ALLOWED_USER_ID)
    PENDING_DELETIONS.pop(ALLOWED_USER_ID, None)
    await update.message.reply_text("🔄 Memory + pending Deletes geleert. Frischer Anfang.")


@require_auth
async def handle_today(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/today — heutige Agenda (Tasks + Reminders + Meetings) plus Daily-Note."""
    agenda = await asyncio.to_thread(get_today_agenda)
    parts = [agenda]
    # Plus Daily-Note Body falls vorhanden — als Sekundär-Info
    path = DAILY_DIR / f"{today_iso()}.md"
    if path.exists():
        content = path.read_text(encoding="utf-8")
        body = re.sub(r"^---\n.*?\n---\n", "", content, count=1, flags=re.DOTALL).strip()
        if body and len(body) > 10:  # nicht nur Template-Reste
            parts.append("\n━━━━━━━━━━━━━━━━━━━━━━━━")
            parts.append("📓 **Heutige Daily-Notes:**\n")
            parts.append(body)
    await safe_reply(update, "\n".join(parts))


@require_auth
async def handle_tasks(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/tasks [filter] — offene Tasks. Optional Filter: today/tomorrow/week/overdue/nodate."""
    args = ctx.args if ctx.args else []
    when = args[0].strip().lower() if args else None
    if when and when not in ("today", "tomorrow", "week", "overdue", "nodate"):
        await update.message.reply_text(
            "Unbekannter Filter. Erlaubt: today, tomorrow, week, overdue, nodate (oder leer = alle gruppiert)."
        )
        return
    text = await asyncio.to_thread(list_open_tasks, when)
    await safe_reply(update, text)


@require_auth
async def handle_usage(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    """/usage [days] — Token-Usage + geschätzte Kosten der letzten N Tage (default 7)."""
    args = ctx.args if ctx.args else []
    days = 7
    if args:
        try:
            days = max(1, min(90, int(args[0])))
        except ValueError:
            pass
    text = await asyncio.to_thread(get_usage_summary, days)
    await safe_reply(update, text)


@require_auth
async def handle_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 <b>Vault-Assistent bereit</b>\n\n"
        "Standardmäßig <b>antworte ich nur</b> — speichere nichts ins Vault, "
        "außer du sagst's mir explizit.\n\n"
        "<b>Speicher-Verben</b> die ich erkenne:\n"
        "• \"speicher / merk dir / notiere / schreib auf …\"\n"
        "• \"task: …\" / \"todo: …\" / \"morgen X machen\"\n"
        "• \"meeting: …\" / \"war im termin mit …\"\n"
        "• \"X erledigt\" → markiert Task als done\n"
        "• \"lösche X\" → fragt nach Bestätigung\n\n"
        "<b>Multimedia:</b>\n"
        "🎤 Sprachnachricht → transkribiert + sortiert\n"
        "🖼 Foto → in 09_Attachments + Vision-Caption + OCR (Tesseract de+en)\n"
        "📄 .md/.txt → in 01_Raw/uploads/\n"
        "📑 .pdf → in 01_Raw/papers/ + Volltext-Extraktion + .md-Wrapper für Suche\n"
        "📝 .docx → in 01_Raw/uploads/ + Text+Tabellen-Extraktion + .md-Wrapper\n"
        "📎 sonstige Files → in 09_Attachments\n"
        "🔗 URL allein → fragt ob clippen\n\n"
        "<b>Commands:</b>\n"
        "/today — heutige Daily anzeigen\n"
        "/briefing — Tagesbriefing (überfällig + offen + heute)\n"
        "/reminders — alle aktiven Erinnerungen\n"
        "/backup — Vault in GitHub-Repo pushen\n"
        "/reset — Conversation-Memory leeren\n\n"
        "<i>3-Tier-Memory aktiv:\n"
        "• ~30 Turns aktiver Konversation (RAM)\n"
        "• gesamte History persistiert (überlebt Restart)\n"
        "• Long-term Fakten in 06_Meta/bot-memory/facts.md (immer im Kontext)\n"
        "Sag 'merk dir dass …' um persistente Fakten anzulegen.</i>",
        parse_mode=constants.ParseMode.HTML,
    )


# ============================================================================
# Main
# ============================================================================

def main():
    log.info(f"Vault: {VAULT}")
    if ALLOWED_USER_ID == 0:
        log.warning("⚠️  ALLOWED_USER_ID=0 → Setup-Modus aktiv. Erste Nachricht im Telegram triggert Anleitung.")
    else:
        log.info(f"Allowed user: {ALLOWED_USER_ID}")
    log.info(f"LLM: {LLM_MODEL} @ {LLM_BASE_URL}")
    log.info(f"Vision: {VISION_MODEL} @ {VISION_BASE_URL}")
    if not VAULT.exists():
        log.error(f"VAULT_PATH existiert nicht: {VAULT}")
        return
    if not TEMPLATES_DIR.exists():
        log.error(f"Templates-Ordner fehlt: {TEMPLATES_DIR}")
        return

    # ─── Tool-Konsistenz-Check: TOOLS-Schema ↔ TOOL_HANDLERS-Dispatch ───
    # Bei Drift wird der Bot beim Boot abgewiesen statt erst zur Laufzeit
    # mit "Tool nicht bekannt"-Errors zu enttäuschen.
    declared = {t["function"]["name"] for t in TOOLS if t.get("function", {}).get("name")}
    handled = set(TOOL_HANDLERS.keys())
    missing_handlers = declared - handled
    orphan_handlers = handled - declared
    if missing_handlers:
        log.error(f"❌ Tools im Schema OHNE Handler: {sorted(missing_handlers)}")
        return
    if orphan_handlers:
        log.warning(f"⚠️  Handler ohne Tool-Schema (für LLM unsichtbar): {sorted(orphan_handlers)}")
    log.info(f"Tools: {len(declared)} declared, {len(handled)} handled — alle aligned")
    app = Application.builder().token(TG_TOKEN).build()

    # Globale Referenz fuer Tools die JobQueue brauchen (z.B. create_reminder)
    global BOT_APP
    BOT_APP = app

    app.add_handler(CommandHandler("start", handle_start))
    app.add_handler(CommandHandler("today", handle_today))
    app.add_handler(CommandHandler("tasks", handle_tasks))
    app.add_handler(CommandHandler("usage", handle_usage))
    app.add_handler(CommandHandler("briefing", handle_briefing))
    app.add_handler(CommandHandler("reminders", handle_reminders))
    app.add_handler(CommandHandler("reset", handle_reset))
    app.add_handler(CommandHandler("backup", handle_backup))

    # Persistente Reminders aus JSON laden + neu schedulen
    persisted_reminders = _load_reminders()
    if persisted_reminders:
        cleaned = []
        for r in persisted_reminders:
            try:
                _schedule_reminder(app, r)
                # nur behalten wenn Reminder noch aktiv (einmalige in Vergangenheit
                # werden durch _schedule_reminder aus JSON entfernt — daher reload)
                cleaned.append(r)
            except Exception as e:
                log.warning(f"Reminder {r.get('id')} konnte nicht reactiviert werden: {e}")
        log.info(f"Reminders geladen: {len(cleaned)} aus {REMINDERS_FILE}")
    else:
        log.info("Keine persistierten Reminders.")

    # Daily-Briefing JobQueue (wenn BRIEFING_HOUR > 0 gesetzt)
    if BRIEFING_HOUR and ALLOWED_USER_ID > 0:
        try:
            briefing_time = dtime(hour=BRIEFING_HOUR, minute=0, tzinfo=TIMEZONE)
            app.job_queue.run_daily(
                daily_briefing_job,
                time=briefing_time,
                name="daily-briefing",
            )
            log.info(f"Daily-Briefing scheduled für {BRIEFING_HOUR}:00 {TIMEZONE.key}")
        except Exception as e:
            log.warning(f"JobQueue-Setup fehlgeschlagen (BRIEFING_HOUR={BRIEFING_HOUR}): {e}")
    else:
        # Briefing aus, aber recurring-Reset trotzdem schedulen — so funktionieren
        # wiederkehrende Tasks auch ohne Briefing. Default-Slot: 5:00 morgens.
        if ALLOWED_USER_ID > 0:
            try:
                reset_time = dtime(hour=5, minute=0, tzinfo=TIMEZONE)
                app.job_queue.run_daily(
                    recurring_task_reset_job,
                    time=reset_time,
                    name="recurring-task-reset",
                )
                log.info(f"Recurring-Task-Reset scheduled für 05:00 {TIMEZONE.key} (Briefing ist aus)")
            except Exception as e:
                log.warning(f"Recurring-Reset-Setup fehlgeschlagen: {e}")
        log.info("Daily-Briefing deaktiviert (BRIEFING_HOUR=0 oder Setup-Modus)")

    # Nightly Memory-Suggestion-Briefing
    if SUGGESTION_HOUR and ALLOWED_USER_ID > 0:
        try:
            sug_time = dtime(hour=SUGGESTION_HOUR, minute=0, tzinfo=TIMEZONE)
            app.job_queue.run_daily(
                nightly_suggestion_job,
                time=sug_time,
                name="nightly-memory-briefing",
            )
            log.info(f"Nightly-Memory-Vorschläge scheduled für {SUGGESTION_HOUR}:00 {TIMEZONE.key}")
        except Exception as e:
            log.warning(f"Suggestion-Job-Setup fehlgeschlagen: {e}")
    else:
        log.info("Nightly-Memory-Vorschläge deaktiviert (SUGGESTION_HOUR=0)")

    # Nightly Vault-Health-Check (immer auf 02:00 — vor Briefing/Memory/Reset)
    if ALLOWED_USER_ID > 0:
        try:
            health_time = dtime(hour=2, minute=0, tzinfo=TIMEZONE)
            app.job_queue.run_daily(
                nightly_health_job,
                time=health_time,
                name="nightly-health-check",
            )
            log.info(f"Nightly-Health-Check scheduled für 02:00 {TIMEZONE.key}")
        except Exception as e:
            log.warning(f"Health-Job-Setup fehlgeschlagen: {e}")
    app.add_handler(MessageHandler(filters.VOICE | filters.AUDIO, handle_voice))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    log.info("Polling started.")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
